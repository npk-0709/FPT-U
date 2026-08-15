"""Sinh file Word 'BaoCao_dbCOMPANY.docx' tu script DBC.sql.

Cach dung:  python build_docx.py
"""

import os
import sys

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

import docx_helpers as dh
import figures
import stats as stats_mod

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH = os.path.join(BASE_DIR, "BaoCao_dbCOMPANY.docx")
REPORT_DATE = "Tháng 8 năm 2026"


class Captions(object):
    """Danh so bang/hinh theo chuong va thu thap de tao danh muc."""

    def __init__(self):
        self._t, self._f = {}, {}
        self.tables, self.figures = [], []

    def table(self, ch, text):
        self._t[ch] = self._t.get(ch, 0) + 1
        label = "Bảng %d.%d. %s" % (ch, self._t[ch], text)
        self.tables.append(label)
        return label

    def fig(self, ch, text):
        self._f[ch] = self._f.get(ch, 0) + 1
        label = "Hình %d.%d. %s" % (ch, self._f[ch], text)
        self.figures.append(label)
        return label

    def raw_table(self, label):
        """Danh so thu cong, dung cho cac bang trong phu luc."""
        self.tables.append(label)
        return label


class Ctx(object):
    def __init__(self, doc, s, figs, cap):
        self.doc = doc
        self.s = s
        self.figs = figs
        self.cap = cap


def cover(doc):
    def line(text, size, bold=False, space=6, color=None, caps=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(text.upper() if caps else text)
        run.font.name = dh.BODY_FONT
        run.font.size = Pt(size)
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color
        return p

    line("TRƯỜNG ĐẠI HỌC FPT", 14, bold=True, space=2)
    line("Bộ môn Công nghệ Thông tin", 12.5, space=18)
    line("HỌC PHẦN DBI202 – DATABASE SYSTEMS", 12.5, bold=True, space=60)

    line("BÁO CÁO PHÂN TÍCH VÀ ĐÁNH GIÁ", 17, bold=True, space=4,
         color=dh.ACCENT)
    line("CƠ SỞ DỮ LIỆU dbCOMPANY", 20, bold=True, space=10, color=dh.ACCENT)
    line("(Phân tích trên script DBC.sql – Microsoft SQL Server)", 12.5,
         space=54)
    for p in doc.paragraphs[-1:]:
        for r in p.runs:
            r.italic = True

    info = [
        ("Sinh viên thực hiện", "............................................"),
        ("Mã số sinh viên", "............................................"),
        ("Lớp", "............................................"),
        ("Giảng viên hướng dẫn", "............................................"),
        ("Nguồn phân tích", "Script DBC.sql (cơ sở dữ liệu dbCOMPANY)"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = 1
    for k, v in info:
        cells = t.add_row().cells
        for idx, (txt, bold, align) in enumerate(
            ((k, True, "left"), (v, False, "left"))
        ):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(txt)
            run.font.name = dh.BODY_FONT
            run.font.size = Pt(12.5)
            run.bold = bold
        cells[0].width = Cm(5.2)
        cells[1].width = Cm(8.4)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    line(REPORT_DATE, 12.5, space=0)


def front_matter(doc, cap):
    dh.heading(doc, "MỤC LỤC", level=1, new_page=True)
    dh.note_box(
        doc,
        "Mục lục và số trang đã được cập nhật sẵn. Nếu bạn chỉnh sửa nội "
        "dung khiến số trang thay đổi, hãy bấm `Ctrl+A` rồi `F9` và chọn "
        "__Update entire table__ để Word điền lại.",
        title="Về mục lục",
    )
    dh.add_toc(doc)

    dh.heading(doc, "DANH MỤC BẢNG", level=1, new_page=True)
    marker_tables = dh.para(doc, "")
    dh.heading(doc, "DANH MỤC HÌNH", level=1)
    marker_figures = dh.para(doc, "")
    return marker_tables, marker_figures


def fill_index(marker, labels):
    for lab in labels:
        p = marker.insert_paragraph_before()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.left_indent = Cm(0.4)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        dh.add_runs(p, lab, size=11.5)
    marker._element.getparent().remove(marker._element)


def main():
    print("[1/4] Doc va phan tich DBC.sql ...")
    s = stats_mod.build()
    print("      %d bang, %d ban ghi du lieu mau." % (
        len(s["row_counts"]), s["total_rows"]))

    print("[2/4] Ve ERD va bieu do ...")
    figs = figures.build_all(s)

    print("[3/4] Soan noi dung Word ...")
    doc = dh.new_document()
    cap = Captions()
    ctx = Ctx(doc, s, figs, cap)

    cover(doc)
    marker_tables, marker_figures = front_matter(doc, cap)

    sec = dh.portrait_section(doc)
    dh.add_footer_page_numbers(sec)

    import content_part1
    import content_part2
    import content_part3
    import content_part4

    for part in (content_part1, content_part2, content_part3, content_part4):
        part.build(ctx)

    fill_index(marker_tables, cap.tables)
    fill_index(marker_figures, cap.figures)

    print("[4/4] Ghi file ...")
    doc.save(OUT_PATH)
    print("Hoan tat: %s" % OUT_PATH)
    print("  - %d bang, %d hinh." % (len(cap.tables), len(cap.figures)))
    return OUT_PATH


if __name__ == "__main__":
    sys.exit(0 if main() else 1)
