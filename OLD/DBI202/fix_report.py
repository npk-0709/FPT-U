"""
Script to add Lab 1 title heading and a professional Table of Contents
to final_report.docx
"""

import docx
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from lxml import etree
import sys, io

# Fix Windows console encoding
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

INPUT_FILE = r"final_report.docx"
OUTPUT_FILE = r"final_report.docx"

doc = docx.Document(INPUT_FILE)
body = doc.element.body

# Get the style ID for 'Heading 1' from existing heading
existing_h1 = doc.paragraphs[77]
h1_style_id = existing_h1.style.style_id
print(f"Heading 1 style ID: {h1_style_id}")

def make_heading1_para(text):
    """Create a Heading 1 paragraph XML element."""
    return parse_xml(
        f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'  <w:pPr><w:pStyle w:val="{h1_style_id}"/></w:pPr>'
        f'  <w:r><w:t>{text}</w:t></w:r>'
        f'</w:p>'
    )

def make_centered_heading1_para(text):
    """Create a centered Heading 1 paragraph XML element."""
    return parse_xml(
        f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'  <w:pPr>'
        f'    <w:pStyle w:val="{h1_style_id}"/>'
        f'    <w:jc w:val="center"/>'
        f'  </w:pPr>'
        f'  <w:r><w:t>{text}</w:t></w:r>'
        f'</w:p>'
    )

# == Step 1: Add "I. Lab 1 - Data Model Study" heading before paragraph 37 ==
target_para = doc.paragraphs[37]  # "1. Objective" of Lab 1
lab1_heading = make_heading1_para('I. Lab 1 - Data Model Study')
target_para._element.addprevious(lab1_heading)
print("[OK] Added Lab 1 title heading")

# == Step 2: Add Table of Contents after cover page ==
# Insert before paragraph 26 (after all Title-style cover paragraphs)
insert_before = doc.paragraphs[26]

# 2a. Add "Table of Contents" title
toc_title = make_centered_heading1_para('Table of Contents')
insert_before._element.addprevious(toc_title)

# 2b. Add TOC field (Word auto-TOC)
toc_sdt = parse_xml(
    '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '  <w:sdtPr>'
    '    <w:docPartObj>'
    '      <w:docPartGallery w:val="Table of Contents"/>'
    '      <w:docPartUnique/>'
    '    </w:docPartObj>'
    '  </w:sdtPr>'
    '  <w:sdtContent>'
    '    <w:p>'
    '      <w:r>'
    '        <w:fldChar w:fldCharType="begin"/>'
    '      </w:r>'
    '      <w:r>'
    '        <w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    '      </w:r>'
    '      <w:r>'
    '        <w:fldChar w:fldCharType="separate"/>'
    '      </w:r>'
    '      <w:r>'
    '        <w:rPr><w:noProof/></w:rPr>'
    '        <w:t>[Right-click here and select Update Field to generate Table of Contents]</w:t>'
    '      </w:r>'
    '      <w:r>'
    '        <w:fldChar w:fldCharType="end"/>'
    '      </w:r>'
    '    </w:p>'
    '  </w:sdtContent>'
    '</w:sdt>'
)
insert_before._element.addprevious(toc_sdt)

# 2c. Add page break after TOC
page_break = parse_xml(
    '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '  <w:r><w:br w:type="page"/></w:r>'
    '</w:p>'
)
insert_before._element.addprevious(page_break)

print("[OK] Added Table of Contents with page break")

# == Step 3: Save ==
doc.save(OUTPUT_FILE)
print(f"[OK] Saved to {OUTPUT_FILE}")
print()
print("IMPORTANT: Open the file in Microsoft Word, then:")
print("  1. Right-click on the Table of Contents area")
print("  2. Select 'Update Field'")
print("  3. Choose 'Update entire table'")
print("  This will populate the TOC with all headings and page numbers.")
