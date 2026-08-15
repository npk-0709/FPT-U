# -*- coding: utf-8 -*-
"""
Script to generate the Word (.docx) report for the
Football Club & Player Management project - J1.L.P0036 - LAB211 - FPT University.

Run:  python generate_report.py
(Requires: pip install python-docx)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_img_placeholder(doc, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[IMG: {caption}]\n")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 0, 0)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"(Insert image: {caption})")
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(128, 128, 128)


def add_code_block(doc, lines):
    for item in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(item)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    header_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2E74B5')
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(text)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(row.cells[c_idx], 'D6E4F0')
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table


def build_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FPT UNIVERSITY")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xF4, 0x7B, 0x20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SOFTWARE ENGINEERING DEPARTMENT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LAB REPORT")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LAB211 - BASIC JAVA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("J1.L.P0036 - Football Club & Player Management")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info_items = [
        ("Student", "Nguyen Phu Khuong"),
        ("Student ID", "SE203056"),
        ("Class", "SE06203"),
        ("Course", "LAB211 - Basic Java"),
        ("Topic", "32 - YouthPlayer Extension"),
        ("Instructor", "............................."),
        ("Semester", "Summer 2026"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run(f"{label}: ")
        run1.bold = True
        run1.font.size = Pt(13)
        run2 = p.add_run(value)
        run2.font.size = Pt(13)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ho Chi Minh City, June 2026")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_page_break()


def build_toc(doc):
    doc.add_heading('TABLE OF CONTENTS', level=1)
    toc_items = [
        ("1.", "Project Overview"),
        ("2.", "Class Diagram & UML"),
        ("  2.1.", "Class Diagram"),
        ("  2.2.", "Use Case Diagram"),
        ("3.", "Project Structure"),
        ("4.", "Class Descriptions"),
        ("  4.1.", "Model Layer"),
        ("  4.2.", "Business Layer"),
        ("  4.3.", "Dispatcher Layer"),
        ("  4.4.", "Tools Layer"),
        ("  4.5.", "Security Layer (Authorization)"),
        ("5.", "How to Run the Program"),
        ("6.", "Feature Usage Guide"),
        ("7.", "OOP Concepts Applied"),
        ("8.", "Conclusion"),
        ("A.", "APPENDIX A: Extension - Topic 32: YouthPlayer"),
        ("B.", "APPENDIX B: Authorization & Access Control (RBAC)"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run_num = p.add_run(f"{num} ")
        run_num.bold = True
        run_num.font.size = Pt(12)
        run_title = p.add_run(f"{title}")
        run_title.font.size = Pt(12)
    doc.add_page_break()


def build_overview(doc):
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        'The "Football Club & Player Management" project (J1.L.P0036) is a console-based '
        'system developed in Java using a 3-layer architecture. It provides full CRUD '
        'functionality for Clubs and Players, along with search, filter, sort, and file-based '
        'data persistence. The system also includes a Youth Player extension (Topic 32) and a '
        'role-based authorization layer that controls which features each user is allowed to use.'
    )

    doc.add_heading('Main Features:', level=2)
    add_styled_table(doc,
        ['#', 'Feature', 'Target'],
        [
            ['1', 'List all clubs', 'Club'],
            ['2', 'Add a new club', 'Club'],
            ['3', 'Search for a club by ID', 'Club'],
            ['4', 'Update a club by ID', 'Club'],
            ['5', 'List clubs with budget <= input value', 'Club'],
            ['6', 'List players sorted by club name, shirt number', 'Player'],
            ['7', 'Search players by partial name', 'Player'],
            ['8', 'Add a new player', 'Player'],
            ['9', 'Remove a player by ID', 'Player'],
            ['10', 'Update a player by ID', 'Player'],
            ['11', 'List players by position', 'Player'],
            ['12', 'List all youth players', 'YouthPlayer'],
            ['13', 'Add a new youth player', 'YouthPlayer'],
            ['14', 'Update a youth player by ID', 'YouthPlayer'],
            ['15', 'Remove a youth player by ID', 'YouthPlayer'],
            ['16', 'List youth players eligible for first team', 'YouthPlayer'],
            ['17', 'Save data to files', 'System'],
            ['18', 'Load data from files', 'System'],
            ['19', 'Quit', 'System'],
        ],
        col_widths=[1.2, 10, 4]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Note: After the authorization layer is added, the menu is no longer fixed. Each user '
        'logs in first and only sees the features permitted by their role (see Appendix B).'
    )

    doc.add_heading('Data Constraints:', level=2)
    add_styled_table(doc,
        ['Entity', 'Attribute', 'Constraint'],
        [
            ['Club', 'clubID', 'Format ^CL-\\d{4}$ (CL-0001..CL-9999), unique'],
            ['Club', 'name', 'Non-empty'],
            ['Club', 'sponsor', 'Non-empty'],
            ['Club', 'budget', 'Positive double (million EUR)'],
            ['Player', 'id', 'Format ^P\\d{4}$ (P0001..P9999), unique'],
            ['Player', 'clubID', 'Must exist in Club list (foreign key)'],
            ['Player', 'name', 'Non-empty'],
            ['Player', 'position', 'Goalkeeper / Defender / Midfielder / Forward / Winger'],
            ['Player', 'shirtNumber', '1-99, unique within the same Club'],
            ['YouthPlayer', 'id', 'Format ^AC-\\d{4}$, unique'],
            ['YouthPlayer', 'clubId', 'Must exist in Club list (foreign key)'],
            ['YouthPlayer', 'age', '8-21; age >= 18 suggests first-team promotion'],
        ],
        col_widths=[3, 3, 9]
    )
    doc.add_page_break()


def build_uml(doc):
    doc.add_heading('2. Class Diagram & UML', level=1)

    doc.add_heading('2.1. Class Diagram', level=2)
    doc.add_paragraph(
        'The Class Diagram below illustrates the relationships between all classes in the system, '
        'including the YouthPlayer extension and the authorization classes.'
    )
    add_img_placeholder(doc, "Class Diagram - System class relationships")

    doc.add_paragraph('Class Relationships:')
    add_styled_table(doc,
        ['Relationship', 'From', 'To', 'Type'],
        [
            ['Inheritance', 'Player', 'Person', 'extends'],
            ['Inheritance', 'YouthPlayer', 'Person', 'extends'],
            ['Implementation', 'Player', 'Comparable<Player>', 'implements'],
            ['Implementation', 'Club', 'Comparable<Club>', 'implements'],
            ['Implementation', 'Person', 'Serializable', 'implements'],
            ['Association', 'PlayersManager', 'List<Player>', 'has-a'],
            ['Association', 'ClubsManager', 'List<Club>', 'has-a'],
            ['Association', 'AuthManager', 'List<User>', 'has-a'],
            ['Association', 'User', 'Role', 'has-a'],
            ['Association', 'Role', 'Set<Permission>', 'has-a'],
            ['Association', 'MenuItem', 'Permission', 'has-a'],
            ['Dependency', 'Menu', 'AuthManager', 'uses'],
            ['Dependency', 'Menu', 'ClubsManager / PlayersManager', 'uses'],
        ],
        col_widths=[3.5, 3.5, 5, 3]
    )

    doc.add_heading('2.2. Use Case Diagram', level=2)
    doc.add_paragraph(
        'The Use Case Diagram shows all system functions from the user perspective. With '
        'authorization, actors are split by role: Admin (all use cases), Manager (player & youth '
        'management plus read-only clubs), and Viewer (read-only).'
    )
    add_img_placeholder(doc, "Use Case Diagram - Use cases grouped by role")
    doc.add_page_break()


def build_structure(doc):
    doc.add_heading('3. Project Structure', level=1)
    doc.add_paragraph('The project follows a 3-layer architecture with an additional security layer:')
    add_code_block(doc, [
        'FootballClub_n_PlayerManagement/',
        '|-- src/',
        '|   |-- model/                 <-- Data Layer',
        '|   |   |-- Person.java         (Abstract base class)',
        '|   |   |-- Player.java         (extends Person)',
        '|   |   |-- YouthPlayer.java    (extends Person)',
        '|   |   |-- Club.java           (Serializable, Comparable)',
        '|   |   |-- Validatable.java    (Interface - validation constants)',
        '|   |   |-- Permission.java     (enum - access permissions)   * NEW',
        '|   |   |-- Role.java           (enum - role -> permissions)  * NEW',
        '|   |   +-- User.java           (account: username/password/role) * NEW',
        '|   |-- business/              <-- Business Layer',
        '|   |   |-- ClubsManager.java',
        '|   |   |-- PlayersManager.java',
        '|   |   |-- YouthPlayersManager.java',
        '|   |   +-- AuthManager.java    (login + user store)          * NEW',
        '|   |-- dispatcher/            <-- Presentation Layer',
        '|   |   |-- Menu.java           (Main entry - permission-based menu)',
        '|   |   +-- MenuItem.java       (label + permission + action) * NEW',
        '|   +-- tools/                 <-- Utility',
        '|       +-- Inputter.java',
        '|-- clubs.txt',
        '|-- players.txt',
        '+-- youth_players.txt',
    ])
    doc.add_paragraph()
    doc.add_paragraph(
        'The 3-layer architecture ensures clear separation of concerns: Model (data) - Business '
        '(logic) - Dispatcher (presentation). Authorization classes live in model (Permission, '
        'Role, User) and business (AuthManager), keeping security concerns isolated.'
    )
    doc.add_page_break()


def build_class_descriptions(doc):
    doc.add_heading('4. Class Descriptions', level=1)

    doc.add_heading('4.1. Model Layer', level=2)
    doc.add_paragraph(
        'Person.java (Abstract Class): Abstract base class representing a person. Contains common '
        'attributes (id, name) and the abstract method getDisplayInfo() for subclass-specific '
        'display (Polymorphism).'
    )
    doc.add_paragraph(
        'Player.java (extends Person): Represents a football player. Adds clubID, position and '
        'shirtNumber. Implements Comparable for ordering by shirt number.'
    )
    doc.add_paragraph(
        'YouthPlayer.java (extends Person): Represents a youth academy player with clubId and age '
        '(8-21). Method isEligibleForFirstTeam() returns true when age >= 18.'
    )
    doc.add_paragraph(
        'Club.java: Represents a football club. Implements Serializable and Comparable (sorted by '
        'name, case-insensitive). setBudget() only accepts positive values (Encapsulation).'
    )
    doc.add_paragraph(
        'Validatable.java (Interface): Centralizes all validation rules in one place: ID regex '
        'patterns, age bounds, the set of valid positions, isValid() and isPosition().'
    )

    doc.add_heading('4.2. Business Layer', level=2)
    add_styled_table(doc,
        ['Class', 'Responsibility'],
        [
            ['ClubsManager', 'CRUD for Club, filter by budget, save/load with strict validation, dirty flag'],
            ['PlayersManager', 'CRUD for Player, sort via Comparator chaining, unique shirt check, save/load'],
            ['YouthPlayersManager', 'CRUD for YouthPlayer, first-team promotion suggestion, save/load'],
            ['AuthManager', 'Stores user accounts and performs login authentication (NEW)'],
        ],
        col_widths=[4.5, 10.5]
    )

    doc.add_heading('4.3. Dispatcher Layer', level=2)
    doc.add_paragraph(
        'Menu.java: Main entry point containing main(). It first requires the user to log in '
        'through AuthManager, then builds a list of MenuItem objects and displays only the items '
        'the logged-in user is permitted to run. Auto-saves on exit when changes are detected.'
    )
    doc.add_paragraph(
        'MenuItem.java (NEW): A small value object that bundles a menu label, the Permission it '
        'requires, and a Runnable action. This turns the menu into a data-driven, permission-aware '
        'structure instead of a large switch statement.'
    )

    doc.add_heading('4.4. Tools Layer', level=2)
    add_styled_table(doc,
        ['Method', 'Description'],
        [
            ['inputInt(prompt, min, max)', 'Input integer within [min, max] range'],
            ['inputLoop(prompt, regex, errMsg)', 'Input string matching a regex pattern'],
            ['inputNonEmpty(prompt)', 'Input non-empty string'],
            ['inputOptional(prompt)', 'Optional input (Enter to skip)'],
            ['inputPositiveDouble(prompt)', 'Input positive double'],
            ['inputYesNo(prompt)', 'Y/N confirmation'],
        ],
        col_widths=[6.5, 8.5]
    )

    doc.add_heading('4.5. Security Layer (Authorization)', level=2)
    doc.add_paragraph(
        'The security layer implements Role-Based Access Control (RBAC). It is fully described in '
        'Appendix B; the table below summarizes the participating types.'
    )
    add_styled_table(doc,
        ['Type', 'Kind', 'Responsibility'],
        [
            ['Permission', 'enum', 'Atomic capabilities (CLUB_VIEW, PLAYER_MANAGE, ...)'],
            ['Role', 'enum', 'Maps a role to an EnumSet of permissions; can(permission)'],
            ['User', 'class', 'Account (username, password, role); can() and matches()'],
            ['AuthManager', 'class', 'Holds users and performs login()'],
            ['MenuItem', 'class', 'Couples a feature with the permission it requires'],
        ],
        col_widths=[3, 2.5, 9.5]
    )
    doc.add_page_break()


def build_how_to_run(doc):
    doc.add_heading('5. How to Run the Program', level=1)

    doc.add_heading('5.1. System Requirements', level=2)
    for item in [
        'JDK 1.8 or higher (installed and JAVA_HOME configured)',
        'IDE: IntelliJ IDEA / Eclipse / NetBeans',
        'OS: Windows / macOS / Linux',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('5.2. Running with Command Line', level=2)
    doc.add_paragraph('Step 1: Open a terminal and navigate to the project directory.')
    doc.add_paragraph('Step 2: Compile all Java files:')
    add_code_block(doc, ['javac -encoding UTF-8 -d out -sourcepath src src/dispatcher/Menu.java'])
    doc.add_paragraph('Step 3: Run the program:')
    add_code_block(doc, ['java -cp out dispatcher.Menu'])

    doc.add_heading('5.3. Startup & Login Screen', level=2)
    doc.add_paragraph(
        'On startup the system asks the user to log in. After a successful login, data is loaded '
        'and the menu shows only the features the role allows.'
    )
    add_code_block(doc, [
        '=== European Elite League (EEL) Management System ===',
        '',
        '--- Please log in ---',
        'Available accounts:',
        '  admin   / admin123    (full access)',
        '  manager / manager123  (players & youth)',
        '  viewer  / viewer123   (read only)',
        '',
        'Username: admin',
        'Password: admin123',
        'Welcome, admin [ADMIN]',
        'Loading data...',
        'Load data successfully!',
    ])
    add_img_placeholder(doc, "Console login screen and welcome message")
    doc.add_page_break()


def build_feature_guide(doc):
    doc.add_heading('6. Feature Usage Guide', level=1)

    doc.add_heading('6.1. Club Management', level=2)
    for title, desc in [
        ('List all clubs', 'Displays all clubs in table format, sorted by club name (Comparable).'),
        ('Add a new club', 'Input Club ID (CL-xxxx), name, sponsor, budget; duplicates rejected.'),
        ('Search club by ID', 'Input Club ID; shows club info or "This club does not exist!".'),
        ('Update a club by ID', 'Update name/sponsor/budget; press Enter to keep a value.'),
        ('List clubs with budget <= input', 'Input a max budget; lists all clubs within it.'),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(f'{title}: ')
        r.bold = True
        p.add_run(desc)

    doc.add_heading('6.2. Player Management', level=2)
    for title, desc in [
        ('List players sorted by club, shirt', 'Sort by club name then shirt number (Comparator chaining).'),
        ('Search players by partial name', 'Lists players whose name contains the keyword (case-insensitive).'),
        ('Add a new player', 'Validate ID, clubId (FK), position, and unique shirt number within the club.'),
        ('Remove a player by ID', 'Removes a player or reports that the player does not exist.'),
        ('Update a player by ID', 'Update name/position/shirt; values re-validated when entered.'),
        ('List players by position', 'Lists all players in a chosen position.'),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(f'{title}: ')
        r.bold = True
        p.add_run(desc)

    doc.add_heading('6.3. Youth Player Management', level=2)
    for title, desc in [
        ('List all youth players', 'Lists every youth player sorted by name, with a total count.'),
        ('Add a youth player', 'Validate ID (AC-xxxx), clubId, name and age (8-21).'),
        ('Update a youth player', 'Update name/age; if age >= 18 the system suggests promotion.'),
        ('Remove a youth player', 'Removes a youth player after a Y/N confirmation.'),
        ('List eligible for first team', 'Lists youth players with age >= 18 (FIRST_TEAM_AGE).'),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(f'{title}: ')
        r.bold = True
        p.add_run(desc)

    doc.add_heading('6.4. System - Save / Load / Quit', level=2)
    for title, desc in [
        ('Save data to files', 'Saves Club, Player and YouthPlayer data to their .txt files.'),
        ('Load data from files', 'Reloads all data with strict per-line validation.'),
        ('Quit', 'Exits; if unsaved changes exist (dirty flag), auto-saves before exit.'),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(f'{title}: ')
        r.bold = True
        p.add_run(desc)
    doc.add_page_break()


def build_oop(doc):
    doc.add_heading('7. OOP Concepts Applied', level=1)
    add_styled_table(doc,
        ['OOP Concept', 'Application', 'Specific Example'],
        [
            ['Abstraction', 'Person is an abstract class defining a common interface',
             'abstract getDisplayInfo() - each subclass defines its own display'],
            ['Inheritance', 'Player and YouthPlayer reuse id and name from Person',
             'class Player extends Person; class YouthPlayer extends Person'],
            ['Polymorphism', 'Subclasses override getDisplayInfo() and compareTo()',
             'YouthPlayer adds a "FIRST TEAM ELIGIBLE" tag'],
            ['Encapsulation', 'Private fields accessed through getters/setters with validation',
             'Club.setBudget() only accepts positive values'],
            ['Comparable', 'Club and Player implement Comparable for natural sorting',
             'Club by name; Player by shirt number'],
            ['Enum + EnumSet', 'Role aggregates permissions in a type-safe set',
             'Role.ADMIN holds EnumSet.allOf(Permission.class)'],
            ['Composition', 'A MenuItem couples behaviour with the permission it needs',
             'new MenuItem("Add a new club", CLUB_MANAGE, clubs::add)'],
        ],
        col_widths=[3, 6, 6]
    )
    doc.add_page_break()


def build_conclusion(doc):
    doc.add_heading('8. Conclusion', level=1)

    doc.add_heading('8.1. Achievements', level=2)
    for r in [
        'Implemented all 19 CRUD features for Club, Player and YouthPlayer.',
        'Followed OOP principles: Abstraction, Inheritance, Polymorphism, Encapsulation.',
        'Strict validation: regex, range, unique constraint, foreign key.',
        'Flexible sorting via Comparable and Comparator chaining.',
        'File-based persistence with strict validation on load and auto-save on exit.',
        'Added a Role-Based Access Control layer so each user only sees permitted features.',
    ]:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading('8.2. Challenges & Solutions', level=2)
    add_styled_table(doc,
        ['Challenge', 'Solution'],
        [
            ['Sorting players by multiple criteria', 'Comparator chaining: byClubName.thenComparingInt()'],
            ['Unique shirt number within a club', 'shirtTakenInClub() iterates the list to detect duplicates'],
            ['Strict per-line file validation', 'loadStrict() checks field count, format, existence, duplicates'],
            ['Hiding features per user role', 'RBAC: filter MenuItem list by user.can(permission)'],
            ['Unicode on Windows (JDK 1.8)', 'Compile with -encoding UTF-8'],
        ],
        col_widths=[7, 8]
    )
    doc.add_page_break()


def build_appendix_youth(doc):
    doc.add_heading('APPENDIX A: Extension - Topic 32: YouthPlayer', level=1)
    doc.add_paragraph(
        'Topic 32 requires adding a Youth Player management feature. A Youth Player has an ID '
        '(AC-xxxx), belongs to a club, has a name and an age (8-21). If age >= 18 the system '
        'suggests promotion to the first team.'
    )

    doc.add_heading('A1. Test Cases from Topic', level=2)
    add_styled_table(doc,
        ['Input Age', 'Expected Result', 'Explanation'],
        [
            ['19', 'Valid + suggest first-team promotion', 'age >= 18'],
            ['7', 'Rejected', 'age < 8 (out of range [8, 21])'],
            ['22', 'Rejected', 'age > 21 (out of range [8, 21])'],
        ],
        col_widths=[3, 7, 5]
    )

    doc.add_heading('A2. System Changes', level=2)
    add_styled_table(doc,
        ['File', 'Change'],
        [
            ['model/YouthPlayer.java', 'New model extending Person; isEligibleForFirstTeam()'],
            ['business/YouthPlayersManager.java', 'Full CRUD + promotion suggestion + save/load'],
            ['youth_players.txt', 'Sample youth player data file'],
            ['model/Validatable.java', '+YOUTH_PLAYER_ID_REGEX, +YOUTH_MIN_AGE, +YOUTH_MAX_AGE, +FIRST_TEAM_AGE'],
            ['tools/Inputter.java', '+inputYesNo(prompt)'],
            ['dispatcher/Menu.java', 'Menu expanded from 14 to 19 features'],
        ],
        col_widths=[6, 9]
    )
    doc.add_page_break()


def build_appendix_auth(doc):
    doc.add_heading('APPENDIX B: Authorization & Access Control (RBAC)', level=1)

    doc.add_heading('B1. Overview', level=2)
    doc.add_paragraph(
        'This extension adds an authorization layer using the Role-Based Access Control (RBAC) '
        'model. Each user logs in with a username and password and is assigned a Role. A Role owns '
        'a set of Permissions, and every menu feature declares the Permission it requires. The '
        'menu then displays only the features the logged-in user is allowed to perform, so '
        'unauthorized actions are never even shown.'
    )
    doc.add_paragraph(
        'The design favours OOP principles: permissions are an enum (type-safe), a role aggregates '
        'permissions through composition, and the menu is data-driven via MenuItem objects rather '
        'than a hard-coded switch. Adding a feature or a role does not require touching the '
        'dispatch logic.'
    )

    doc.add_heading('B2. Roles & Permissions Matrix', level=2)
    doc.add_paragraph('Three roles are provided with the following permissions:')
    add_styled_table(doc,
        ['Permission', 'ADMIN', 'MANAGER', 'VIEWER'],
        [
            ['CLUB_VIEW', 'Yes', 'Yes', 'Yes'],
            ['CLUB_MANAGE', 'Yes', 'No', 'No'],
            ['PLAYER_VIEW', 'Yes', 'Yes', 'Yes'],
            ['PLAYER_MANAGE', 'Yes', 'Yes', 'No'],
            ['YOUTH_VIEW', 'Yes', 'Yes', 'Yes'],
            ['YOUTH_MANAGE', 'Yes', 'Yes', 'No'],
            ['DATA_PERSIST', 'Yes', 'Yes', 'No'],
        ],
        col_widths=[6, 3, 3, 3]
    )
    doc.add_paragraph()
    doc.add_paragraph('Default demo accounts:')
    add_styled_table(doc,
        ['Username', 'Password', 'Role', 'Access'],
        [
            ['admin', 'admin123', 'ADMIN', 'All 18 features'],
            ['manager', 'manager123', 'MANAGER', 'View clubs; manage players & youth; save/load'],
            ['viewer', 'viewer123', 'VIEWER', 'Read-only: list/search only'],
        ],
        col_widths=[3.5, 3.5, 3, 5]
    )

    doc.add_heading('B3. New & Modified Files', level=2)
    add_styled_table(doc,
        ['File', 'Type', 'Description'],
        [
            ['model/Permission.java', 'NEW', 'enum of atomic permissions'],
            ['model/Role.java', 'NEW', 'enum mapping each role to an EnumSet<Permission>; can()'],
            ['model/User.java', 'NEW', 'account holding username/password/role; can(), matches()'],
            ['business/AuthManager.java', 'NEW', 'user store + login() with up to 3 attempts'],
            ['dispatcher/MenuItem.java', 'NEW', 'label + required Permission + Runnable action'],
            ['dispatcher/Menu.java', 'MODIFIED', 'login first, dynamic menu filtered by permission'],
        ],
        col_widths=[5.5, 2.5, 7]
    )

    doc.add_heading('B4. Key Design Points', level=2)
    for item in [
        'Permission is an enum, giving type-safe, self-documenting capability names.',
        'Role stores permissions in an EnumSet and exposes can(Permission) - the single source of truth for access checks.',
        'User delegates can() to its Role, so callers never inspect role internals.',
        'MenuItem couples each feature with the permission it needs, enabling a uniform filter.',
        'Menu computes the allowed items once per session and renders a compact, renumbered menu.',
        'Because unauthorized items are hidden, there is no scattered "Access denied" handling.',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('B5. Sample Menus per Role', level=2)
    doc.add_paragraph('VIEWER sees read-only features only:')
    add_code_block(doc, [
        '===== FOOTBALL CLUB & PLAYER MANAGEMENT (EEL) =====',
        'Logged in as: viewer [VIEWER]',
        '---------------------------------------------------',
        ' 1. List all clubs',
        ' 2. Search for a club by ID',
        ' 3. List clubs with budget <= input value',
        ' 4. List players sorted by club name, shirt number',
        ' 5. Search players by partial name',
        ' 6. List players by position',
        ' 7. List all youth players',
        ' 8. List youth players eligible for first team',
        ' 9. Quit',
    ])
    doc.add_paragraph()
    doc.add_paragraph('ADMIN sees the full menu (all 18 features + Quit).')
    add_img_placeholder(doc, "Console menus for VIEWER vs ADMIN")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('--- End of Report ---')
    run.bold = True
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)


def create_report():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.5

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    build_cover(doc)
    build_toc(doc)
    build_overview(doc)
    build_uml(doc)
    build_structure(doc)
    build_class_descriptions(doc)
    build_how_to_run(doc)
    build_feature_guide(doc)
    build_oop(doc)
    build_conclusion(doc)
    build_appendix_youth(doc)
    build_appendix_auth(doc)

    output_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        'Report_FootballClub_n_PlayerManagement.docx'
    )
    doc.save(output_path)
    print(f"[OK] Report created: {output_path}")
    return output_path


if __name__ == '__main__':
    create_report()
