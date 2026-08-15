# -*- coding: utf-8 -*-
"""Generate BAO_CAO_MountainHiking.docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "BAO_CAO_MountainHiking.docx"


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return h


def add_para(doc, text, bold=False, italic=False, align=None, size=13):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_image_placeholder(doc, caption, note=""):
    """Empty box placeholder for user to insert screenshot."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F5F5F5")
    cell.width = Cm(15)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n[ CHÈN ẢNH TẠI ĐÂY ]\n\n{caption}")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    if note:
        add_para(doc, note, italic=True, size=11)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                r.font.size = Pt(11)
        set_cell_shading(hdr[i], "D9E2F3")
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    r.font.size = Pt(11)
    doc.add_paragraph()


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # ===== TRANG BÌA =====
    for _ in range(4):
        doc.add_paragraph()
    add_para(doc, "ĐẠI HỌC FPT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_para(doc, "MÔN: LAB211 – LẬP TRÌNH JAVA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    doc.add_paragraph()
    add_para(doc, "BÁO CÁO DỰ ÁN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_para(doc, "MOUNTAIN HIKING CHALLENGE REGISTRATION", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18)
    add_para(doc, "(Hệ thống quản lý đăng ký thử thách leo núi)", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    doc.add_paragraph()
    add_para(doc, "Mã đề tài: J1.L.P0027", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    add_para(doc, "Ngôn ngữ: Java (Console Application)", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "Sinh viên thực hiện: ....................................", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    add_para(doc, "MSSV: ....................................", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    add_para(doc, "Lớp: ....................................", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    doc.add_paragraph()
    add_para(doc, "Giảng viên hướng dẫn: ....................................", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "Hồ Chí Minh, tháng ..... năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    doc.add_page_break()

    # ===== MỤC LỤC =====
    add_heading(doc, "MỤC LỤC", 1)
    toc_items = [
        "LỜI MỞ ĐẦU",
        "CHƯƠNG 1. TỔNG QUAN DỰ ÁN",
        "CHƯƠNG 2. PHÂN TÍCH YÊU CẦU",
        "CHƯƠNG 3. THIẾT KẾ HỆ THỐNG",
        "CHƯƠNG 4. CÀI ĐẶT VÀ TRIỂN KHAI",
        "CHƯƠNG 5. KIỂM THỬ VÀ KẾT QUẢ",
        "KẾT LUẬN",
        "TÀI LIỆU THAM KHẢO",
    ]
    for i, item in enumerate(toc_items, 1):
        add_para(doc, f"{i}. {item}", size=13)
    add_para(doc, "(Cập nhật số trang sau khi hoàn thiện báo cáo)", italic=True, size=11)
    doc.add_page_break()

    # ===== LỜI MỞ ĐẦU =====
    add_heading(doc, "LỜI MỞ ĐẦU", 1)
    add_para(doc,
        "Trong bối cảnh FPT University tổ chức các hoạt động ngoại khóa, thử thách leo núi "
        "(Mountain Hiking Challenge) là sự kiện thu hút đông đảo sinh viên tham gia. Việc quản lý "
        "đăng ký, thông tin sinh viên, phí tham gia và tình nguyện viên hỗ trợ bằng phương pháp "
        "thủ công gây khó khăn trong kiểm soát dữ liệu, tìm kiếm và thống kê.")
    add_para(doc,
        "Dự án Mountain Hiking Challenge Registration được xây dựng bằng Java thuần (console application), "
        "áp dụng các nguyên lý lập trình hướng đối tượng (OOP) để quản lý toàn bộ quy trình: đăng nhập "
        "theo vai trò, đăng ký sinh viên, cập nhật/xóa/tìm kiếm/lọc dữ liệu, thống kê theo đỉnh núi, "
        "quản lý tình nguyện viên, quản lý tài khoản và lưu trữ dữ liệu.")
    add_para(doc,
        "Báo cáo trình bày từ phân tích yêu cầu, thiết kế kiến trúc, sơ đồ UML, thuật toán, "
        "cài đặt chi tiết đến kết quả kiểm thử — phản ánh đúng source code hiện tại của dự án.")
    doc.add_page_break()

    # ===== CHƯƠNG 1 =====
    add_heading(doc, "CHƯƠNG 1. TỔNG QUAN DỰ ÁN", 1)

    add_heading(doc, "1.1. Bối cảnh và mục tiêu", 2)
    add_para(doc,
        "Hệ thống giải quyết bài toán quản lý đăng ký leo núi cho sinh viên FPT University tại "
        "nhiều campus (SE, HE, DE, QE, CE). Mỗi sinh viên đăng ký tham gia tại một đỉnh núi cụ thể, "
        "đóng phí tham gia (có chính sách giảm giá theo nhà mạng điện thoại). Ngoài ra, hệ thống hỗ trợ "
        "quản lý tình nguyện viên phục vụ sự kiện và phân quyền truy cập theo vai trò người dùng.")
    add_para(doc, "Mục tiêu chính:", bold=True)
    goals = [
        "Xây dựng ứng dụng console Java quản lý đăng ký leo núi đầy đủ chức năng CRUD.",
        "Áp dụng OOP: kế thừa, đa hình, đóng gói, trừu tượng trong thiết kế class.",
        "Xác thực người dùng và phân quyền theo Role (ADMIN, STAFF, VIEWER).",
        "Lưu trữ dữ liệu bằng Serialization (.dat) và xuất báo cáo CSV.",
        "Đảm bảo validation đầu vào chặt chẽ bằng regex.",
    ]
    for g in goals:
        add_para(doc, f"• {g}")

    add_heading(doc, "1.2. Phạm vi dự án", 2)
    add_para(doc,
        "Dự án bao gồm 18 file Java trong thư mục src/, 3 file dữ liệu (MountainList.csv, "
        "registrations.dat/csv, volunteers.dat/csv, accounts.dat). Ứng dụng chạy trên console, "
        "không sử dụng GUI hay framework bên ngoài.")

    add_heading(doc, "1.3. Công nghệ sử dụng", 2)
    add_table(doc, ["Thành phần", "Mô tả"], [
        ["Ngôn ngữ", "Java (JDK 8+)"],
        ["Kiểu ứng dụng", "Console Application"],
        ["Lưu trữ", "Object Serialization (.dat), CSV (.csv)"],
        ["Validation", "Java Regex (Pattern.matches)"],
        ["Collection", "ArrayList, HashMap, Collections.sort()"],
        ["IDE", "IntelliJ IDEA / NetBeans / VS Code"],
    ])

    add_heading(doc, "1.4. Cấu trúc thư mục dự án", 2)
    add_table(doc, ["Thư mục / File", "Vai trò"], [
        ["src/Main.java", "Điểm vào, menu chính, điều phối logic"],
        ["src/Person.java, Student.java, Volunteer.java, Account.java", "Model – thực thể dữ liệu"],
        ["src/Students.java, Volunteers.java, Mountains.java, Accounts.java", "Manager – quản lý collection"],
        ["src/Statistics.java, StatisticalInfo.java", "Thống kê theo đỉnh núi"],
        ["src/Inputter.java, Acceptable.java", "Utility – nhập liệu và validation"],
        ["src/Role.java, Permission.java, MenuItem.java, Skill.java", "Enum và hỗ trợ phân quyền"],
        ["MountainList.csv", "Danh sách đỉnh núi (read-only)"],
        ["diagrams/*.drawio", "Sơ đồ Class Diagram và Sequence Diagram"],
    ])
    doc.add_page_break()

    # ===== CHƯƠNG 2 =====
    add_heading(doc, "CHƯƠNG 2. PHÂN TÍCH YÊU CẦU", 1)

    add_heading(doc, "2.1. Tác nhân (Actor)", 2)
    add_table(doc, ["Actor", "Mô tả"], [
        ["ADMIN", "Quản trị viên — toàn quyền truy cập mọi chức năng"],
        ["STAFF", "Nhân viên — quản lý đăng ký, tình nguyện viên, lưu dữ liệu"],
        ["VIEWER", "Người xem — chỉ xem danh sách và thống kê"],
    ])

    add_heading(doc, "2.2. Yêu cầu chức năng", 2)
    add_table(doc, ["STT", "Chức năng", "Mô tả", "Quyền"], [
        ["1", "Đăng nhập", "Xác thực username/password, tối đa 3 lần thử", "Tất cả"],
        ["2", "New Registration", "Đăng ký sinh viên mới leo núi", "ADMIN, STAFF"],
        ["3", "Update Registration", "Cập nhật thông tin (Enter giữ giá trị cũ)", "ADMIN, STAFF"],
        ["4", "Display Registered List", "Hiển thị bảng đã sắp xếp theo ID", "Tất cả (có quyền VIEW)"],
        ["5", "Delete Registration", "Xóa đăng ký (xác nhận Y/N)", "ADMIN"],
        ["6", "Search Participants", "Tìm theo ID hoặc tên (partial match)", "Tất cả (có quyền VIEW)"],
        ["7", "Filter by Campus", "Lọc theo mã campus SE/HE/DE/QE/CE", "Tất cả (có quyền VIEW)"],
        ["8", "Statistics", "Thống kê số lượng và tổng phí theo đỉnh núi", "Tất cả (có quyền VIEW)"],
        ["9", "Save Data to File", "Lưu .dat và xuất .csv", "ADMIN, STAFF"],
        ["10", "Volunteer Management", "Sub-menu CRUD + phân ca tình nguyện viên", "ADMIN, STAFF"],
        ["11", "Account Management", "Sub-menu quản lý tài khoản", "ADMIN"],
        ["12", "Exit", "Thoát, cảnh báo nếu chưa lưu", "Tất cả"],
    ])

    add_heading(doc, "2.3. Yêu cầu phi chức năng", 2)
    reqs = [
        "Validation đầu vào bằng regex — chương trình không crash khi nhập sai.",
        "Dữ liệu được serialize an toàn, hỗ trợ đọc lại khi khởi động.",
        "Menu động theo quyền — chỉ hiển thị chức năng user được phép.",
        "Dirty flag (isSaved) — cảnh báo trước khi thoát nếu chưa lưu.",
        "Code tuân thủ OOP: private fields, getter/setter, equals/hashCode, Comparable.",
    ]
    for r in reqs:
        add_para(doc, f"• {r}")

    add_heading(doc, "2.4. Sơ đồ Use Case", 2)
    add_para(doc,
        "Use Case mô tả tương tác giữa người dùng và hệ thống. Các use case chính bao gồm "
        "Login, quản lý đăng ký (CRUD + Search + Filter), Statistics, Save, Volunteer Management, "
        "Account Management và Exit. Mỗi use case nhập liệu đều «include» Validate Input.")
    add_image_placeholder(doc,
        "Hình 2.1 – Sơ đồ Use Case Diagram",
        "Gợi ý: Vẽ bằng draw.io hoặc StarUML. Actor: Admin, Staff, Viewer. "
        "Include relationship với «Validate Input».")

    add_heading(doc, "2.5. Quy tắc nghiệp vụ", 2)
    add_table(doc, ["Quy tắc", "Chi tiết"], [
        ["Student ID", "2 ký tự campus (SE/HE/DE/QE/CE) + 6 chữ số"],
        ["Phí đăng ký", "Mặc định 6,000,000 VND; giảm 35% nếu SĐT Viettel/VNPT → 3,900,000 VND"],
        ["Campus code", "Lấy 2 ký tự đầu của Student ID"],
        ["Volunteer ID", "VL + 3 chữ số (VD: VL001)"],
        ["Phân ca", "Tối đa maxShiftsPerDay (1-3); slot MEDIC yêu cầu skill MEDIC"],
        ["Tài khoản mặc định", "admin/staff/viewer — mật khẩu: 123456"],
    ])
    doc.add_page_break()

    # ===== CHƯƠNG 3 =====
    add_heading(doc, "CHƯƠNG 3. THIẾT KẾ HỆ THỐNG", 1)

    add_heading(doc, "3.1. Kiến trúc tổng thể", 2)
    add_para(doc,
        "Hệ thống được tổ chức theo mô hình phân lớp đơn giản phù hợp ứng dụng console:")
    layers = [
        ("Presentation Layer", "Main.java — menu engine, điều phối, xử lý I/O người dùng"),
        ("Business Layer", "Students, Volunteers, Mountains, Accounts, Statistics — CRUD, tìm kiếm, thống kê"),
        ("Model Layer", "Person, Student, Volunteer, Account, Mountain, Skill — thực thể dữ liệu"),
        ("Utility Layer", "Inputter, Acceptable — nhập liệu và validation"),
        ("Security Layer", "Role, Permission, MenuItem — xác thực và phân quyền"),
        ("Data Layer", "File .dat (Serialization), .csv (export), MountainList.csv (master data)"),
    ]
    add_table(doc, ["Lớp", "Thành phần"], layers)

    add_heading(doc, "3.2. Sơ đồ lớp (Class Diagram)", 2)
    add_para(doc,
        "Sơ đồ lớp mô tả 18 class/interface/enum và quan hệ giữa chúng. "
        "Thiết kế trung tâm là hệ phân cấp Person (abstract) với 3 lớp con: Student, Volunteer, Account. "
        "Các manager class (Students, Volunteers, Mountains, Accounts) kế thừa ArrayList để tận dụng "
        "Collection Framework. Statistics kế thừa HashMap<String, StatisticalInfo>.")
    add_para(doc,
        "File draw.io: diagrams/ClassDiagram.drawio — mở bằng https://app.diagrams.net, "
        "export PNG/SVG rồi chèn vào báo cáo.")
    add_image_placeholder(doc,
        "Hình 3.1 – Class Diagram (export từ diagrams/ClassDiagram.drawio)",
        "Cách export: Mở file .drawio → File → Export as → PNG (300 DPI) → chèn ảnh vào đây.")

    add_heading(doc, "3.3. Mô tả quan hệ OOP chính", 2)
    add_table(doc, ["Quan hệ", "Ví dụ trong dự án"], [
        ["Kế thừa (extends)", "Student extends Person; Students extends ArrayList<Student>"],
        ["Triển khai (implements)", "Person implements Serializable; Student implements Comparable<Student>"],
        ["Đa hình", "getDisplayInfo() — mỗi lớp con hiển thị khác nhau"],
        ["Đóng gói", "private phone, email; setPhone() tự tính lại tuitionFee"],
        ["Trừu tượng", "Person abstract class; Acceptable interface chứa regex constants"],
        ["Enum", "Skill (MEDIC, LOGISTIC, GUIDE_ASSIST); Role; Permission"],
    ])

    add_heading(doc, "3.4. Sơ đồ tuần tự (Sequence Diagram)", 2)

    add_heading(doc, "3.4.1. Sequence Diagram – Đăng nhập", 3)
    add_para(doc,
        "Luồng đăng nhập: User nhập username/password → Main gọi Accounts.login() → "
        "searchByUsername() → Account.authenticate(). Tối đa 3 lần thử. "
        "File: diagrams/SequenceDiagram_Login.drawio")
    add_image_placeholder(doc,
        "Hình 3.2 – Sequence Diagram: Đăng nhập (diagrams/SequenceDiagram_Login.drawio)")

    add_heading(doc, "3.4.2. Sequence Diagram – Đăng ký mới", 3)
    add_para(doc,
        "Luồng đăng ký: Validate ID → kiểm tra trùng → nhập thông tin → validate mountain code → "
        "Student.calculateFee() → tạo Student → Students.add() → markUnsaved(). "
        "File: diagrams/SequenceDiagram_NewRegistration.drawio")
    add_image_placeholder(doc,
        "Hình 3.3 – Sequence Diagram: New Registration (diagrams/SequenceDiagram_NewRegistration.drawio)")

    add_heading(doc, "3.5. Thuật toán tổng quát", 2)
    add_para(doc, "Thuật toán chương trình chính:", bold=True)
    algo = """1. Khởi tạo Mountains → đọc MountainList.csv
2. Khởi tạo Students → đọc registrations.dat (nếu có)
3. Khởi tạo Volunteers → đọc volunteers.dat (nếu có)
4. Khởi tạo Accounts → đọc accounts.dat (hoặc seed 3 tài khoản mặc định)
5. Gọi login() — nếu thất bại sau 3 lần → thoát
6. Lặp menu chính (runMenu):
   a. Lọc MenuItem theo quyền currentUser (visibleItems)
   b. Hiển thị menu + thông tin user đang đăng nhập
   c. Đọc lựa chọn hợp lệ (getMenuChoice)
   d. Thực thi chức năng tương ứng hoặc xử lý Exit
7. Khi Exit: kiểm tra isSaved → hỏi lưu trước khi thoát"""
    p = doc.add_paragraph()
    run = p.add_run(algo)
    run.font.name = "Courier New"
    run.font.size = Pt(10)

    add_heading(doc, "3.6. Cơ chế Menu Engine và phân quyền", 2)
    add_para(doc,
        "Main sử dụng generic menu engine (runMenu) kết hợp MenuItem và Permission. "
        "Mỗi MenuItem gắn một Permission (hoặc null cho sub-menu nội bộ). "
        "Phương thức isAllowedFor(Account) kiểm tra account.can(permission) thông qua Role. "
        "Role enum sử dụng EnumSet<Permission> để định nghĩa quyền: ADMIN = all, STAFF = CRUD trừ delete/account, "
        "VIEWER = chỉ xem.")

    add_heading(doc, "3.7. Cơ chế lưu trữ dữ liệu", 2)
    add_table(doc, ["File", "Cơ chế", "Mục đích"], [
        ["registrations.dat", "ObjectOutputStream", "Lưu runtime List<Student>"],
        ["registrations.csv", "BufferedWriter", "Export đọc được bằng Excel"],
        ["volunteers.dat / .csv", "Tương tự Students", "Quản lý tình nguyện viên"],
        ["accounts.dat", "Serialization", "Lưu danh sách tài khoản"],
        ["MountainList.csv", "BufferedReader", "Master data đỉnh núi (read-only)"],
    ])
    doc.add_page_break()

    # ===== CHƯƠNG 4 =====
    add_heading(doc, "CHƯƠNG 4. CÀI ĐẶT VÀ TRIỂN KHAI", 1)

    add_heading(doc, "4.1. Danh sách class và vai trò", 2)
    add_table(doc, ["Class", "Loại", "Vai trò"], [
        ["Person", "abstract class", "Lớp cha: id, name, getDisplayInfo()"],
        ["Student", "class", "Sinh viên đăng ký leo núi + tính phí"],
        ["Volunteer", "class", "Tình nguyện viên + phân ca"],
        ["Account", "class", "Tài khoản đăng nhập + phân quyền"],
        ["Mountain", "class", "Thông tin đỉnh núi"],
        ["Students", "manager", "CRUD, search, filter, save/load Student"],
        ["Volunteers", "manager", "CRUD, save/load Volunteer"],
        ["Mountains", "manager", "Đọc CSV, validate mountain code"],
        ["Accounts", "manager", "Login, CRUD Account, seed default"],
        ["Statistics", "class", "HashMap thống kê theo mountainCode"],
        ["StatisticalInfo", "class", "Đếm số SV và tổng phí mỗi đỉnh"],
        ["Inputter", "utility", "Scanner wrapper + validation loop"],
        ["Acceptable", "interface", "Regex patterns + isValid()"],
        ["Skill", "enum", "MEDIC, LOGISTIC, GUIDE_ASSIST"],
        ["Role", "enum", "ADMIN, STAFF, VIEWER + EnumSet permission"],
        ["Permission", "enum", "8 quyền chức năng"],
        ["MenuItem", "class", "Label + Permission + Runnable action"],
        ["Main", "entry point", "Điều phối toàn bộ chương trình"],
    ])

    add_heading(doc, "4.2. Chi tiết các class quan trọng", 2)

    add_heading(doc, "4.2.1. Person và hệ phân cấp", 3)
    add_para(doc,
        "Person là abstract class implements Serializable, chứa protected id và name. "
        "Phương thức abstract getDisplayInfo() buộc lớp con định nghĩa cách hiển thị. "
        "Student, Volunteer và Account đều kế thừa Person — thể hiện tái sử dụng code và đa hình.")

    add_heading(doc, "4.2.2. Student – Tính phí đăng ký", 3)
    add_para(doc,
        "Phương thức static calculateFee(phone) kiểm tra số điện thoại thuộc Viettel hoặc VNPT "
        "(qua Acceptable.VIETTEL_VALID / VNPT_VALID). Nếu đúng → phí = 6,000,000 × 0.65 = 3,900,000 VND. "
        "setPhone() tự động gọi calculateFee() — đảm bảo phí luôn đồng bộ khi cập nhật SĐT.")

    add_heading(doc, "4.2.3. Hệ thống phân quyền", 3)
    add_table(doc, ["Role", "Quyền"], [
        ["ADMIN", "Tất cả Permission (EnumSet.allOf)"],
        ["STAFF", "CREATE, UPDATE, VIEW, STATISTICS, SAVE, MANAGE_VOLUNTEER"],
        ["VIEWER", "VIEW_REGISTRATION, VIEW_STATISTICS"],
    ])

    add_heading(doc, "4.2.4. Volunteer – Phân ca", 3)
    add_para(doc,
        "assign() tăng shiftsToday nếu chưa vượt maxShiftsPerDay. hasSkillFor(Skill) kiểm tra "
        "kỹ năng: slot GENERAL (null) cho phép mọi volunteer; slot MEDIC yêu cầu skill MEDIC.")

    add_heading(doc, "4.2.5. Inputter và Acceptable", 3)
    add_para(doc,
        "Inputter giữ một Scanner duy nhất (singleton trong Main). inputAndLoop() lặp đến khi "
        "input khớp regex. inputAndLoopAllowEmpty() cho phép Enter để giữ giá trị cũ khi update. "
        "Acceptable tập trung mọi pattern regex — tránh rải rác validation trong code.")

    add_heading(doc, "4.3. Hướng dẫn biên dịch và chạy", 2)
    steps = [
        "Mở project trong IDE (IntelliJ IDEA / NetBeans).",
        "Đảm bảo MountainList.csv nằm cùng thư mục working directory với Main.class.",
        "Compile: javac -d out src/*.java",
        "Run: java -cp out Main",
        "Đăng nhập bằng admin/123456, staff/123456 hoặc viewer/123456.",
    ]
    for i, s in enumerate(steps, 1):
        add_para(doc, f"{i}. {s}")

    add_image_placeholder(doc,
        "Hình 4.1 – Màn hình cấu trúc project trong IDE",
        "Chụp ảnh cây thư mục project trong IntelliJ/NetBeans.")
    doc.add_page_break()

    # ===== CHƯƠNG 5 =====
    add_heading(doc, "CHƯƠNG 5. KIỂM THỬ VÀ KẾT QUẢ", 1)

    add_heading(doc, "5.1. Kịch bản kiểm thử", 2)
    add_table(doc, ["TC", "Chức năng", "Input mẫu", "Kết quả mong đợi"], [
        ["TC01", "Login thành công", "admin / 123456", "Welcome Administrator [ADMIN]"],
        ["TC02", "Login thất bại", "admin / wrong", "Invalid, còn 2 attempts"],
        ["TC03", "New Registration", "SE123456, hợp lệ", "Added, hiển thị phí"],
        ["TC04", "Trùng ID", "SE123456 lần 2", "Student ID already exists"],
        ["TC05", "Update", "Enter giữ giá trị cũ", "Cập nhật field đã nhập"],
        ["TC06", "Search by Name", "Nguyen", "Partial match, case-insensitive"],
        ["TC07", "Filter Campus", "SE", "Chỉ SV có ID bắt đầu SE"],
        ["TC08", "Statistics", "Có dữ liệu", "Bảng Code | Peak | Count | Cost"],
        ["TC09", "Save", "Chọn Save", "registrations.dat + .csv được tạo"],
        ["TC10", "Exit chưa save", "Thoát sau khi add", "Hỏi có muốn save không"],
        ["TC11", "Viewer role", "viewer login", "Không thấy menu Create/Delete"],
        ["TC12", "Assign Volunteer MEDIC", "VL không có MEDIC", "Cannot assign to MEDIC slot"],
    ])

    add_heading(doc, "5.2. Kết quả chạy chương trình", 2)
    add_para(doc, "Dưới đây là các màn hình cần chụp ảnh minh chứng kết quả kiểm thử:")

    screenshots = [
        ("Hình 5.1", "Màn hình đăng nhập thành công"),
        ("Hình 5.2", "Menu chính (đăng nhập bằng ADMIN — hiển thị đầy đủ chức năng)"),
        ("Hình 5.3", "Menu chính (đăng nhập bằng VIEWER — chỉ chức năng xem)"),
        ("Hình 5.4", "Chức năng New Registration — nhập liệu và kết quả"),
        ("Hình 5.5", "Chức năng Display Registered List — bảng dữ liệu"),
        ("Hình 5.6", "Chức năng Update Registration — Enter giữ giá trị cũ"),
        ("Hình 5.7", "Chức năng Search by Name — kết quả tìm kiếm"),
        ("Hình 5.8", "Chức năng Filter by Campus"),
        ("Hình 5.9", "Chức năng Statistics by Mountain Peak"),
        ("Hình 5.10", "Chức năng Save Data — thông báo lưu file thành công"),
        ("Hình 5.11", "Sub-menu Volunteer Management"),
        ("Hình 5.12", "Chức năng Assign Volunteer to Shift"),
        ("Hình 5.13", "Sub-menu Account Management"),
        ("Hình 5.14", "Chức năng Exit — cảnh báo chưa lưu dữ liệu"),
        ("Hình 5.15", "File registrations.csv mở bằng Excel"),
    ]
    for fig, desc in screenshots:
        add_image_placeholder(doc, f"{fig} – {desc}")

    add_heading(doc, "5.3. Đánh giá kết quả", 2)
    add_para(doc,
        "Hệ thống đáp ứng đầy đủ yêu cầu đề tài J1.L.P0027: quản lý đăng ký leo núi với validation "
        "chặt chẽ, thống kê, lưu file, mở rộng quản lý tình nguyện viên và phân quyền người dùng. "
        "Code áp dụng đúng các nguyên lý OOP và Collection Framework của Java.")
    doc.add_page_break()

    # ===== KẾT LUẬN =====
    add_heading(doc, "KẾT LUẬN", 1)
    add_para(doc,
        "Dự án Mountain Hiking Challenge Registration đã xây dựng thành công hệ thống quản lý "
        "đăng ký leo núi bằng Java console, với kiến trúc phân lớp rõ ràng và 18 class được thiết kế "
        "theo nguyên tắc OOP. Các chức năng CRUD, tìm kiếm, lọc, thống kê, lưu trữ file, quản lý "
        "tình nguyện viên và phân quyền đều hoạt động ổn định.")
    add_para(doc, "Hướng phát triển trong tương lai:", bold=True)
    future = [
        "Tách Main thành Controller riêng để giảm độ dài file.",
        "Chuyển từ extends ArrayList sang composition (List<T> bên trong).",
        "Thêm giao diện GUI (JavaFX/Swing).",
        "Sử dụng database (MySQL/SQLite) thay Serialization.",
        "Viết unit test (JUnit) cho các business logic quan trọng.",
    ]
    for f in future:
        add_para(doc, f"• {f}")

    # ===== TÀI LIỆU THAM KHẢO =====
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "FPT University – LAB211 Course Materials, Assignment J1.L.P0027.",
        "Oracle Java Documentation – Collections Framework, Serialization, Regex.",
        "Martin Fowler – UML Distilled (Class Diagram, Sequence Diagram).",
        "docs.oracle.com – java.io.Serializable, java.util.Comparable.",
        "draw.io (diagrams.net) – UML Diagram Editor.",
    ]
    for i, r in enumerate(refs, 1):
        add_para(doc, f"[{i}] {r}")

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build_report()
