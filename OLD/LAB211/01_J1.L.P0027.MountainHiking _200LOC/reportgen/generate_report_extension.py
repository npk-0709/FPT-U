# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_img_placeholder(doc, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[CHÈN ẢNH: {caption}]\n")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 0, 0)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"(Vui lòng chèn ảnh: {caption})")
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(128, 128, 128)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    header_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2E74B5')
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(text)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(row.cells[c_idx], 'D6E4F0')
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

def create_report_extension():
    doc = Document()

    # PAGE SETUP
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # STYLES
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.5

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        if level == 1:
            heading_style.font.size = Pt(16)
        elif level == 2:
            heading_style.font.size = Pt(14)
        elif level == 3:
            heading_style.font.size = Pt(13)

    # DOC START
    doc.add_heading('NỘI DUNG BỔ SUNG BÁO CÁO - PHẦN MỞ RỘNG TÍNH NĂNG', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run("Lưu ý: ")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run("File này chứa nội dung cần thêm vào file báo cáo hiện tại. Các vị trí cần chèn ảnh đã được đánh dấu rõ ràng.")

    # 9.
    doc.add_heading('9. Mở rộng tính năng – Volunteer Management', level=1)
    
    # 9.1
    doc.add_heading('9.1. Giới thiệu tính năng mới', level=2)
    doc.add_paragraph('Trong phiên bản mở rộng này, hệ thống Mountain Hiking Challenge Registration được bổ sung module Volunteer Management (Quản lý Tình nguyện viên). Mục tiêu của phần mở rộng:')
    
    reqs = [
        'Tái cấu trúc (Refactoring): Áp dụng nguyên tắc kế thừa (Inheritance) bằng cách tạo lớp trừu tượng Person làm lớp cha chung cho Student và Volunteer, giúp tái sử dụng mã nguồn và tuân thủ nguyên lý OOP.',
        'Thêm module Volunteer: Cho phép quản lý danh sách tình nguyện viên hỗ trợ các chuyến leo núi, bao gồm các chức năng CRUD và chức năng phân ca trực (Assign to Shift).',
        'Mở rộng Acceptable: Bổ sung các regex pattern mới để validate dữ liệu đầu vào cho Volunteer.',
        'Mở rộng Main Menu: Thêm menu option 9 (Volunteer Management) với sub-menu riêng.'
    ]
    for r in reqs:
        doc.add_paragraph(r, style='List Bullet')
        
    doc.add_paragraph('Các file mới được thêm:')
    add_styled_table(doc, ['STT', 'Tên file', 'Mô tả'], [
        ['1', 'Person.java', 'Abstract class – lớp cha chung'],
        ['2', 'Skill.java', 'Enum – danh sách kỹ năng của tình nguyện viên'],
        ['3', 'Volunteer.java', 'Class – đối tượng Tình nguyện viên'],
        ['4', 'Volunteers.java', 'Class – quản lý danh sách Tình nguyện viên']
    ], col_widths=[1.5, 4, 9])
    
    doc.add_paragraph('Các file đã sửa đổi:')
    add_styled_table(doc, ['STT', 'Tên file', 'Nội dung sửa đổi'], [
        ['1', 'Student.java', 'Refactor: kế thừa từ Person thay vì khai báo id, name riêng'],
        ['2', 'Acceptable.java', 'Thêm 3 regex pattern cho Volunteer'],
        ['3', 'Main.java', 'Thêm menu Volunteer Management + các phương thức xử lý']
    ], col_widths=[1.5, 4, 9])
    
    # 9.2
    doc.add_heading('9.2. Thiết kế hệ thống phân cấp lớp (Inheritance – Abstract Class)', level=2)
    doc.add_paragraph('Trước khi mở rộng, lớp Student có các thuộc tính id và name khai báo trực tiếp. Khi thêm đối tượng Volunteer (cũng cần id và name), việc trùng lặp code là không thể tránh khỏi nếu không refactor.')
    doc.add_paragraph('Giải pháp: Tạo lớp trừu tượng Person chứa các thuộc tính và phương thức chung, sau đó cả Student và Volunteer đều kế thừa từ Person.')
    
    add_img_placeholder(doc, "Class Diagram tổng quan hệ thống sau khi refactor – bao gồm tất cả các lớp")
    
    # 9.3
    doc.add_heading('9.3. Mô tả các lớp mới', level=2)
    
    doc.add_heading('9.3.1. Abstract Class – Person', level=3)
    doc.add_paragraph('File: Person.java | Dòng code: 36 dòng')
    add_styled_table(doc, ['Thành phần', 'Mô tả'], [
        ['abstract class Person', 'Lớp trừu tượng, không thể khởi tạo trực tiếp'],
        ['implements Serializable', 'Cho phép serialize/deserialize đối tượng để lưu vào file .dat'],
        ['protected String id, name', 'Thuộc tính chung, truy cập được từ lớp con'],
        ['abstract getDisplayInfo()', 'Phương thức trừu tượng – buộc các lớp con phải override']
    ], col_widths=[5, 9])
    
    doc.add_heading('9.3.2. Enum – Skill', level=3)
    doc.add_paragraph('File: Skill.java | Dòng code: 20 dòng')
    add_styled_table(doc, ['Giá trị', 'Ý nghĩa'], [
        ['MEDIC', 'Tình nguyện viên có kỹ năng y tế – sơ cứu'],
        ['LOGISTIC', 'Tình nguyện viên phụ trách hậu cần – vận chuyển'],
        ['GUIDE_ASSIST', 'Tình nguyện viên hỗ trợ hướng dẫn viên']
    ], col_widths=[4, 10])
    
    doc.add_heading('9.3.3. Class – Volunteer', level=3)
    doc.add_paragraph('File: Volunteer.java | Dòng code: 115 dòng')
    doc.add_paragraph('Được khai báo kế thừa từ Person (public class Volunteer extends Person implements Comparable<Volunteer>).')
    add_styled_table(doc, ['Thuộc tính', 'Kiểu', 'Mô tả'], [
        ['skill', 'Skill', 'Kỹ năng chuyên môn'],
        ['maxShiftsPerDay', 'int', 'Số ca trực tối đa mỗi ngày (1–3)'],
        ['shiftsToday', 'int', 'Số ca đã nhận trong ngày hiện tại']
    ], col_widths=[4, 3, 7])
    
    add_styled_table(doc, ['Phương thức', 'Mô tả'], [
        ['assign()', 'Phân ca trực: tăng shiftsToday lên 1 nếu chưa đạt max. Trả về true nếu thành công.'],
        ['hasSkillFor(Skill)', 'Kiểm tra volunteer có kỹ năng phù hợp không. slot GENERAL thì ai cũng phù hợp.'],
        ['getDisplayInfo()', 'Override từ Person – hiển thị thông tin dạng bảng có format.'],
        ['toCsv()', 'Chuyển đổi thông tin sang định dạng CSV.']
    ], col_widths=[4, 10])

    doc.add_heading('9.3.4. Class – Volunteers', level=3)
    doc.add_paragraph('File: Volunteers.java | Dòng code: 150 dòng')
    doc.add_paragraph('Kế thừa từ ArrayList<Volunteer> và quản lý danh sách Volunteer. Hỗ trợ các phương thức add, searchById, delete, showAll, readFromFile, saveToFile (lưu vào volunteers.dat và xuất CSV).')
    
    # 9.4
    doc.add_heading('9.4. Các thay đổi ở lớp cũ', level=2)
    
    doc.add_heading('9.4.1. Student (refactored)', level=3)
    add_styled_table(doc, ['Trước (cũ)', 'Sau (mới)'], [
        ['implements Serializable, Comparable<Student>', 'extends Person implements Comparable<Student>'],
        ['Khai báo private String id; private String name;', 'Kế thừa id và name từ Person (protected)'],
        ['Constructor: this.id = id; this.name = name;', 'Constructor: super(id, name);'],
        ['Không có getDisplayInfo()', 'Override getDisplayInfo() trả về toString()']
    ], col_widths=[7, 7])
    
    doc.add_heading('9.4.2. Acceptable (updated)', level=3)
    add_styled_table(doc, ['Pattern', 'Giải thích'], [
        ['VOLUNTEER_ID', 'Bắt đầu bằng "VL" (case-insensitive) + 3 chữ số'],
        ['VOLUNTEER_NAME_VALID', '3–30 ký tự chữ cái (hỗ trợ tiếng Việt) + khoảng trắng'],
        ['SHIFT_VALID', 'Chỉ chấp nhận giá trị 1, 2 hoặc 3']
    ], col_widths=[5, 9])
    
    doc.add_heading('9.4.3. Main (updated)', level=3)
    doc.add_paragraph('Thay đổi Menu chính: Mở rộng từ 9 option thành 10 option, thêm "9. Volunteer Management".')
    add_img_placeholder(doc, "Ảnh chụp Main Menu hiển thị đầy đủ 10 lựa chọn")
    
    doc.add_paragraph('Thêm Sub-menu Volunteer Management với 6 chức năng (Add, Display, Update, Assign, Delete, Back).')
    add_img_placeholder(doc, "Ảnh chụp Volunteer Management sub-menu")
    
    # 9.5
    doc.add_heading('9.5. Class Diagram cập nhật', level=2)
    add_img_placeholder(doc, "Class Diagram hoàn chỉnh (PlantUML hoặc Draw.io) – bao gồm tất cả 13 class/interface/enum")
    
    # 9.6
    doc.add_heading('9.6. Chức năng Volunteer Management', level=2)
    
    doc.add_heading('9.6.1. Add New Volunteer', level=3)
    add_img_placeholder(doc, "Ảnh chụp quá trình thêm Volunteer mới – thành công")
    add_img_placeholder(doc, "Ảnh chụp trường hợp nhập trùng ID – hiển thị thông báo lỗi")
    
    doc.add_heading('9.6.2. Display Volunteer List', level=3)
    add_img_placeholder(doc, "Ảnh chụp danh sách Volunteer hiển thị đầy đủ")
    add_img_placeholder(doc, "Ảnh chụp trường hợp danh sách trống")
    
    doc.add_heading('9.6.3. Update Volunteer', level=3)
    add_img_placeholder(doc, "Ảnh chụp quá trình Update Volunteer – hiển thị thông tin cũ và mới")
    
    doc.add_heading('9.6.4. Assign Volunteer to Shift', level=3)
    add_img_placeholder(doc, "Ảnh chụp Assign thành công")
    add_img_placeholder(doc, "Ảnh chụp trường hợp Over shift limit")
    add_img_placeholder(doc, "Ảnh chụp trường hợp không đủ skill cho MEDIC slot")
    
    doc.add_heading('9.6.5. Delete Volunteer', level=3)
    add_img_placeholder(doc, "Ảnh chụp quá trình xóa Volunteer – xác nhận Y")
    add_img_placeholder(doc, "Ảnh chụp trường hợp hủy xóa – nhấn N")
    
    # 9.7
    doc.add_heading('9.7. Tích hợp lưu trữ dữ liệu Volunteer', level=2)
    doc.add_paragraph('Hàm saveDataToFile() được cập nhật để lưu cả Student (registrations.dat) và Volunteer (volunteers.dat).')
    add_img_placeholder(doc, "Ảnh chụp kết quả Save Data – hiển thị cả Student và Volunteer đều lưu thành công")
    doc.add_paragraph('Exit Program: Kiểm tra cả hai nguồn dữ liệu. Cảnh báo nếu một trong hai chưa lưu.')
    add_img_placeholder(doc, "Ảnh chụp Exit khi có unsaved changes")
    
    # 9.8
    doc.add_heading('9.8. Kết luận phần mở rộng', level=2)
    doc.add_paragraph('Phần mở rộng Volunteer Management đã thành công trong việc:')
    kl = [
        'Áp dụng tính kế thừa (Inheritance): Tạo abstract class Person làm lớp cha.',
        'Áp dụng tính đa hình (Polymorphism): getDisplayInfo() được override ở mỗi lớp con.',
        'Áp dụng tính đóng gói (Encapsulation): Skill enum, Volunteers class quản lý danh sách.',
        'Tích hợp đầy đủ và hoạt động mượt mà với hệ thống Student cũ.'
    ]
    for k in kl:
        doc.add_paragraph(k, style='List Bullet')
        
    doc.add_paragraph('Tổng số dòng code sau mở rộng đạt 1,518 dòng với 13 file class/enum/interface.')

    # ============================================================
    # 10. PHÂN QUYỀN (RBAC)
    # ============================================================
    doc.add_page_break()
    doc.add_heading('10. Mở rộng tính năng – Phân quyền theo vai trò (RBAC)', level=1)

    # 10.1
    doc.add_heading('10.1. Giới thiệu tính năng phân quyền', level=2)
    doc.add_paragraph('Ở phiên bản mở rộng tiếp theo, hệ thống được bổ sung cơ chế Đăng nhập (Authentication) và Phân quyền theo vai trò (RBAC – Role-Based Access Control). Trước đây, bất kỳ ai chạy chương trình đều sử dụng được toàn bộ chức năng; nay mỗi người dùng phải đăng nhập và chỉ thấy/được dùng các chức năng phù hợp với vai trò của mình.')

    reqs = [
        'Authentication: Người dùng phải đăng nhập bằng username/password trước khi vào hệ thống (tối đa 3 lần thử).',
        'Authorization (RBAC): Mỗi tài khoản gắn với một vai trò (Role); mỗi vai trò sở hữu một tập quyền (Permission). Menu được sinh động – chỉ hiển thị các chức năng mà người dùng có quyền.',
        'Áp dụng OOP: Tách quyền (Permission) – vai trò (Role) – tài khoản (Account) thành các thành phần độc lập, mỗi lớp tự chịu trách nhiệm về nhiệm vụ của mình (Single Responsibility).',
        'Quản lý tài khoản: Bổ sung menu Account Management cho quản trị viên (ADMIN) để xem, thêm, xóa tài khoản.',
        'Refactor code sạch: Thay thế switch-case menu cồng kềnh bằng một "menu engine" dùng danh sách MenuItem gắn quyền.'
    ]
    for r in reqs:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_paragraph('Các file mới được thêm:')
    add_styled_table(doc, ['STT', 'Tên file', 'Số dòng', 'Mô tả'], [
        ['1', 'Permission.java', '20', 'Enum – liệt kê các quyền chức năng của hệ thống'],
        ['2', 'Role.java', '44', 'Enum – vai trò, mỗi vai trò gắn một tập Permission'],
        ['3', 'Account.java', '70', 'Class – tài khoản người dùng, kế thừa từ Person'],
        ['4', 'Accounts.java', '128', 'Class – quản lý danh sách tài khoản + xử lý đăng nhập'],
        ['5', 'MenuItem.java', '24', 'Class – một mục menu gắn với quyền và hành động'],
    ], col_widths=[1.5, 3.5, 2, 8])

    doc.add_paragraph('Các file đã sửa đổi:')
    add_styled_table(doc, ['STT', 'Tên file', 'Nội dung sửa đổi'], [
        ['1', 'Main.java', 'Thêm luồng đăng nhập, menu động theo quyền và menu Account Management; thay switch-case bằng menu engine'],
        ['2', 'Acceptable.java', 'Thêm 2 regex pattern: USERNAME_VALID và PASSWORD_VALID'],
        ['3', 'Person.java', 'Được tái sử dụng làm lớp cha cho Account (id = username, name = tên hiển thị)']
    ], col_widths=[1.5, 3.5, 10])

    # 10.2
    doc.add_heading('10.2. Mô hình phân quyền RBAC', level=2)
    doc.add_paragraph('RBAC (Role-Based Access Control) là mô hình kiểm soát truy cập dựa trên vai trò. Thay vì gán trực tiếp quyền cho từng người dùng, hệ thống gán: Người dùng → Vai trò → Tập quyền. Cách này giúp việc quản lý quyền tập trung, dễ mở rộng và đúng tinh thần OOP.')
    doc.add_paragraph('Quan hệ giữa các thành phần:')
    add_styled_table(doc, ['Thành phần', 'Vai trò trong mô hình'], [
        ['Account', 'Đại diện một người dùng đăng nhập; giữ tham chiếu tới một Role.'],
        ['Role', 'Đại diện một vai trò; nắm giữ một Set<Permission> (dùng EnumSet).'],
        ['Permission', 'Đại diện một quyền truy cập đến một nhóm chức năng cụ thể.'],
        ['MenuItem', 'Mỗi mục menu khai báo quyền cần có; tự kiểm tra account có được phép hay không.']
    ], col_widths=[4, 10])
    add_img_placeholder(doc, "Sơ đồ mô hình RBAC: Account → Role → Set<Permission>")

    # 10.3
    doc.add_heading('10.3. Mô tả các thành phần mới', level=2)

    doc.add_heading('10.3.1. Enum – Permission', level=3)
    doc.add_paragraph('File: Permission.java | Dòng code: 20 dòng')
    doc.add_paragraph('Liệt kê 8 quyền tương ứng các nhóm chức năng, mỗi quyền kèm mô tả (description):')
    add_styled_table(doc, ['Permission', 'Cho phép'], [
        ['CREATE_REGISTRATION', 'Thêm đăng ký mới'],
        ['UPDATE_REGISTRATION', 'Cập nhật thông tin đăng ký'],
        ['VIEW_REGISTRATION', 'Xem / tìm kiếm / lọc danh sách đăng ký'],
        ['DELETE_REGISTRATION', 'Xóa đăng ký'],
        ['VIEW_STATISTICS', 'Xem thống kê theo địa điểm'],
        ['SAVE_DATA', 'Lưu dữ liệu ra file'],
        ['MANAGE_VOLUNTEER', 'Quản lý tình nguyện viên'],
        ['MANAGE_ACCOUNT', 'Quản lý tài khoản (chỉ ADMIN)']
    ], col_widths=[5, 9])

    doc.add_heading('10.3.2. Enum – Role', level=3)
    doc.add_paragraph('File: Role.java | Dòng code: 44 dòng')
    doc.add_paragraph('Mỗi hằng số vai trò khởi tạo kèm một EnumSet<Permission>. Phương thức has(Permission) kiểm tra vai trò có chứa quyền đó không.')
    add_styled_table(doc, ['Vai trò', 'Phạm vi quyền'], [
        ['ADMIN', 'Toàn quyền (EnumSet.allOf) – mọi chức năng kể cả quản lý tài khoản'],
        ['STAFF', 'Tạo, cập nhật, xem, thống kê, lưu file và quản lý tình nguyện viên (không xóa, không quản lý tài khoản)'],
        ['VIEWER', 'Chỉ xem danh sách và xem thống kê (read-only)']
    ], col_widths=[3, 11])

    doc.add_heading('10.3.3. Class – Account', level=3)
    doc.add_paragraph('File: Account.java | Dòng code: 70 dòng')
    doc.add_paragraph('Kế thừa từ Person (public class Account extends Person), tái sử dụng id làm username và name làm tên hiển thị.')
    add_styled_table(doc, ['Thành phần', 'Mô tả'], [
        ['password (private)', 'Mật khẩu được đóng gói kín, KHÔNG có getter để tránh lộ.'],
        ['role (Role)', 'Vai trò của tài khoản.'],
        ['authenticate(pwd)', 'Tài khoản tự xác thực mật khẩu nhập vào (trả về true/false).'],
        ['can(Permission)', 'Ủy quyền cho Role kiểm tra: role.has(permission).'],
        ['getDisplayInfo()', 'Override từ Person – hiển thị username | name | role.']
    ], col_widths=[5, 9])

    doc.add_heading('10.3.4. Class – Accounts', level=3)
    doc.add_paragraph('File: Accounts.java | Dòng code: 128 dòng')
    doc.add_paragraph('Kế thừa ArrayList<Account>, quản lý danh sách tài khoản và lưu/đọc file accounts.dat (Serialization). Nếu file rỗng, hệ thống tự seed 3 tài khoản mặc định. Cung cấp login(username, password) trả về Account nếu hợp lệ, searchByUsername, add, delete và showAll.')

    doc.add_heading('10.3.5. Class – MenuItem', level=3)
    doc.add_paragraph('File: MenuItem.java | Dòng code: 24 dòng')
    doc.add_paragraph('Đóng gói một mục menu gồm: nhãn hiển thị (label), quyền yêu cầu (Permission) và hành động thực thi (Runnable). Phương thức isAllowedFor(account) trả về true khi mục không yêu cầu quyền hoặc account có quyền tương ứng – nhờ đó menu chỉ hiển thị các mục mà người dùng được phép.')

    # 10.4
    doc.add_heading('10.4. Các thay đổi ở lớp cũ', level=2)

    doc.add_heading('10.4.1. Main (refactored – menu engine + đăng nhập)', level=3)
    add_styled_table(doc, ['Trước (cũ)', 'Sau (mới)'], [
        ['Vào thẳng menu, không đăng nhập', 'Bắt buộc login() trước, tối đa 3 lần thử'],
        ['switch-case lớn (case 1..10)', 'runMenu() duyệt danh sách MenuItem, sinh menu động'],
        ['Menu cố định cho mọi người', 'Chỉ hiển thị mục mà currentUser có quyền (visibleItems)'],
        ['Không có quản lý tài khoản', 'Thêm menu Account Management (ADMIN)']
    ], col_widths=[7, 7])
    add_img_placeholder(doc, "Ảnh chụp màn hình đăng nhập (LOGIN)")

    doc.add_heading('10.4.2. Acceptable (updated)', level=3)
    add_styled_table(doc, ['Pattern', 'Giải thích'], [
        ['USERNAME_VALID', '3–20 ký tự gồm chữ cái, chữ số hoặc dấu gạch dưới'],
        ['PASSWORD_VALID', '6–20 ký tự, không chứa khoảng trắng']
    ], col_widths=[5, 9])

    # 10.5
    doc.add_heading('10.5. Ma trận phân quyền (Role × Permission)', level=2)
    doc.add_paragraph('Bảng dưới đây tổng hợp quyền của từng vai trò (✔ = được phép, ✘ = không):')
    add_styled_table(doc, ['Chức năng', 'ADMIN', 'STAFF', 'VIEWER'], [
        ['New Registration', '✔', '✔', '✘'],
        ['Update Registration', '✔', '✔', '✘'],
        ['View / Search / Filter', '✔', '✔', '✔'],
        ['Delete Registration', '✔', '✘', '✘'],
        ['Statistics', '✔', '✔', '✔'],
        ['Save Data', '✔', '✔', '✘'],
        ['Volunteer Management', '✔', '✔', '✘'],
        ['Account Management', '✔', '✘', '✘'],
    ], col_widths=[7, 2.3, 2.3, 2.3])

    # 10.6
    doc.add_heading('10.6. Tài khoản mặc định', level=2)
    doc.add_paragraph('Khi chạy lần đầu (file accounts.dat chưa tồn tại), hệ thống tạo sẵn 3 tài khoản để minh họa phân quyền:')
    add_styled_table(doc, ['Username', 'Password', 'Vai trò'], [
        ['admin', '123456', 'ADMIN'],
        ['staff', '123456', 'STAFF'],
        ['viewer', '123456', 'VIEWER'],
    ], col_widths=[4, 4, 4])

    # 10.7
    doc.add_heading('10.7. Minh họa các chức năng phân quyền', level=2)

    doc.add_heading('10.7.1. Đăng nhập', level=3)
    add_img_placeholder(doc, "Ảnh chụp đăng nhập thành công với tài khoản admin")
    add_img_placeholder(doc, "Ảnh chụp đăng nhập sai mật khẩu – hiển thị số lần thử còn lại")

    doc.add_heading('10.7.2. Menu hiển thị theo vai trò', level=3)
    add_img_placeholder(doc, "Ảnh chụp menu của ADMIN – hiển thị đầy đủ 10 chức năng + Exit")
    add_img_placeholder(doc, "Ảnh chụp menu của VIEWER – chỉ hiển thị các chức năng xem/thống kê")

    doc.add_heading('10.7.3. Quản lý tài khoản (Account Management)', level=3)
    add_img_placeholder(doc, "Ảnh chụp danh sách tài khoản")
    add_img_placeholder(doc, "Ảnh chụp thêm tài khoản mới và chọn vai trò")
    add_img_placeholder(doc, "Ảnh chụp xóa tài khoản (và trường hợp không cho xóa tài khoản đang đăng nhập)")

    # 10.8
    doc.add_heading('10.8. Kết luận phần phân quyền', level=2)
    doc.add_paragraph('Phần mở rộng phân quyền đã đạt được:')
    kl2 = [
        'Áp dụng RBAC chuẩn: tách bạch Permission – Role – Account, dễ thêm vai trò/quyền mới mà không sửa logic cũ (Open/Closed Principle).',
        'Đóng gói (Encapsulation): mật khẩu được giấu kín, việc xác thực và kiểm tra quyền do chính đối tượng đảm nhiệm.',
        'Kế thừa & đa hình: Account kế thừa Person và override getDisplayInfo().',
        'Code sạch hơn: menu engine dựa trên MenuItem loại bỏ switch-case dài, đồng thời tự ẩn các chức năng người dùng không có quyền.'
    ]
    for k in kl2:
        doc.add_paragraph(k, style='List Bullet')

    doc.add_paragraph('Sau khi bổ sung phân quyền, dự án có tổng cộng 18 file class/enum/interface với khoảng 1,844 dòng code.')

    # SAVE
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'Report_Extension_MountainHiking.docx'
    )
    doc.save(output_path)
    print(f"[OK] Báo cáo phần mở rộng đã được tạo tại: {output_path}")

if __name__ == '__main__':
    create_report_extension()
