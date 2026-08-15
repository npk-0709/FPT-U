# -*- coding: utf-8 -*-
"""
Script to generate a comprehensive Word (.docx) report for the
Mountain Hiking Registration project – LAB211 – FPT University
(English version)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_img_placeholder(doc, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[IMG: {caption}]\n")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 0, 0)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"(Insert image: {caption})")
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(128, 128, 128)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    header_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2E74B5')
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(text)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(row.cells[c_idx], 'D6E4F0')
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table


def create_report():
    doc = Document()

    # ===================== PAGE SETUP =====================
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # ===================== STYLES =====================
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.5

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # ===================== COVER PAGE =====================
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FPT UNIVERSITY")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xF4, 0x7B, 0x20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SOFTWARE ENGINEERING DEPARTMENT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LAB REPORT")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LAB211 - BASIC JAVA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("J1.L.P0027 - Mountain Hiking Challenge Registration")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info_items = [
        ("Student", "Nguyen Phu Khuong"),
        ("Student ID", "SE203056"),
        ("Class", "SE06203"),
        ("Course", "LAB211 - Basic Java"),
        ("Instructor", "............................."),
        ("Semester", "Summer 2026"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run(f"{label}: ")
        run1.bold = True
        run1.font.size = Pt(13)
        run2 = p.add_run(value)
        run2.font.size = Pt(13)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ho Chi Minh City, June 2026")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_page_break()

    # ===================== TABLE OF CONTENTS =====================
    doc.add_heading('TABLE OF CONTENTS', level=1)
    toc_items = [
        ("1.", "Project Introduction", "3"),
        ("  1.1.", "Problem Description", "3"),
        ("  1.2.", "Lab Requirements", "3"),
        ("  1.3.", "Technologies Used", "4"),
        ("2.", "System Design", "5"),
        ("  2.1.", "Project Structure", "5"),
        ("  2.2.", "Class Diagram", "5"),
        ("  2.3.", "Use Case Diagram", "6"),
        ("  2.4.", "Overall Flowchart", "6"),
        ("3.", "Detailed Class Descriptions", "7"),
        ("  3.1.", "Interface Acceptable", "7"),
        ("  3.2.", "Class Inputter", "8"),
        ("  3.3.", "Class Mountain", "9"),
        ("  3.4.", "Class Mountains", "9"),
        ("  3.5.", "Class Student", "10"),
        ("  3.6.", "Class Students", "11"),
        ("  3.7.", "Class StatisticalInfo", "12"),
        ("  3.8.", "Class Statistics", "13"),
        ("  3.9.", "Class Main", "13"),
        ("4.", "User Guide for Each Feature", "14"),
        ("  4.1.", "Main Menu", "14"),
        ("  4.2.", "Feature 1: New Registration", "15"),
        ("  4.3.", "Feature 2: Update Registration", "16"),
        ("  4.4.", "Feature 3: Display Registered List", "17"),
        ("  4.5.", "Feature 4: Delete Registration", "18"),
        ("  4.6.", "Feature 5: Search Participants", "19"),
        ("  4.7.", "Feature 6: Filter by Campus", "20"),
        ("  4.8.", "Feature 7: Statistics by Mountain", "21"),
        ("  4.9.", "Feature 8: Save Data to File", "22"),
        ("  4.10.", "Feature 9: Exit Program", "23"),
        ("5.", "Overall Workflow", "24"),
        ("  5.1.", "Program Overall Flowchart", "24"),
        ("  5.2.", "Flowchart per Feature", "25"),
        ("  5.3.", "Sequence Diagram - New Registration", "26"),
        ("6.", "Sample Data", "27"),
        ("  6.1.", "MountainList.csv File", "27"),
        ("  6.2.", "registrations.csv File", "27"),
        ("7.", "Data Validation Rules", "28"),
        ("8.", "Conclusion", "29"),
    ]
    for num, title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run_num = p.add_run(f"{num} ")
        run_num.bold = True
        run_num.font.size = Pt(12)
        run_title = p.add_run(f"{title}")
        run_title.font.size = Pt(12)
        tab_run = p.add_run(f"  ........ {page}")
        tab_run.font.size = Pt(12)
        tab_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # ===================== CHAPTER 1: PROJECT INTRODUCTION =====================
    doc.add_heading('1. Project Introduction', level=1)

    doc.add_heading('1.1. Problem Description', level=2)
    doc.add_paragraph(
        'The project "J1.L.P0027 - Mountain Hiking Challenge Registration" is a console-based application '
        'built in Java, designed to manage student registrations for the Mountain Hiking Challenge. '
        'The program allows an Operator to perform CRUD (Create, Read, Update, Delete) operations '
        'on the registration list, as well as search, filter by campus, generate statistics by mountain peak, '
        'and save/load data to/from files.'
    )
    doc.add_paragraph(
        'This application was developed as part of the LAB211 - Basic Java course at FPT University, '
        'aimed at practicing Object-Oriented Programming (OOP), file I/O handling, '
        'input data validation, and organizing source code with a clear class structure.'
    )

    doc.add_heading('1.2. Lab Requirements', level=2)
    doc.add_paragraph('According to the LAB211 assignment, students must fulfill the following requirements:')
    requirements = [
        'Build a Java console application to manage mountain hiking registrations for FPT students.',
        'The menu consists of 9 features: New Registration, Update, Display, Delete, Search, Filter, Statistics, Save, Exit.',
        'Apply OOP: use classes, interfaces, extend ArrayList, implement Serializable and Comparable.',
        'Validate input data using Regex: Student ID, Phone, Email, Name, Campus Code.',
        'Calculate registration fee: default 6,000,000 VND, 35% discount for Viettel or VNPT carriers.',
        'Read the mountain list from the MountainList.csv file.',
        'Save registration data to registrations.dat (binary) and export to registrations.csv.',
        'Prompt for confirmation when exiting if there is unsaved data.',
        'Code must be at least 200 LOC (Lines of Code).',
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('1.3. Technologies Used', level=2)
    add_styled_table(doc,
        ['No.', 'Technology', 'Description'],
        [
            ['1', 'Java SE', 'Primary programming language (Java 8+)'],
            ['2', 'Java I/O', 'Read/write CSV files, serialization (.dat)'],
            ['3', 'Java Collections', 'ArrayList, HashMap, Collections.sort()'],
            ['4', 'Regex (Pattern)', 'Input data format validation'],
            ['5', 'Serialization', 'Store Student objects to binary file'],
            ['6', 'IntelliJ IDEA', 'Development IDE'],
        ],
        col_widths=[1.5, 4, 10]
    )

    doc.add_page_break()

    # ===================== CHAPTER 2: SYSTEM DESIGN =====================
    doc.add_heading('2. System Design', level=1)

    doc.add_heading('2.1. Project Structure', level=2)
    doc.add_paragraph('The project is organized with the following directory structure:')
    structure_items = [
        'Project1_MountainHiking/',
        '|-- src/',
        '|   |-- Main.java              (Main class, menu and orchestration)',
        '|   |-- Acceptable.java         (Interface with Regex constants)',
        '|   |-- Inputter.java           (Console input handling class)',
        '|   |-- Mountain.java           (Mountain entity class)',
        '|   |-- Mountains.java          (Mountain list, reads from CSV)',
        '|   |-- Student.java            (Student entity class)',
        '|   |-- Students.java           (Student list, CRUD, file I/O)',
        '|   |-- StatisticalInfo.java     (Statistics for one mountain peak)',
        '|   |-- Statistics.java          (Aggregated statistics table)',
        '|-- MountainList.csv            (Data: 13 mountain peaks)',
        '|-- registrations.dat           (Binary file for registrations)',
        '|-- registrations.csv           (CSV export of registrations)',
        '|-- Class Diagram.drawio        (Class diagram)',
        '|-- FlowChart.drawio            (Overall flowchart)',
        '|-- FlowChart_PerFunc.drawio    (Flowchart per feature)',
    ]
    for item in structure_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(item)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

    doc.add_paragraph()

    doc.add_heading('2.2. Class Diagram', level=2)
    doc.add_paragraph(
        'Below is the Class Diagram illustrating the relationships between classes in the system. '
        'The classes are designed following OOP principles with inheritance, interface implementation, '
        'and references between objects.'
    )
    add_img_placeholder(doc, "Class Diagram - System class relationships")

    doc.add_paragraph('Key relationships in the Class Diagram:')
    relationships = [
        'Student implements Serializable: enables serializing Student objects for .dat file storage.',
        'Student implements Comparable<Student>: enables comparing and sorting students by ID.',
        'Students extends ArrayList<Student>: the student list inherits from ArrayList.',
        'Mountains extends ArrayList<Mountain>: the mountain list inherits from ArrayList.',
        'Statistics extends HashMap<String, StatisticalInfo>: the statistics table uses HashMap.',
        'Inputter uses Acceptable: calls isValid() method for data validation.',
        'Students references Student; Statistics references StatisticalInfo.',
    ]
    for rel in relationships:
        doc.add_paragraph(rel, style='List Bullet')

    doc.add_heading('2.3. Use Case Diagram', level=2)
    doc.add_paragraph(
        'The Use Case Diagram describes the functionalities available to the Operator. '
        'All input-related features include a Validate Input step.'
    )
    add_img_placeholder(doc, "Use Case Diagram - System use cases")

    doc.add_heading('2.4. Overall Flowchart', level=2)
    doc.add_paragraph(
        'The overall flowchart describes the main workflow of the program from startup to exit. '
        'The program operates in a menu-driven loop.'
    )
    add_img_placeholder(doc, "Overall Flowchart - Main program workflow")

    doc.add_paragraph('Workflow description:')
    flow_steps = [
        'Step 1: Startup - Load the mountain list from MountainList.csv.',
        'Step 2: Load registration data from registrations.dat (if it exists).',
        'Step 3: Display the main menu with 9 options.',
        'Step 4: The user enters a choice (validated 1-9).',
        'Step 5: Execute the corresponding feature.',
        'Step 6: Mark data as unsaved if any changes were made (Add/Update/Delete).',
        'Step 7: Return to Step 3 until the user chooses Exit (9).',
        'Step 8: If there is unsaved data, prompt to save -> End.',
    ]
    for step in flow_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_page_break()

    # ===================== CHAPTER 3: DETAILED CLASS DESCRIPTIONS =====================
    doc.add_heading('3. Detailed Class Descriptions', level=1)

    # --- 3.1. Acceptable ---
    doc.add_heading('3.1. Interface Acceptable', level=2)
    doc.add_paragraph(
        'Acceptable is an interface that contains Regex (Regular Expression) constants '
        'used for input data validation, and a static method isValid() '
        'to check whether a string matches a given regex pattern.'
    )
    add_styled_table(doc,
        ['Constant', 'Regex Pattern', 'Meaning'],
        [
            ['STUDENT_ID', '^(?i)(SE|HE|DE|QE|CE)\\d{6}$', 'Student ID: 2-char campus code + 6 digits'],
            ['CAMPUS_CODE', '^(?i)(SE|HE|DE|QE|CE)$', 'Campus code: SE/HE/DE/QE/CE'],
            ['NAME_VALID', '^[A-Za-zAa-yy\\s]{2,20}$', 'Name: 2-20 letters and spaces only'],
            ['PHONE_VALID', '^0\\d{9}$', 'Phone: starts with 0, exactly 10 digits'],
            ['VIETTEL_VALID', '^(032|033|...|086)\\d{7}$', 'Viettel carrier phone (discounted fee)'],
            ['VNPT_VALID', '^(081|082|...|094)\\d{7}$', 'VNPT carrier phone (discounted fee)'],
            ['EMAIL_VALID', '^[A-Za-z0-9+_.-]+@...+\\.[A-Za-z]{2,}$', 'Valid email format'],
            ['YES_NO_VALID', '^[YyNn]$', 'Confirmation: Y or N'],
        ],
        col_widths=[3.5, 5.5, 6]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Method isValid(String data, String pattern): Uses Pattern.matches() '
        'to check whether the data string matches the pattern. Returns true if it matches, false otherwise.'
    )

    # --- 3.2. Inputter ---
    doc.add_heading('3.2. Class Inputter', level=2)
    doc.add_paragraph(
        'Inputter is a utility class for handling keyboard (console) input. '
        'This class ensures all input data is validated before being returned.'
    )
    add_styled_table(doc,
        ['Method', 'Description', 'Notes'],
        [
            ['getString(message)', 'Displays a prompt and reads a string', 'Returns trimmed string'],
            ['getInt(message)', 'Reads an integer, loops until valid', 'Uses INTEGER_VALID regex'],
            ['getDouble(message)', 'Reads a double, loops until valid', 'Uses DOUBLE_VALID regex'],
            ['getMenuChoice(msg, min, max)', 'Reads a menu choice within [min, max]', 'Loops if out of range'],
            ['inputAndLoop(msg, pattern)', 'Reads input and validates with regex', 'Does not allow empty input'],
            ['inputAndLoopAllowEmpty(msg, pattern)', 'Same as above but allows empty (Enter)', 'Used during Update'],
            ['confirmYesNo(message)', 'Asks for Y/N confirmation', 'Returns true if Y'],
        ],
        col_widths=[5.5, 5, 4.5]
    )

    # --- 3.3. Mountain ---
    doc.add_heading('3.3. Class Mountain', level=2)
    doc.add_paragraph(
        'Mountain represents a mountain peak in the program. '
        'Each Mountain object stores the mountain code, name, province, and description.'
    )
    add_styled_table(doc,
        ['Attribute', 'Data Type', 'Description'],
        [
            ['mountainCode', 'String', 'Mountain code (e.g., 1, 2, 3...)'],
            ['mountain', 'String', 'Mountain name (e.g., Ham Rong Mountain)'],
            ['province', 'String', 'Province/city where the mountain is located'],
            ['description', 'String', 'Detailed description of the mountain'],
        ],
        col_widths=[3.5, 3, 8.5]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Special methods: toString() displays information in columnar format, '
        'equals() compares by mountainCode (case-insensitive), '
        'hashCode() is based on the uppercase mountainCode.'
    )

    # --- 3.4. Mountains ---
    doc.add_heading('3.4. Class Mountains', level=2)
    doc.add_paragraph(
        'Mountains extends ArrayList<Mountain> and manages the entire list of mountain peaks. '
        'It automatically reads data from MountainList.csv upon initialization.'
    )
    add_styled_table(doc,
        ['Method', 'Description'],
        [
            ['Mountains()', 'Default constructor, reads MountainList.csv'],
            ['Mountains(pathFile)', 'Constructor with custom file path'],
            ['get(mountainCode)', 'Finds a Mountain by code (case-insensitive)'],
            ['isValidMountainCode(code)', 'Checks whether a mountain code exists'],
            ['dataToObject(text)', 'Converts a CSV line into a Mountain object'],
            ['readFromFile()', 'Reads the entire CSV file, skips header and blank lines'],
            ['showAll()', 'Displays the mountain list in a table with header'],
        ],
        col_widths=[5, 10]
    )

    # --- 3.5. Student ---
    doc.add_heading('3.5. Class Student', level=2)
    doc.add_paragraph(
        'Student represents a student who registered for the hiking challenge. '
        'This class implements Serializable (for binary file storage) and Comparable<Student> (for sorting by ID).'
    )
    add_styled_table(doc,
        ['Attribute', 'Data Type', 'Description'],
        [
            ['id', 'String', 'Student ID (e.g., SE203056)'],
            ['name', 'String', 'Student full name'],
            ['phone', 'String', 'Phone number (10 digits, starts with 0)'],
            ['email', 'String', 'Email address'],
            ['mountainCode', 'String', 'Registered mountain peak code'],
            ['tuitionFee', 'double', 'Registration fee (VND)'],
        ],
        col_widths=[3, 3, 9]
    )
    doc.add_paragraph()
    doc.add_paragraph('Important constants:')
    doc.add_paragraph('DEFAULT_FEE = 6,000,000 VND - Default registration fee.', style='List Bullet')
    doc.add_paragraph('DISCOUNT_RATE = 0.35 (35%) - Discount rate for Viettel/VNPT phone numbers.', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Fee calculation logic (calculateFee):')
    doc.add_paragraph('If the phone number belongs to Viettel or VNPT -> Fee = 6,000,000 x (1 - 0.35) = 3,900,000 VND.', style='List Bullet')
    doc.add_paragraph('If the phone number belongs to another carrier -> Fee = 6,000,000 VND (no discount).', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: When calling setPhone(), the registration fee is automatically recalculated. '
        'This ensures the fee is always accurate when the phone number is updated.'
    )

    # --- 3.6. Students ---
    doc.add_heading('3.6. Class Students', level=2)
    doc.add_paragraph(
        'Students extends ArrayList<Student> and manages the entire registration list. '
        'It supports CRUD operations, search, filter, statistics, and file I/O.'
    )
    add_styled_table(doc,
        ['Method', 'Description', 'Notes'],
        [
            ['add(student)', 'Adds a student, checks for duplicate ID', 'Calls markUnsaved() on success'],
            ['update(student)', 'Updates a student by ID', 'Calls markUnsaved() on success'],
            ['delete(id)', 'Deletes a student by ID', 'Calls markUnsaved() on success'],
            ['searchById(id)', 'Exact search by ID (case-insensitive)', 'Returns Student or null'],
            ['searchByName(name)', 'Search by name (partial, case-insensitive)', 'Returns List<Student>'],
            ['filterByCampusCode(campus)', 'Filters by first 2 characters of ID', 'e.g., SE, HE, DE...'],
            ['showAll()', 'Displays the entire list', 'Sorted by ID'],
            ['showAll(list)', 'Displays a sub-list', 'Used for search/filter results'],
            ['readFromFile()', 'Reads .dat file (ObjectInputStream)', 'Called in constructor'],
            ['saveToFile()', 'Saves .dat file + exports CSV', 'Marks as saved afterward'],
            ['statisticalizeByMountainPeak()', 'Generates statistics by peak', 'Creates Statistics object'],
        ],
        col_widths=[5, 5, 5]
    )

    # --- 3.7. StatisticalInfo ---
    doc.add_heading('3.7. Class StatisticalInfo', level=2)
    doc.add_paragraph(
        'StatisticalInfo stores statistics for a specific mountain peak, '
        'including the mountain code, mountain name, number of registered students, and total fee.'
    )
    add_styled_table(doc,
        ['Attribute', 'Data Type', 'Description'],
        [
            ['mountainCode', 'String', 'Mountain peak code'],
            ['mountainName', 'String', 'Mountain peak name'],
            ['numOfStudent', 'int', 'Number of registered students'],
            ['totalCost', 'double', 'Total registration fee (VND)'],
        ],
        col_widths=[3.5, 3, 8.5]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Method addStudent(fee): Increments numOfStudent by 1 and adds fee to totalCost.'
    )

    # --- 3.8. Statistics ---
    doc.add_heading('3.8. Class Statistics', level=2)
    doc.add_paragraph(
        'Statistics extends HashMap<String, StatisticalInfo> and '
        'performs statistical analysis of registration count and total fees grouped by mountain peak.'
    )
    doc.add_paragraph('Statistics logic (statisticalize):')
    stats_steps = [
        'Iterates through the entire student list.',
        'For each student, retrieves the mountainCode.',
        'If no StatisticalInfo exists for that mountainCode, creates a new one and looks up the mountain name from Mountains.',
        'Calls addStudent(fee) to accumulate the count and total.',
        'Displays a statistics table with columns: Code, Peak Name, Number of Participants, Total Cost.',
    ]
    for step in stats_steps:
        doc.add_paragraph(step, style='List Number')

    # --- 3.9. Main ---
    doc.add_heading('3.9. Class Main', level=2)
    doc.add_paragraph(
        'Main is the primary orchestration class containing the main() method and static methods '
        'corresponding to each menu feature. Main initializes 3 core objects: '
        'Inputter (input handling), Mountains (mountain list), and Students (student list).'
    )
    add_styled_table(doc,
        ['Method', 'Menu', 'Feature'],
        [
            ['addNewRegistration()', '1', 'Register a new student'],
            ['updateRegistration()', '2', 'Update registration information'],
            ['displayRegisteredList()', '3', 'Display the registered list'],
            ['deleteRegistration()', '4', 'Delete a registration'],
            ['searchByName()', '5', 'Search by ID or name'],
            ['filterByCampus()', '6', 'Filter by campus'],
            ['showStatistics()', '7', 'Statistics by mountain peak'],
            ['saveDataToFile()', '8', 'Save data to file'],
            ['exitProgram()', '9', 'Exit the program'],
        ],
        col_widths=[5, 1.5, 8.5]
    )

    doc.add_page_break()

    # ===================== CHAPTER 4: USER GUIDE =====================
    doc.add_heading('4. User Guide for Each Feature', level=1)

    # --- 4.1. Main Menu ---
    doc.add_heading('4.1. Main Menu', level=2)
    doc.add_paragraph(
        'When the program starts, the main menu is displayed. '
        'The user enters a number from 1 to 9 to select the corresponding feature. '
        'If an invalid input is entered (letters, out-of-range numbers, special characters), '
        'the program prompts for re-entry without crashing.'
    )
    add_img_placeholder(doc, "Main Menu interface upon program startup")

    doc.add_paragraph('Menu features:')
    menu_items = [
        ('1. New Registration', 'Register a new student for the hiking challenge.'),
        ('2. Update Registration Information', 'Update information of an already registered student.'),
        ('3. Display Registered List', 'Display the entire registration list in table format.'),
        ('4. Delete Registration Information', 'Delete a student\'s registration.'),
        ('5. Search Participants by Name', 'Search for students by ID or name.'),
        ('6. Filter Data by Campus', 'Filter the list by campus (SE/HE/DE/QE/CE).'),
        ('7. Statistics of Registration Numbers by Location', 'Statistics of count and total fee by mountain peak.'),
        ('8. Save Data to File', 'Save data to .dat file and export CSV.'),
        ('9. Exit', 'Exit the program (prompts to save if there is unsaved data).'),
    ]
    for title, desc in menu_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)

    # --- 4.2. New Registration ---
    doc.add_heading('4.2. Feature 1: New Registration', level=2)
    doc.add_paragraph('Description: Allows entering new student information to register for the hiking challenge.')
    doc.add_paragraph('Steps:')
    steps = [
        'Select option 1 from the main menu.',
        'Enter Student ID in the format: SE/HE/DE/QE/CE + 6 digits (e.g., SE203056). If the ID already exists, an error is displayed and re-entry is required.',
        'Enter Name: 2-20 characters, letters and spaces only.',
        'Enter Phone: 10 digits, starting with 0.',
        'Enter Email: valid email format (e.g., example@gmail.com).',
        'The list of 13 mountain peaks is displayed. Enter the desired mountain code (1-13).',
        'The system automatically calculates the fee: 3,900,000 VND (Viettel/VNPT) or 6,000,000 VND (other).',
        'A success message with the registration fee is displayed.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')

    add_img_placeholder(doc, "New Registration screen - entering complete student information")
    add_img_placeholder(doc, "Successful registration result with auto-calculated fee")

    doc.add_paragraph('Error handling during registration:')
    doc.add_paragraph('Invalid ID format -> "Invalid format, please re-enter."', style='List Bullet')
    doc.add_paragraph('Duplicate ID -> "Student ID already exists. Please try again."', style='List Bullet')
    doc.add_paragraph('Invalid name, phone, or email -> "Invalid format, please re-enter."', style='List Bullet')
    doc.add_paragraph('Invalid mountain code -> "Invalid mountain code. Please choose a code from the list."', style='List Bullet')

    add_img_placeholder(doc, "Error message screen when entering invalid data")

    # --- 4.3. Update ---
    doc.add_heading('4.3. Feature 2: Update Registration Information', level=2)
    doc.add_paragraph('Description: Updates the registration information of an existing student.')
    doc.add_paragraph('Steps:')
    steps = [
        'Select option 2 from the main menu.',
        'Enter the Student ID to update.',
        'If the ID does not exist -> "This student has not registered yet." -> Return to menu.',
        'Display the student\'s current information.',
        'For each field (Name, Phone, Email, Mountain Code): enter a new value or press Enter to keep the current value.',
        'If the Phone is changed, the registration fee is automatically recalculated.',
        'Display the updated information.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')

    add_img_placeholder(doc, "Update screen - displaying old information and entering new values")
    add_img_placeholder(doc, "Successful update result")

    doc.add_paragraph(
        'Note: The "press Enter to keep current value" feature is very convenient '
        'when only 1-2 fields need to be modified without re-entering everything.'
    )

    # --- 4.4. Display ---
    doc.add_heading('4.4. Feature 3: Display Registered List', level=2)
    doc.add_paragraph(
        'Description: Displays the entire list of registered students in table format, '
        'sorted by Student ID in ascending order.'
    )
    doc.add_paragraph('Steps:')
    doc.add_paragraph('Step 1: Select option 3 from the main menu.', style='List Number')
    doc.add_paragraph('Step 2: The table is displayed with columns: StudentID, Name, Phone, Email, PeakCode, Fee.', style='List Number')
    doc.add_paragraph('If no one has registered: "No students have registered yet."', style='List Bullet')

    add_img_placeholder(doc, "Display screen showing the full registration list in table format")

    # --- 4.5. Delete ---
    doc.add_heading('4.5. Feature 4: Delete Registration', level=2)
    doc.add_paragraph('Description: Removes a student\'s registration from the list.')
    doc.add_paragraph('Steps:')
    steps = [
        'Select option 4 from the main menu.',
        'Enter the Student ID to delete.',
        'If the ID does not exist -> error message -> return to menu.',
        'Display the student\'s detailed information for confirmation.',
        'Prompt for confirmation: "Are you sure you want to delete this registration? (Y/N)".',
        'If Y -> delete and display "The registration has been successfully deleted."',
        'If N -> "Deletion cancelled." -> return to menu.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')

    add_img_placeholder(doc, "Delete screen - confirmation before deletion")
    add_img_placeholder(doc, "Result after successful deletion")

    # --- 4.6. Search ---
    doc.add_heading('4.6. Feature 5: Search Participants', level=2)
    doc.add_paragraph('Description: Search for students by exact Student ID or by partial name (case-insensitive).')
    doc.add_paragraph('Steps:')
    steps = [
        'Select option 5 from the main menu.',
        'A sub-menu appears: 1. Search by Student ID / 2. Search by Name.',
        'If option 1: Enter Student ID -> exact search -> display result or "No one matches the search criteria!"',
        'If option 2: Enter (partial) name -> find all students whose name contains the input (case-insensitive) -> display result table.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {step}', style='List Number')

    add_img_placeholder(doc, "Search by ID screen - result found")
    add_img_placeholder(doc, "Search by Name screen - multiple matching results")
    add_img_placeholder(doc, "Search screen - no results found")

    # --- 4.7. Filter ---
    doc.add_heading('4.7. Feature 6: Filter by Campus', level=2)
    doc.add_paragraph('Description: Filters the student list by campus code (first 2 characters of the Student ID).')

    add_styled_table(doc,
        ['Campus Code', 'Campus Name'],
        [
            ['CE', 'FPT Can Tho'],
            ['DE', 'FPT Da Nang'],
            ['HE', 'FPT Ha Noi'],
            ['SE', 'FPT Ho Chi Minh'],
            ['QE', 'FPT Quy Nhon'],
        ],
        col_widths=[3, 12]
    )

    doc.add_paragraph()
    doc.add_paragraph('Steps:')
    doc.add_paragraph('Step 1: Select option 6 from the main menu.', style='List Number')
    doc.add_paragraph('Step 2: Enter the campus code (e.g., SE).', style='List Number')
    doc.add_paragraph('Step 3: Display the list of students under that campus, or "No students have registered under this campus." if none found.', style='List Number')

    add_img_placeholder(doc, "Filter by Campus screen - displaying results")

    # --- 4.8. Statistics ---
    doc.add_heading('4.8. Feature 7: Statistics by Mountain', level=2)
    doc.add_paragraph(
        'Description: Generates statistics of the number of registered students and total fee per mountain peak.'
    )
    doc.add_paragraph('Steps:')
    doc.add_paragraph('Step 1: Select option 7 from the main menu.', style='List Number')
    doc.add_paragraph('Step 2: The statistics table is displayed with columns: Code, Peak Name, Number of Participants, Total Cost.', style='List Number')
    doc.add_paragraph('If there is no data: "No registration data available for statistics."', style='List Bullet')

    add_img_placeholder(doc, "Statistics by Mountain screen - results table")

    # --- 4.9. Save ---
    doc.add_heading('4.9. Feature 8: Save Data to File', level=2)
    doc.add_paragraph(
        'Description: Saves all registration data to 2 files:'
    )
    doc.add_paragraph('registrations.dat - Binary file, uses ObjectOutputStream to serialize ArrayList<Student>.', style='List Bullet')
    doc.add_paragraph('registrations.csv - Text CSV file, easily opened with Excel. Columns: StudentID, Name, Phone, Email, MountainCode, TuitionFee.', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Steps:')
    doc.add_paragraph('Step 1: Select option 8 from the main menu.', style='List Number')
    doc.add_paragraph('Step 2: The system automatically saves and displays "Registration data has been successfully saved."', style='List Number')

    add_img_placeholder(doc, "Successful data save screen")
    add_img_placeholder(doc, "Contents of registrations.csv opened in Excel or Notepad")

    # --- 4.10. Exit ---
    doc.add_heading('4.10. Feature 9: Exit Program', level=2)
    doc.add_paragraph('Description: Exits the program with unsaved data checking.')
    doc.add_paragraph('Exit workflow:')
    exit_steps = [
        'Select option 9 from the main menu.',
        'If there is NO unsaved data -> Exit immediately, display "Goodbye!".',
        'If there IS unsaved data:',
        '    a. Prompt: "You have unsaved changes. Do you want to save before exiting? (Y/N)"',
        '    b. If Y -> Save to file -> Exit.',
        '    c. If N -> Prompt again: "Are you sure you want to exit without saving? (Y/N)"',
        '    d. If Y -> Exit without saving.',
        '    e. If N -> Return to menu (do not exit).',
    ]
    for step in exit_steps:
        doc.add_paragraph(step, style='List Number')

    add_img_placeholder(doc, "Exit screen - prompt to save when there is unsaved data")
    add_img_placeholder(doc, "Exit flow: chose not to save -> second confirmation")

    doc.add_page_break()

    # ===================== CHAPTER 5: OVERALL WORKFLOW =====================
    doc.add_heading('5. Overall Workflow', level=1)

    doc.add_heading('5.1. Program Overall Flowchart', level=2)
    doc.add_paragraph(
        'The flowchart below describes the overall workflow of the program from start to finish. '
        'The program operates in a menu-driven model (loop), meaning after completing a feature, '
        'the system returns to the main menu until the user selects Exit.'
    )
    add_img_placeholder(doc, "Overall Flowchart - Complete program workflow")

    doc.add_paragraph('Workflow explanation:')
    doc.add_paragraph(
        'Start -> Load MountainList.csv into Mountains -> Load registrations.dat into Students (if file exists) '
        '-> Display Menu -> Read user choice -> Execute feature -> Return to Menu (loop) -> '
        'When Exit is selected -> Check unsaved data -> Prompt to save -> End.'
    )

    doc.add_heading('5.2. Flowchart per Feature', level=2)
    doc.add_paragraph(
        'Below are detailed flowcharts for each feature. '
        'Each flowchart clearly shows the processing flow, validation steps, and error cases.'
    )

    functions_flow = [
        ("Feature 1 - New Registration",
         "Flowchart: Enter ID -> Check duplicate -> Enter Name/Phone/Email -> Validate -> Display mountain list -> Enter Mountain Code -> Calculate fee -> Add to list"),
        ("Feature 2 - Update Registration",
         "Flowchart: Enter ID -> Find student -> Display current info -> Enter new values for each field (Enter to keep) -> Validate -> Recalculate fee if phone changed -> Update"),
        ("Feature 3 - Display Registered List",
         "Flowchart: Check if list is empty -> Sort by ID -> Print table"),
        ("Feature 4 - Delete Registration",
         "Flowchart: Enter ID -> Find student -> Display info -> Confirm Y/N -> Delete or cancel"),
        ("Feature 5 - Search Participants",
         "Flowchart: Choose search by ID or Name -> Enter keyword -> Search -> Display results"),
        ("Feature 6 - Filter by Campus",
         "Flowchart: Enter campus code -> Validate -> Filter by first 2 characters of ID -> Display"),
        ("Feature 7 - Statistics",
         "Flowchart: Iterate student list -> Group by mountainCode -> Count quantity + sum fee -> Display table"),
        ("Feature 8 - Save Data",
         "Flowchart: Serialize -> ObjectOutputStream -> registrations.dat + Export CSV"),
        ("Feature 9 - Exit",
         "Flowchart: Check isSaved -> If false: prompt to save (Y/N) -> If N: confirm again -> Exit or return to menu"),
    ]
    for title, desc in functions_flow:
        p = doc.add_paragraph()
        run = p.add_run(f"> {title}")
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(desc)
        add_img_placeholder(doc, f"Flowchart - {title}")
        doc.add_paragraph()

    doc.add_heading('5.3. Sequence Diagram - New Registration', level=2)
    doc.add_paragraph(
        'The Sequence Diagram describes the interaction sequence between objects when performing the '
        'New Registration feature. Participating objects: Operator, Main, Inputter, Mountains, Students, Student.'
    )
    add_img_placeholder(doc, "Sequence Diagram - New Registration workflow")

    doc.add_paragraph('Sequence Diagram explanation:')
    seq_steps = [
        'Operator selects New Registration from the menu.',
        'Main calls Inputter to input and validate ID, Name, Phone, Email.',
        'Main calls Students.searchById() to check for duplicate ID.',
        'Main calls Mountains.isValidMountainCode() to validate the mountain code.',
        'Main creates a new Student object and calls calculateFee() to compute the fee.',
        'Main calls Students.add() to add the student to the list and marks unsaved.',
        'Main displays a success message to the Operator.',
    ]
    for i, step in enumerate(seq_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_page_break()

    # ===================== CHAPTER 6: SAMPLE DATA =====================
    doc.add_heading('6. Sample Data', level=1)

    doc.add_heading('6.1. MountainList.csv File', level=2)
    doc.add_paragraph(
        'The MountainList.csv file contains a list of 13 mountain peaks, loaded when the program starts. '
        'Structure: Code, Mountain, Province, Description.'
    )
    mountains_data = [
        ['1', 'Ham Rong Mountain', 'Lao Cai', 'Near Sa Pa center, highest point 1850m'],
        ['2', 'Doi Bo Mountain', 'Lao Cai', ''],
        ['3', 'Pha Luong Mountain', 'Son La', 'Nearly 2000m, the Roof of Moc Chau'],
        ['4', 'Hon Vuon Mountain', 'Hue', ''],
        ['5', 'Da Do Mountain', 'Ninh Thuan', ''],
        ['6', 'Da Bia Mountain', 'Phu Yen', ''],
        ['7', 'Chu Hreng Mountain', 'Kon Tum', ''],
        ['8', 'Lang Biang Mountain', 'Lam Dong', ''],
        ['9', 'Ta Nang Mountain', 'Lam Dong', ''],
        ['10', 'Cam Mountain', 'An Giang', ''],
        ['11', 'Thi Vai Mountain', 'Vung Tau', ''],
        ['12', 'Dinh Mountain', 'Vung Tau', ''],
        ['13', 'Co Tien Mountain', 'Khanh Hoa', 'Has 3 peaks; conquer peak 1 only'],
    ]
    add_styled_table(doc,
        ['Code', 'Mountain', 'Province', 'Description'],
        mountains_data,
        col_widths=[1.5, 5, 3, 5.5]
    )

    doc.add_heading('6.2. registrations.csv File', level=2)
    doc.add_paragraph(
        'The registrations.csv file is exported when the user selects Save Data. '
        'This is a text file that can be opened with Excel.'
    )
    add_styled_table(doc,
        ['StudentID', 'Name', 'Phone', 'Email', 'MountainCode', 'TuitionFee'],
        [
            ['SE203056', 'Nguyen Phu Khuong', '0363561629', 'khuong@gmail.com', '1', '3,900,000'],
        ],
        col_widths=[2.5, 4, 2.5, 4, 2.5, 2]
    )

    doc.add_page_break()

    # ===================== CHAPTER 7: VALIDATION RULES =====================
    doc.add_heading('7. Data Validation Rules', level=1)
    doc.add_paragraph(
        'The system applies strict input data validation using Regular Expressions (Regex). '
        'The table below summarizes the validation rules for each data field.'
    )
    add_styled_table(doc,
        ['Field', 'Rule', 'Valid Examples', 'Invalid Examples'],
        [
            ['Student ID', '2-char campus code (SE/HE/DE/QE/CE)\n+ 6 digits', 'SE203056\nHE180001', 'AB123456\nSE12345\nSE1234567'],
            ['Name', '2-20 characters, letters\nand spaces only', 'Nguyen Van A\nLe Thi B', 'A\n12345\nName longer than 20 chars abc'],
            ['Phone', '10 digits, starts with 0', '0363561629\n0912345678', '123456789\n0123\nabcdefghij'],
            ['Email', 'user@domain.ext\n(letters, digits, +, _, ., -)', 'abc@gmail.com\ntest_1@fpt.edu.vn', '@gmail.com\nabc@\nabc.gmail.com'],
            ['Mountain Code', 'Must exist in the\nMountainList.csv list (1-13)', '1\n5\n13', '0\n14\nabc'],
            ['Campus Code', 'SE, HE, DE, QE, or CE\n(case-insensitive)', 'SE\nhe\nDE', 'AB\nSF\n123'],
            ['Menu Choice', 'Integer from 1 to 9', '1\n5\n9', '0\n10\nabc'],
            ['Confirm (Y/N)', 'Y or N\n(case-insensitive)', 'Y\nn', 'yes\nno\n1'],
        ],
        col_widths=[3, 4, 4, 4]
    )

    doc.add_paragraph()
    doc.add_paragraph('Registration fee calculation rules:')
    add_styled_table(doc,
        ['Carrier', 'Prefix Codes', 'Registration Fee'],
        [
            ['Viettel', '032, 033, 034, 035, 036, 037, 038, 039, 096, 097, 098, 086', '3,900,000 VND (35% discount)'],
            ['VNPT', '081, 082, 083, 084, 085, 088, 091, 094', '3,900,000 VND (35% discount)'],
            ['Other (Mobifone, ...)', 'Other prefix codes', '6,000,000 VND (no discount)'],
        ],
        col_widths=[3.5, 8, 3.5]
    )

    doc.add_page_break()

    # ===================== CHAPTER 8: CONCLUSION =====================
    doc.add_heading('8. Conclusion', level=1)

    doc.add_heading('8.1. Achievements', level=2)
    results = [
        'Successfully completed all 9 features as required by the LAB211 assignment.',
        'Applied Object-Oriented Programming (OOP) with a clear structure of 9 classes/interfaces.',
        'Implemented strict input data validation using Regex; the program does not crash on invalid input.',
        'Automatic fee calculation based on phone carrier (Viettel/VNPT receive 35% discount).',
        'Reads mountain data from CSV file; saves registrations to binary file (.dat) and exports CSV.',
        'Smart exit handling: warns the user when there is unsaved data.',
        'Displays data in formatted tables with headers, dividers, and sorted by ID.',
        'Code exceeds 200 LOC with clear variable/method naming following conventions.',
    ]
    for r in results:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading('8.2. Applied Knowledge', level=2)
    knowledge = [
        'OOP: Class, Interface, Inheritance, Polymorphism, Encapsulation.',
        'Java Collections: ArrayList, HashMap, Collections.sort().',
        'Java I/O: BufferedReader, FileReader, ObjectInputStream/ObjectOutputStream, BufferedWriter.',
        'Serialization: Storing Java objects to binary file.',
        'Comparable: Sorting objects by custom criteria.',
        'Regex: Validating input data formats.',
        'Design Pattern: Menu-driven architecture, Separation of Concerns.',
    ]
    for k in knowledge:
        doc.add_paragraph(k, style='List Bullet')

    doc.add_heading('8.3. Limitations and Future Improvements', level=2)
    doc.add_paragraph('Limitations:')
    limitations = [
        'No graphical user interface (GUI); the application is console-based only.',
        'No multi-language support (Vietnamese accented names may encounter encoding issues).',
        'No automated unit tests.',
    ]
    for l in limitations:
        doc.add_paragraph(l, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Future improvements:')
    improvements = [
        'Build a GUI using JavaFX or Swing.',
        'Connect to a database (MySQL, SQLite) instead of files.',
        'Add PDF report export functionality.',
        'Write unit tests using JUnit.',
        'Support multi-language and full UTF-8 encoding.',
    ]
    for imp in improvements:
        doc.add_paragraph(imp, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('--- End of Report ---')
    run.bold = True
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # ===================== SAVE =====================
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'Report_MountainHiking_LAB211_EN.docx'
    )
    doc.save(output_path)
    print(f"[OK] Report created: {output_path}")
    return output_path


if __name__ == '__main__':
    create_report()
