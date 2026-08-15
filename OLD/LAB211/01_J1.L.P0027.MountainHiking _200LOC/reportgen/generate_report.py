# -*- coding: utf-8 -*-
"""
Script tạo báo cáo Word (.docx) cho dự án Mountain Hiking Registration
LAB211 - FPT University
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Tô màu nền cho ô trong bảng."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Đặt border cho ô."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr in kwargs[edge]:
                element.set(qn(f'w:{attr}'), str(kwargs[edge][attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_img_placeholder(doc, caption=""):
    """Thêm placeholder [IMG] cho người dùng chèn ảnh."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[IMG: {caption}]\n")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 0, 0)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"(Chèn ảnh: {caption})")
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(128, 128, 128)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Tạo bảng có style đẹp."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    header_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2E74B5')
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(text)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(row.cells[c_idx], 'D6E4F0')
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

def add_code_block(doc, code, language="Java"):
    """Thêm khối code có format."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"📄 {language} Code:")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    code_p = doc.add_paragraph()
    code_p.style = 'No Spacing'
    code_p.paragraph_format.left_indent = Cm(1)
    code_run = code_p.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)
    code_run.font.color.rgb = RGBColor(40, 40, 40)

def create_report():
    doc = Document()

    # ===================== PAGE SETUP =====================
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # ===================== STYLES =====================
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.5

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # ===================== TRANG BÌA =====================
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FPT UNIVERSITY")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xF4, 0x7B, 0x20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SOFTWARE ENGINEERING DEPARTMENT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BÁO CÁO BÀI LAB")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LAB211 – BASIC JAVA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("J1.L.P0027 – Mountain Hiking Challenge Registration")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info_items = [
        ("Sinh viên thực hiện", "Nguyễn Phú Khương"),
        ("MSSV", "SE203056"),
        ("Lớp", "SE06203"),
        ("Môn học", "LAB211 – Basic Java"),
        ("Giảng viên hướng dẫn", "............................."),
        ("Thời gian", "Summer 2026"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run(f"{label}: ")
        run1.bold = True
        run1.font.size = Pt(13)
        run2 = p.add_run(value)
        run2.font.size = Pt(13)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hồ Chí Minh, tháng 06 năm 2026")
    run.italic = True
    run.font.size = Pt(12)

    # ===================== PAGE BREAK =====================
    doc.add_page_break()

    # ===================== MỤC LỤC =====================
    doc.add_heading('MỤC LỤC', level=1)
    toc_items = [
        ("1.", "Giới thiệu dự án", "3"),
        ("  1.1.", "Mô tả đề bài", "3"),
        ("  1.2.", "Yêu cầu bài Lab", "3"),
        ("  1.3.", "Công nghệ sử dụng", "4"),
        ("2.", "Thiết kế hệ thống", "5"),
        ("  2.1.", "Cấu trúc dự án", "5"),
        ("  2.2.", "Class Diagram", "5"),
        ("  2.3.", "Use Case Diagram", "6"),
        ("  2.4.", "Flowchart tổng quan", "6"),
        ("3.", "Mô tả chi tiết các lớp (Class)", "7"),
        ("  3.1.", "Interface Acceptable", "7"),
        ("  3.2.", "Class Inputter", "8"),
        ("  3.3.", "Class Mountain", "9"),
        ("  3.4.", "Class Mountains", "9"),
        ("  3.5.", "Class Student", "10"),
        ("  3.6.", "Class Students", "11"),
        ("  3.7.", "Class StatisticalInfo", "12"),
        ("  3.8.", "Class Statistics", "13"),
        ("  3.9.", "Class Main", "13"),
        ("4.", "Hướng dẫn sử dụng từng chức năng", "14"),
        ("  4.1.", "Menu chính", "14"),
        ("  4.2.", "Chức năng 1: Đăng ký mới (New Registration)", "15"),
        ("  4.3.", "Chức năng 2: Cập nhật thông tin (Update Registration)", "16"),
        ("  4.4.", "Chức năng 3: Hiển thị danh sách (Display Registered List)", "17"),
        ("  4.5.", "Chức năng 4: Xoá đăng ký (Delete Registration)", "18"),
        ("  4.6.", "Chức năng 5: Tìm kiếm (Search Participants)", "19"),
        ("  4.7.", "Chức năng 6: Lọc theo Campus (Filter by Campus)", "20"),
        ("  4.8.", "Chức năng 7: Thống kê (Statistics by Mountain)", "21"),
        ("  4.9.", "Chức năng 8: Lưu dữ liệu (Save Data to File)", "22"),
        ("  4.10.", "Chức năng 9: Thoát chương trình (Exit)", "23"),
        ("5.", "Luồng hoạt động tổng thể", "24"),
        ("  5.1.", "Flowchart tổng quan chương trình", "24"),
        ("  5.2.", "Flowchart từng chức năng", "25"),
        ("  5.3.", "Sequence Diagram – Đăng ký mới", "26"),
        ("6.", "Dữ liệu mẫu", "27"),
        ("  6.1.", "File MountainList.csv", "27"),
        ("  6.2.", "File registrations.csv", "27"),
        ("7.", "Quy tắc validation dữ liệu", "28"),
        ("8.", "Kết luận", "29"),
    ]
    for num, title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run_num = p.add_run(f"{num} ")
        run_num.bold = True
        run_num.font.size = Pt(12)
        run_title = p.add_run(f"{title}")
        run_title.font.size = Pt(12)
        tab_run = p.add_run(f"  ........ {page}")
        tab_run.font.size = Pt(12)
        tab_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # ===================== CHƯƠNG 1: GIỚI THIỆU DỰ ÁN =====================
    doc.add_heading('1. Giới thiệu dự án', level=1)

    doc.add_heading('1.1. Mô tả đề bài', level=2)
    doc.add_paragraph(
        'Dự án "J1.L.P0027 – Mountain Hiking Challenge Registration" là một ứng dụng console được xây dựng bằng '
        'ngôn ngữ Java, dùng để quản lý việc đăng ký của sinh viên tham gia Thử thách Leo núi (Mountain Hiking Challenge). '
        'Chương trình cho phép người vận hành (Operator) thực hiện các thao tác CRUD (Create, Read, Update, Delete) '
        'trên danh sách đăng ký, đồng thời hỗ trợ tìm kiếm, lọc theo campus, thống kê theo đỉnh núi, và lưu/đọc dữ liệu từ file.'
    )
    doc.add_paragraph(
        'Ứng dụng được phát triển trong khuôn khổ môn LAB211 – Basic Java tại Đại học FPT, '
        'nhằm rèn luyện kỹ năng lập trình hướng đối tượng (OOP), xử lý file I/O, '
        'kiểm tra dữ liệu đầu vào (validation), và tổ chức mã nguồn theo cấu trúc lớp rõ ràng.'
    )

    doc.add_heading('1.2. Yêu cầu bài Lab', level=2)
    doc.add_paragraph('Theo đề bài LAB211, sinh viên cần thực hiện các yêu cầu sau:')
    requirements = [
        'Xây dựng ứng dụng console Java quản lý đăng ký leo núi cho sinh viên FPT.',
        'Menu gồm 9 chức năng: Đăng ký mới, Cập nhật, Hiển thị, Xoá, Tìm kiếm, Lọc, Thống kê, Lưu file, Thoát.',
        'Áp dụng OOP: sử dụng class, interface, kế thừa ArrayList, Serializable, Comparable.',
        'Validate đầu vào bằng Regex: Student ID, Phone, Email, Name, Campus Code.',
        'Tính phí đăng ký: mặc định 6,000,000 VND, giảm 35% nếu thuộc nhà mạng Viettel hoặc VNPT.',
        'Đọc danh sách núi từ file MountainList.csv.',
        'Lưu dữ liệu đăng ký ra file registrations.dat (binary) và xuất ra registrations.csv.',
        'Hỏi xác nhận khi thoát nếu có dữ liệu chưa lưu.',
        'Code tối thiểu 200 LOC (Lines of Code).',
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('1.3. Công nghệ sử dụng', level=2)
    add_styled_table(doc,
        ['STT', 'Công nghệ', 'Mô tả'],
        [
            ['1', 'Java SE', 'Ngôn ngữ lập trình chính (Java 8+)'],
            ['2', 'Java I/O', 'Đọc/ghi file CSV, serialization (.dat)'],
            ['3', 'Java Collections', 'ArrayList, HashMap, Collections.sort()'],
            ['4', 'Regex (Pattern)', 'Kiểm tra định dạng dữ liệu đầu vào'],
            ['5', 'Serialization', 'Lưu trữ đối tượng Student ra file nhị phân'],
            ['6', 'IntelliJ IDEA', 'IDE phát triển'],
        ],
        col_widths=[1.5, 4, 10]
    )

    doc.add_page_break()

    # ===================== CHƯƠNG 2: THIẾT KẾ HỆ THỐNG =====================
    doc.add_heading('2. Thiết kế hệ thống', level=1)

    doc.add_heading('2.1. Cấu trúc dự án', level=2)
    doc.add_paragraph('Dự án được tổ chức theo cấu trúc thư mục như sau:')
    structure_items = [
        'Project1_MountainHiking/',
        '├── src/',
        '│   ├── Main.java              (Lớp chính, chứa menu và điều phối)',
        '│   ├── Acceptable.java         (Interface chứa các hằng Regex)',
        '│   ├── Inputter.java           (Lớp xử lý nhập liệu từ console)',
        '│   ├── Mountain.java           (Lớp đối tượng Núi)',
        '│   ├── Mountains.java          (Danh sách các núi, đọc từ CSV)',
        '│   ├── Student.java            (Lớp đối tượng Sinh viên)',
        '│   ├── Students.java           (Danh sách sinh viên, CRUD, file I/O)',
        '│   ├── StatisticalInfo.java     (Thông tin thống kê 1 đỉnh núi)',
        '│   └── Statistics.java          (Bảng thống kê tổng hợp)',
        '├── MountainList.csv            (Dữ liệu danh sách 13 đỉnh núi)',
        '├── registrations.dat           (File nhị phân lưu đăng ký)',
        '├── registrations.csv           (File CSV xuất dữ liệu đăng ký)',
        '├── Class Diagram.drawio        (Sơ đồ lớp)',
        '├── FlowChart.drawio            (Lưu đồ tổng quan)',
        '└── FlowChart_PerFunc.drawio    (Lưu đồ từng chức năng)',
    ]
    for item in structure_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(item)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

    doc.add_paragraph()

    doc.add_heading('2.2. Class Diagram', level=2)
    doc.add_paragraph(
        'Dưới đây là sơ đồ lớp (Class Diagram) thể hiện mối quan hệ giữa các lớp trong hệ thống. '
        'Các lớp được thiết kế theo nguyên tắc OOP với các mối quan hệ kế thừa, triển khai interface, '
        'và tham chiếu giữa các đối tượng.'
    )
    add_img_placeholder(doc, "Class Diagram – Sơ đồ lớp của hệ thống")

    doc.add_paragraph('Giải thích các mối quan hệ chính trong Class Diagram:')
    relationships = [
        'Student implements Serializable: cho phép serialize đối tượng Student để lưu vào file .dat.',
        'Student implements Comparable<Student>: cho phép so sánh và sắp xếp sinh viên theo ID.',
        'Students extends ArrayList<Student>: danh sách sinh viên kế thừa từ ArrayList.',
        'Mountains extends ArrayList<Mountain>: danh sách núi kế thừa từ ArrayList.',
        'Statistics extends HashMap<String, StatisticalInfo>: bảng thống kê sử dụng HashMap.',
        'Inputter sử dụng Acceptable: gọi phương thức isValid() để validate dữ liệu.',
        'Students tham chiếu đến Student, Statistics tham chiếu đến StatisticalInfo.',
    ]
    for rel in relationships:
        doc.add_paragraph(rel, style='List Bullet')

    doc.add_heading('2.3. Use Case Diagram', level=2)
    doc.add_paragraph(
        'Sơ đồ Use Case mô tả các chức năng mà Operator (người vận hành) có thể thực hiện trong hệ thống. '
        'Tất cả các chức năng nhập liệu đều bao gồm (include) bước Validate Input.'
    )
    add_img_placeholder(doc, "Use Case Diagram – Sơ đồ ca sử dụng")

    doc.add_heading('2.4. Flowchart tổng quan', level=2)
    doc.add_paragraph(
        'Flowchart tổng quan mô tả luồng hoạt động chính của chương trình từ khi khởi động '
        'đến khi thoát. Chương trình hoạt động theo vòng lặp menu-driven.'
    )
    add_img_placeholder(doc, "Flowchart tổng quan – Luồng hoạt động chính của chương trình")

    doc.add_paragraph('Mô tả luồng tổng quan:')
    flow_steps = [
        'Bước 1: Khởi động – Load danh sách núi từ MountainList.csv.',
        'Bước 2: Load dữ liệu đăng ký từ registrations.dat (nếu tồn tại).',
        'Bước 3: Hiển thị menu chính với 9 lựa chọn.',
        'Bước 4: Người dùng nhập lựa chọn (validate 1-9).',
        'Bước 5: Thực thi chức năng tương ứng.',
        'Bước 6: Đánh dấu dữ liệu unsaved nếu có thay đổi (Add/Update/Delete).',
        'Bước 7: Quay lại Bước 3 cho đến khi chọn Exit (9).',
        'Bước 8: Nếu có dữ liệu chưa lưu → hỏi có muốn lưu không → Kết thúc.',
    ]
    for step in flow_steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_page_break()

    # ===================== CHƯƠNG 3: MÔ TẢ CHI TIẾT CÁC LỚP =====================
    doc.add_heading('3. Mô tả chi tiết các lớp (Class)', level=1)

    # --- 3.1. Acceptable ---
    doc.add_heading('3.1. Interface Acceptable', level=2)
    doc.add_paragraph(
        'Acceptable là một interface chứa các hằng số regex (Regular Expression) '
        'dùng để validate dữ liệu đầu vào, và phương thức static isValid() '
        'để kiểm tra chuỗi có khớp regex hay không.'
    )
    add_styled_table(doc,
        ['Hằng số', 'Regex Pattern', 'Ý nghĩa'],
        [
            ['STUDENT_ID', '^(?i)(SE|HE|DE|QE|CE)\\d{6}$', 'Mã SV: 2 ký tự campus + 6 số'],
            ['CAMPUS_CODE', '^(?i)(SE|HE|DE|QE|CE)$', 'Mã campus: SE/HE/DE/QE/CE'],
            ['NAME_VALID', '^[A-Za-zÀ-ỹ\\s]{2,20}$', 'Tên: 2-20 ký tự chữ và dấu cách'],
            ['PHONE_VALID', '^0\\d{9}$', 'SĐT: bắt đầu bằng 0, đủ 10 số'],
            ['VIETTEL_VALID', '^(032|033|...|086)\\d{7}$', 'SĐT mạng Viettel (giảm phí)'],
            ['VNPT_VALID', '^(081|082|...|094)\\d{7}$', 'SĐT mạng VNPT (giảm phí)'],
            ['EMAIL_VALID', '^[A-Za-z0-9+_.-]+@...+\\.[A-Za-z]{2,}$', 'Định dạng email hợp lệ'],
            ['YES_NO_VALID', '^[YyNn]$', 'Xác nhận Y hoặc N'],
        ],
        col_widths=[3.5, 5.5, 6]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Phương thức isValid(String data, String pattern): Sử dụng Pattern.matches() '
        'để kiểm tra xem chuỗi data có khớp với pattern hay không. Trả về true nếu khớp, false nếu không.'
    )

    # --- 3.2. Inputter ---
    doc.add_heading('3.2. Class Inputter', level=2)
    doc.add_paragraph(
        'Inputter là lớp tiện ích xử lý việc nhập liệu từ bàn phím (console). '
        'Lớp này đảm bảo mọi dữ liệu nhập vào đều được validate trước khi trả về.'
    )
    add_styled_table(doc,
        ['Phương thức', 'Mô tả', 'Đặc điểm'],
        [
            ['getString(message)', 'Hiển thị prompt, đọc chuỗi từ bàn phím', 'Trả về chuỗi đã trim()'],
            ['getInt(message)', 'Nhập số nguyên, lặp đến khi hợp lệ', 'Dùng INTEGER_VALID regex'],
            ['getDouble(message)', 'Nhập số thực, lặp đến khi hợp lệ', 'Dùng DOUBLE_VALID regex'],
            ['getMenuChoice(msg, min, max)', 'Nhập lựa chọn menu trong khoảng [min, max]', 'Lặp nếu ngoài phạm vi'],
            ['inputAndLoop(msg, pattern)', 'Nhập và kiểm tra regex, lặp nếu sai', 'Không cho phép bỏ trống'],
            ['inputAndLoopAllowEmpty(msg, pattern)', 'Như trên nhưng cho phép Enter (bỏ trống)', 'Dùng khi Update'],
            ['confirmYesNo(message)', 'Hỏi xác nhận Y/N', 'Trả về true nếu Y'],
        ],
        col_widths=[5.5, 5, 4.5]
    )

    # --- 3.3. Mountain ---
    doc.add_heading('3.3. Class Mountain', level=2)
    doc.add_paragraph(
        'Mountain đại diện cho một đỉnh núi trong chương trình. '
        'Mỗi đối tượng Mountain chứa thông tin về mã núi, tên núi, tỉnh/thành, và mô tả.'
    )
    add_styled_table(doc,
        ['Thuộc tính', 'Kiểu dữ liệu', 'Mô tả'],
        [
            ['mountainCode', 'String', 'Mã đỉnh núi (VD: 1, 2, 3...)'],
            ['mountain', 'String', 'Tên đỉnh núi (VD: Ham Rong Mountain)'],
            ['province', 'String', 'Tỉnh/thành phố nơi đỉnh núi toạ lạc'],
            ['description', 'String', 'Mô tả chi tiết về đỉnh núi'],
        ],
        col_widths=[3.5, 3, 8.5]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Các phương thức đặc biệt: toString() hiển thị thông tin theo cột, '
        'equals() so sánh bằng mountainCode (case-insensitive), '
        'hashCode() dựa trên mountainCode uppercase.'
    )

    # --- 3.4. Mountains ---
    doc.add_heading('3.4. Class Mountains', level=2)
    doc.add_paragraph(
        'Mountains kế thừa ArrayList<Mountain>, quản lý toàn bộ danh sách các đỉnh núi. '
        'Tự động đọc dữ liệu từ file MountainList.csv khi khởi tạo.'
    )
    add_styled_table(doc,
        ['Phương thức', 'Mô tả'],
        [
            ['Mountains()', 'Constructor mặc định, đọc file MountainList.csv'],
            ['Mountains(pathFile)', 'Constructor có tham số đường dẫn file'],
            ['get(mountainCode)', 'Tìm Mountain theo mã (case-insensitive)'],
            ['isValidMountainCode(code)', 'Kiểm tra mã núi có tồn tại không'],
            ['dataToObject(text)', 'Chuyển đổi dòng CSV thành đối tượng Mountain'],
            ['readFromFile()', 'Đọc toàn bộ file CSV, bỏ qua header, bỏ dòng trống'],
            ['showAll()', 'Hiển thị danh sách núi dạng bảng có header'],
        ],
        col_widths=[5, 10]
    )

    # --- 3.5. Student ---
    doc.add_heading('3.5. Class Student', level=2)
    doc.add_paragraph(
        'Student đại diện cho một sinh viên đăng ký tham gia leo núi. '
        'Lớp này implement Serializable (để lưu file nhị phân) và Comparable<Student> (để sắp xếp theo ID).'
    )
    add_styled_table(doc,
        ['Thuộc tính', 'Kiểu dữ liệu', 'Mô tả'],
        [
            ['id', 'String', 'Mã sinh viên (VD: SE203056)'],
            ['name', 'String', 'Họ tên sinh viên'],
            ['phone', 'String', 'Số điện thoại (10 số, bắt đầu bằng 0)'],
            ['email', 'String', 'Địa chỉ email'],
            ['mountainCode', 'String', 'Mã đỉnh núi đã đăng ký'],
            ['tuitionFee', 'double', 'Phí đăng ký (VND)'],
        ],
        col_widths=[3, 3, 9]
    )
    doc.add_paragraph()
    doc.add_paragraph('Hằng số quan trọng:')
    doc.add_paragraph('DEFAULT_FEE = 6,000,000 VND – Phí đăng ký mặc định.', style='List Bullet')
    doc.add_paragraph('DISCOUNT_RATE = 0.35 (35%) – Tỷ lệ giảm giá cho SĐT Viettel/VNPT.', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Logic tính phí (calculateFee):')
    doc.add_paragraph('Nếu SĐT thuộc nhà mạng Viettel hoặc VNPT → Phí = 6,000,000 × (1 – 0.35) = 3,900,000 VND.', style='List Bullet')
    doc.add_paragraph('Nếu SĐT thuộc nhà mạng khác → Phí = 6,000,000 VND (không giảm).', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'Đặc biệt: Khi gọi setPhone(), phí đăng ký sẽ được tự động tính lại. '
        'Điều này đảm bảo phí luôn chính xác khi cập nhật số điện thoại.'
    )

    # --- 3.6. Students ---
    doc.add_heading('3.6. Class Students', level=2)
    doc.add_paragraph(
        'Students kế thừa ArrayList<Student>, quản lý toàn bộ danh sách đăng ký. '
        'Hỗ trợ CRUD, tìm kiếm, lọc, thống kê, và lưu/đọc file.'
    )
    add_styled_table(doc,
        ['Phương thức', 'Mô tả', 'Ghi chú'],
        [
            ['add(student)', 'Thêm sinh viên, kiểm tra trùng ID', 'markUnsaved() nếu thành công'],
            ['update(student)', 'Cập nhật SV theo ID', 'markUnsaved() nếu thành công'],
            ['delete(id)', 'Xoá SV theo ID', 'markUnsaved() nếu thành công'],
            ['searchById(id)', 'Tìm chính xác theo ID (case-insensitive)', 'Trả về Student hoặc null'],
            ['searchByName(name)', 'Tìm theo tên (partial, case-insensitive)', 'Trả về List<Student>'],
            ['filterByCampusCode(campus)', 'Lọc theo 2 ký tự đầu ID', 'VD: SE, HE, DE...'],
            ['showAll()', 'Hiển thị toàn bộ danh sách', 'Sắp xếp theo ID'],
            ['showAll(list)', 'Hiển thị danh sách con', 'Dùng cho kết quả search/filter'],
            ['readFromFile()', 'Đọc file .dat (ObjectInputStream)', 'Gọi trong constructor'],
            ['saveToFile()', 'Lưu file .dat + xuất CSV', 'markSaved sau khi lưu'],
            ['statisticalizeByMountainPeak()', 'Thống kê theo đỉnh núi', 'Tạo đối tượng Statistics'],
        ],
        col_widths=[5, 5, 5]
    )

    # --- 3.7. StatisticalInfo ---
    doc.add_heading('3.7. Class StatisticalInfo', level=2)
    doc.add_paragraph(
        'StatisticalInfo lưu trữ thông tin thống kê cho một đỉnh núi cụ thể, '
        'bao gồm mã núi, tên núi, số lượng sinh viên đăng ký, và tổng phí.'
    )
    add_styled_table(doc,
        ['Thuộc tính', 'Kiểu dữ liệu', 'Mô tả'],
        [
            ['mountainCode', 'String', 'Mã đỉnh núi'],
            ['mountainName', 'String', 'Tên đỉnh núi'],
            ['numOfStudent', 'int', 'Số lượng sinh viên đã đăng ký'],
            ['totalCost', 'double', 'Tổng phí đăng ký (VND)'],
        ],
        col_widths=[3.5, 3, 8.5]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Phương thức addStudent(fee): Tăng numOfStudent lên 1 và cộng fee vào totalCost.'
    )

    # --- 3.8. Statistics ---
    doc.add_heading('3.8. Class Statistics', level=2)
    doc.add_paragraph(
        'Statistics kế thừa HashMap<String, StatisticalInfo>, '
        'thực hiện việc thống kê số lượng đăng ký và tổng phí theo từng đỉnh núi.'
    )
    doc.add_paragraph('Logic thống kê (statisticalize):')
    stats_steps = [
        'Duyệt qua toàn bộ danh sách sinh viên.',
        'Với mỗi sinh viên, lấy mountainCode.',
        'Nếu chưa có StatisticalInfo cho mountainCode → tạo mới, tra tên núi từ Mountains.',
        'Gọi addStudent(fee) để cộng dồn.',
        'Hiển thị bảng thống kê với các cột: Code, Peak Name, Number of Participants, Total Cost.',
    ]
    for step in stats_steps:
        doc.add_paragraph(step, style='List Number')

    # --- 3.9. Main ---
    doc.add_heading('3.9. Class Main', level=2)
    doc.add_paragraph(
        'Main là lớp điều phối chính, chứa phương thức main() và các phương thức static '
        'tương ứng với từng chức năng trong menu. Main khởi tạo 3 đối tượng chính: '
        'Inputter (nhập liệu), Mountains (danh sách núi), và Students (danh sách SV).'
    )
    add_styled_table(doc,
        ['Phương thức', 'Menu', 'Chức năng'],
        [
            ['addNewRegistration()', '1', 'Đăng ký mới sinh viên'],
            ['updateRegistration()', '2', 'Cập nhật thông tin đăng ký'],
            ['displayRegisteredList()', '3', 'Hiển thị danh sách đã đăng ký'],
            ['deleteRegistration()', '4', 'Xoá đăng ký'],
            ['searchByName()', '5', 'Tìm kiếm theo ID hoặc tên'],
            ['filterByCampus()', '6', 'Lọc theo campus'],
            ['showStatistics()', '7', 'Thống kê theo đỉnh núi'],
            ['saveDataToFile()', '8', 'Lưu dữ liệu ra file'],
            ['exitProgram()', '9', 'Thoát chương trình'],
        ],
        col_widths=[5, 1.5, 8.5]
    )

    doc.add_page_break()

    # ===================== CHƯƠNG 4: HƯỚNG DẪN SỬ DỤNG =====================
    doc.add_heading('4. Hướng dẫn sử dụng từng chức năng', level=1)

    # --- 4.1. Menu chính ---
    doc.add_heading('4.1. Menu chính', level=2)
    doc.add_paragraph(
        'Khi chương trình khởi động, menu chính sẽ được hiển thị. '
        'Người dùng nhập số từ 1 đến 9 để chọn chức năng tương ứng. '
        'Nếu nhập sai (chữ, số ngoài phạm vi, ký tự đặc biệt), '
        'chương trình sẽ yêu cầu nhập lại mà không bị crash.'
    )
    add_img_placeholder(doc, "Giao diện Menu chính khi chương trình khởi động")

    doc.add_paragraph('Các chức năng trong menu:')
    menu_items = [
        ('1. New Registration', 'Đăng ký sinh viên mới tham gia leo núi.'),
        ('2. Update Registration Information', 'Cập nhật thông tin sinh viên đã đăng ký.'),
        ('3. Display Registered List', 'Hiển thị toàn bộ danh sách đăng ký dạng bảng.'),
        ('4. Delete Registration Information', 'Xoá đăng ký của một sinh viên.'),
        ('5. Search Participants by Name', 'Tìm kiếm sinh viên theo ID hoặc tên.'),
        ('6. Filter Data by Campus', 'Lọc danh sách theo campus (SE/HE/DE/QE/CE).'),
        ('7. Statistics of Registration Numbers by Location', 'Thống kê số lượng và tổng phí theo đỉnh núi.'),
        ('8. Save Data to File', 'Lưu dữ liệu ra file .dat và xuất CSV.'),
        ('9. Exit', 'Thoát chương trình (hỏi lưu nếu có dữ liệu chưa lưu).'),
    ]
    for title, desc in menu_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)

    # --- 4.2. New Registration ---
    doc.add_heading('4.2. Chức năng 1: Đăng ký mới (New Registration)', level=2)
    doc.add_paragraph('Mô tả: Cho phép nhập thông tin sinh viên mới để đăng ký tham gia leo núi.')
    doc.add_paragraph('Các bước thực hiện:')
    steps = [
        'Chọn số 1 từ menu chính.',
        'Nhập Student ID theo định dạng: SE/HE/DE/QE/CE + 6 chữ số (VD: SE203056). Nếu ID đã tồn tại → báo lỗi, yêu cầu nhập lại.',
        'Nhập Name (tên): 2-20 ký tự, chỉ chứa chữ cái và dấu cách.',
        'Nhập Phone (SĐT): 10 chữ số, bắt đầu bằng 0.',
        'Nhập Email: định dạng email hợp lệ (VD: example@gmail.com).',
        'Danh sách 13 đỉnh núi sẽ được hiển thị. Nhập mã đỉnh núi muốn đăng ký (1-13).',
        'Hệ thống tự động tính phí: 3,900,000 VND (Viettel/VNPT) hoặc 6,000,000 VND (khác).',
        'Hiển thị thông báo thành công và phí đăng ký.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Bước {i}: {step}', style='List Number')

    add_img_placeholder(doc, "Màn hình đăng ký mới – nhập đầy đủ thông tin sinh viên")
    add_img_placeholder(doc, "Kết quả đăng ký thành công với phí tính tự động")

    doc.add_paragraph('Xử lý lỗi khi đăng ký:')
    doc.add_paragraph('Nếu nhập ID sai định dạng → "Invalid format, please re-enter."', style='List Bullet')
    doc.add_paragraph('Nếu nhập ID đã tồn tại → "Student ID already exists. Please try again."', style='List Bullet')
    doc.add_paragraph('Nếu nhập tên, SĐT, email sai → "Invalid format, please re-enter."', style='List Bullet')
    doc.add_paragraph('Nếu nhập mã núi không hợp lệ → "Invalid mountain code. Please choose a code from the list."', style='List Bullet')

    add_img_placeholder(doc, "Màn hình báo lỗi khi nhập dữ liệu không hợp lệ")

    # --- 4.3. Update ---
    doc.add_heading('4.3. Chức năng 2: Cập nhật thông tin (Update Registration)', level=2)
    doc.add_paragraph('Mô tả: Cập nhật thông tin đăng ký của một sinh viên đã tồn tại.')
    doc.add_paragraph('Các bước thực hiện:')
    steps = [
        'Chọn số 2 từ menu chính.',
        'Nhập Student ID cần cập nhật.',
        'Nếu ID không tồn tại → "This student has not registered yet." → Quay về menu.',
        'Hiển thị thông tin hiện tại của sinh viên.',
        'Với mỗi trường (Name, Phone, Email, Mountain Code): nhập giá trị mới hoặc nhấn Enter để giữ nguyên.',
        'Nếu thay đổi Phone → phí đăng ký tự động được tính lại.',
        'Hiển thị thông tin đã cập nhật.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Bước {i}: {step}', style='List Number')

    add_img_placeholder(doc, "Màn hình cập nhật thông tin – hiển thị thông tin cũ và nhập mới")
    add_img_placeholder(doc, "Kết quả cập nhật thành công")

    doc.add_paragraph(
        'Lưu ý: Tính năng "nhấn Enter để giữ nguyên" rất tiện lợi khi chỉ muốn sửa 1-2 trường mà không cần nhập lại toàn bộ.'
    )

    # --- 4.4. Display ---
    doc.add_heading('4.4. Chức năng 3: Hiển thị danh sách (Display Registered List)', level=2)
    doc.add_paragraph(
        'Mô tả: Hiển thị toàn bộ danh sách sinh viên đã đăng ký dưới dạng bảng, '
        'được sắp xếp theo Student ID tăng dần.'
    )
    doc.add_paragraph('Các bước:')
    doc.add_paragraph('Bước 1: Chọn số 3 từ menu chính.', style='List Number')
    doc.add_paragraph('Bước 2: Bảng danh sách sẽ được hiển thị với các cột: StudentID, Name, Phone, Email, PeakCode, Fee.', style='List Number')
    doc.add_paragraph('Nếu chưa có ai đăng ký: "No students have registered yet."', style='List Bullet')

    add_img_placeholder(doc, "Màn hình hiển thị danh sách đăng ký đầy đủ dạng bảng")

    # --- 4.5. Delete ---
    doc.add_heading('4.5. Chức năng 4: Xoá đăng ký (Delete Registration)', level=2)
    doc.add_paragraph('Mô tả: Xoá đăng ký của một sinh viên khỏi danh sách.')
    doc.add_paragraph('Các bước:')
    steps = [
        'Chọn số 4 từ menu chính.',
        'Nhập Student ID cần xoá.',
        'Nếu ID không tồn tại → thông báo lỗi → quay về menu.',
        'Hiển thị thông tin chi tiết của sinh viên để xác nhận.',
        'Hỏi xác nhận: "Are you sure you want to delete this registration? (Y/N)".',
        'Nếu Y → xoá và thông báo "The registration has been successfully deleted."',
        'Nếu N → "Deletion cancelled." → quay về menu.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Bước {i}: {step}', style='List Number')

    add_img_placeholder(doc, "Màn hình xoá đăng ký – xác nhận trước khi xoá")
    add_img_placeholder(doc, "Kết quả sau khi xoá thành công")

    # --- 4.6. Search ---
    doc.add_heading('4.6. Chức năng 5: Tìm kiếm (Search Participants)', level=2)
    doc.add_paragraph('Mô tả: Tìm kiếm sinh viên theo Student ID (chính xác) hoặc theo tên (một phần, không phân biệt hoa thường).')
    doc.add_paragraph('Các bước:')
    steps = [
        'Chọn số 5 từ menu chính.',
        'Menu con hiện ra: 1. Search by Student ID / 2. Search by Name.',
        'Nếu chọn 1: Nhập Student ID → tìm chính xác → hiển thị kết quả hoặc "No one matches the search criteria!"',
        'Nếu chọn 2: Nhập (một phần) tên → tìm tất cả sinh viên có tên chứa chuỗi nhập (case-insensitive) → hiển thị bảng kết quả.',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'Bước {i}: {step}', style='List Number')

    add_img_placeholder(doc, "Màn hình tìm kiếm theo ID – tìm thấy kết quả")
    add_img_placeholder(doc, "Màn hình tìm kiếm theo tên – nhiều kết quả khớp")
    add_img_placeholder(doc, "Màn hình tìm kiếm không tìm thấy kết quả")

    # --- 4.7. Filter ---
    doc.add_heading('4.7. Chức năng 6: Lọc theo Campus (Filter by Campus)', level=2)
    doc.add_paragraph('Mô tả: Lọc danh sách sinh viên theo mã campus (2 ký tự đầu của Student ID).')

    add_styled_table(doc,
        ['Mã Campus', 'Tên Campus'],
        [
            ['CE', 'FPT Cần Thơ'],
            ['DE', 'FPT Đà Nẵng'],
            ['HE', 'FPT Hà Nội'],
            ['SE', 'FPT Hồ Chí Minh'],
            ['QE', 'FPT Quy Nhơn'],
        ],
        col_widths=[3, 12]
    )

    doc.add_paragraph()
    doc.add_paragraph('Các bước:')
    doc.add_paragraph('Bước 1: Chọn số 6 từ menu chính.', style='List Number')
    doc.add_paragraph('Bước 2: Nhập mã campus (VD: SE).', style='List Number')
    doc.add_paragraph('Bước 3: Hiển thị danh sách sinh viên thuộc campus đó, hoặc "No students have registered under this campus." nếu không có.', style='List Number')

    add_img_placeholder(doc, "Màn hình lọc theo campus – hiển thị kết quả")

    # --- 4.8. Statistics ---
    doc.add_heading('4.8. Chức năng 7: Thống kê (Statistics by Mountain)', level=2)
    doc.add_paragraph(
        'Mô tả: Thống kê số lượng sinh viên đăng ký và tổng phí theo từng đỉnh núi.'
    )
    doc.add_paragraph('Các bước:')
    doc.add_paragraph('Bước 1: Chọn số 7 từ menu chính.', style='List Number')
    doc.add_paragraph('Bước 2: Bảng thống kê hiển thị: Mã núi, Tên núi, Số người đăng ký, Tổng phí.', style='List Number')
    doc.add_paragraph('Nếu không có dữ liệu: "No registration data available for statistics."', style='List Bullet')

    add_img_placeholder(doc, "Màn hình thống kê theo đỉnh núi – bảng kết quả")

    # --- 4.9. Save ---
    doc.add_heading('4.9. Chức năng 8: Lưu dữ liệu (Save Data to File)', level=2)
    doc.add_paragraph(
        'Mô tả: Lưu toàn bộ dữ liệu đăng ký vào 2 file:'
    )
    doc.add_paragraph('registrations.dat – File nhị phân (Binary), sử dụng ObjectOutputStream để serialize ArrayList<Student>.', style='List Bullet')
    doc.add_paragraph('registrations.csv – File CSV dạng text, dễ mở bằng Excel, gồm các cột: StudentID, Name, Phone, Email, MountainCode, TuitionFee.', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Các bước:')
    doc.add_paragraph('Bước 1: Chọn số 8 từ menu chính.', style='List Number')
    doc.add_paragraph('Bước 2: Hệ thống tự động lưu và thông báo "Registration data has been successfully saved."', style='List Number')

    add_img_placeholder(doc, "Màn hình lưu dữ liệu thành công")
    add_img_placeholder(doc, "Nội dung file registrations.csv mở bằng Excel hoặc Notepad")

    # --- 4.10. Exit ---
    doc.add_heading('4.10. Chức năng 9: Thoát chương trình (Exit)', level=2)
    doc.add_paragraph('Mô tả: Thoát chương trình với kiểm tra dữ liệu chưa lưu.')
    doc.add_paragraph('Luồng xử lý khi thoát:')
    exit_steps = [
        'Chọn số 9 từ menu chính.',
        'Nếu KHÔNG có dữ liệu chưa lưu → Thoát ngay, hiển thị "Goodbye!".',
        'Nếu CÓ dữ liệu chưa lưu:',
        '    a. Hỏi: "You have unsaved changes. Do you want to save before exiting? (Y/N)"',
        '    b. Nếu Y → Lưu file → Thoát.',
        '    c. Nếu N → Hỏi tiếp: "Are you sure you want to exit without saving? (Y/N)"',
        '    d. Nếu Y → Thoát không lưu.',
        '    e. Nếu N → Quay lại menu (không thoát).',
    ]
    for step in exit_steps:
        doc.add_paragraph(step, style='List Number')

    add_img_placeholder(doc, "Màn hình thoát – hỏi lưu khi có dữ liệu chưa lưu")
    add_img_placeholder(doc, "Luồng thoát: chọn không lưu → xác nhận lần 2")

    doc.add_page_break()

    # ===================== CHƯƠNG 5: LUỒNG HOẠT ĐỘNG =====================
    doc.add_heading('5. Luồng hoạt động tổng thể', level=1)

    doc.add_heading('5.1. Flowchart tổng quan chương trình', level=2)
    doc.add_paragraph(
        'Flowchart bên dưới mô tả luồng hoạt động tổng quan của chương trình từ khi bắt đầu đến khi kết thúc. '
        'Chương trình hoạt động theo mô hình menu-driven (lặp menu), nghĩa là sau khi thực hiện xong một chức năng, '
        'hệ thống sẽ quay lại menu chính cho đến khi người dùng chọn Exit.'
    )
    add_img_placeholder(doc, "Flowchart tổng quan – Toàn bộ luồng hoạt động của chương trình")

    doc.add_paragraph('Giải thích luồng:')
    doc.add_paragraph(
        '• Start → Load MountainList.csv vào Mountains → Load registrations.dat vào Students (nếu file tồn tại) '
        '→ Hiển thị Menu → Nhập lựa chọn → Thực thi chức năng → Quay lại Menu (lặp) → '
        'Khi chọn Exit → Kiểm tra unsaved → Hỏi lưu → End.'
    )

    doc.add_heading('5.2. Flowchart từng chức năng', level=2)
    doc.add_paragraph(
        'Bên dưới là flowchart chi tiết cho từng chức năng. '
        'Mỗi flowchart thể hiện rõ luồng xử lý, các bước validate, và các trường hợp lỗi.'
    )

    functions_flow = [
        ("Chức năng 1 – New Registration",
         "Flowchart đăng ký mới: Nhập ID → Kiểm tra trùng → Nhập Name/Phone/Email → Validate → Hiển thị danh sách núi → Nhập Mountain Code → Tính phí → Thêm vào danh sách"),
        ("Chức năng 2 – Update Registration",
         "Flowchart cập nhật: Nhập ID → Tìm SV → Hiển thị thông tin cũ → Nhập từng trường mới (Enter để giữ) → Validate → Tính lại phí nếu đổi SĐT → Cập nhật"),
        ("Chức năng 3 – Display Registered List",
         "Flowchart hiển thị: Kiểm tra danh sách rỗng → Sort theo ID → In bảng"),
        ("Chức năng 4 – Delete Registration",
         "Flowchart xoá: Nhập ID → Tìm SV → Hiển thị thông tin → Xác nhận Y/N → Xoá hoặc huỷ"),
        ("Chức năng 5 – Search Participants",
         "Flowchart tìm kiếm: Chọn tìm theo ID hoặc Tên → Nhập keyword → Tìm → Hiển thị kết quả"),
        ("Chức năng 6 – Filter by Campus",
         "Flowchart lọc: Nhập mã campus → Validate → Lọc theo 2 ký tự đầu ID → Hiển thị"),
        ("Chức năng 7 – Statistics",
         "Flowchart thống kê: Duyệt danh sách SV → Nhóm theo mountainCode → Đếm số lượng + tổng phí → Hiển thị bảng"),
        ("Chức năng 8 – Save Data",
         "Flowchart lưu file: Serialize → ObjectOutputStream → registrations.dat + Xuất CSV"),
        ("Chức năng 9 – Exit",
         "Flowchart thoát: Kiểm tra isSaved → Nếu false: hỏi lưu (Y/N) → Nếu N: hỏi xác nhận → Thoát hoặc quay lại menu"),
    ]
    for title, desc in functions_flow:
        p = doc.add_paragraph()
        run = p.add_run(f"▶ {title}")
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(desc)
        add_img_placeholder(doc, f"Flowchart – {title}")
        doc.add_paragraph()

    doc.add_heading('5.3. Sequence Diagram – Đăng ký mới', level=2)
    doc.add_paragraph(
        'Sequence Diagram mô tả trình tự tương tác giữa các đối tượng khi thực hiện chức năng Đăng ký mới (New Registration). '
        'Các đối tượng tham gia: Operator, Main, Inputter, Mountains, Students, Student.'
    )
    add_img_placeholder(doc, "Sequence Diagram – Luồng đăng ký mới (New Registration)")

    doc.add_paragraph('Giải thích Sequence Diagram:')
    seq_steps = [
        'Operator chọn New Registration từ menu.',
        'Main gọi Inputter để nhập và validate ID, Name, Phone, Email.',
        'Main gọi Students.searchById() để kiểm tra ID trùng.',
        'Main gọi Mountains.isValidMountainCode() để kiểm tra mã núi hợp lệ.',
        'Main tạo đối tượng Student mới và gọi calculateFee() để tính phí.',
        'Main gọi Students.add() để thêm sinh viên vào danh sách, đánh dấu unsaved.',
        'Main hiển thị thông báo thành công cho Operator.',
    ]
    for i, step in enumerate(seq_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_page_break()

    # ===================== CHƯƠNG 6: DỮ LIỆU MẪU =====================
    doc.add_heading('6. Dữ liệu mẫu', level=1)

    doc.add_heading('6.1. File MountainList.csv', level=2)
    doc.add_paragraph(
        'File MountainList.csv chứa danh sách 13 đỉnh núi, được đọc khi khởi động chương trình. '
        'Cấu trúc: Code, Mountain, Province, Description.'
    )
    mountains_data = [
        ['1', 'Ham Rong Mountain', 'Lao Cai', 'Gần trung tâm Sa Pa, cao 1850m'],
        ['2', 'Doi Bo Mountain', 'Lao Cai', ''],
        ['3', 'Pha Luong Mountain', 'Son La', 'Cao gần 2000m, mái nhà Mộc Châu'],
        ['4', 'Hon Vuon Mountain', 'Hue', ''],
        ['5', 'Da Do Mountain', 'Ninh Thuan', ''],
        ['6', 'Da Bia Mountain', 'Phu Yen', ''],
        ['7', 'Chu Hreng Mountain', 'Kon Tum', ''],
        ['8', 'Lang Biang Mountain', 'Lam Dong', ''],
        ['9', 'Ta Nang Mountain', 'Lam Dong', ''],
        ['10', 'Cam Mountain', 'An Giang', ''],
        ['11', 'Thi Vai Mountain', 'Vung Tau', ''],
        ['12', 'Dinh Mountain', 'Vung Tau', ''],
        ['13', 'Co Tien Mountain', 'Khanh Hoa', 'Có 3 đỉnh, chỉ cần chinh phục đỉnh 1'],
    ]
    add_styled_table(doc,
        ['Code', 'Mountain', 'Province', 'Description'],
        mountains_data,
        col_widths=[1.5, 5, 3, 5.5]
    )

    doc.add_heading('6.2. File registrations.csv', level=2)
    doc.add_paragraph(
        'File registrations.csv được xuất ra khi người dùng chọn Save Data. '
        'Đây là file dạng text có thể mở bằng Excel.'
    )
    add_styled_table(doc,
        ['StudentID', 'Name', 'Phone', 'Email', 'MountainCode', 'TuitionFee'],
        [
            ['SE203056', 'Nguyen Phu Khuong', '0363561629', 'khuong@gmail.com', '1', '3,900,000'],
        ],
        col_widths=[2.5, 4, 2.5, 4, 2.5, 2]
    )

    doc.add_page_break()

    # ===================== CHƯƠNG 7: QUY TẮC VALIDATION =====================
    doc.add_heading('7. Quy tắc validation dữ liệu', level=1)
    doc.add_paragraph(
        'Hệ thống áp dụng kiểm tra dữ liệu đầu vào (validation) nghiêm ngặt bằng Regular Expression (Regex). '
        'Bảng dưới đây tổng hợp các quy tắc validation cho từng trường dữ liệu.'
    )
    add_styled_table(doc,
        ['Trường', 'Quy tắc', 'Ví dụ hợp lệ', 'Ví dụ không hợp lệ'],
        [
            ['Student ID', '2 ký tự campus (SE/HE/DE/QE/CE)\n+ 6 chữ số', 'SE203056\nHE180001', 'AB123456\nSE12345\nSE1234567'],
            ['Name', '2-20 ký tự, chỉ chữ cái\nvà dấu cách', 'Nguyen Van A\nLe Thi B', 'A\n12345\nTên quá dài hơn 20 ký tự abc'],
            ['Phone', '10 chữ số, bắt đầu bằng 0', '0363561629\n0912345678', '123456789\n0123\nabcdefghij'],
            ['Email', 'user@domain.ext\n(chữ, số, +, _, ., -)', 'abc@gmail.com\ntest_1@fpt.edu.vn', '@gmail.com\nabc@\nabc.gmail.com'],
            ['Mountain Code', 'Phải tồn tại trong danh sách\nMountainList.csv (1-13)', '1\n5\n13', '0\n14\nabc'],
            ['Campus Code', 'SE, HE, DE, QE, hoặc CE\n(không phân biệt hoa thường)', 'SE\nhe\nDE', 'AB\nSF\n123'],
            ['Menu Choice', 'Số nguyên từ 1 đến 9', '1\n5\n9', '0\n10\nabc'],
            ['Confirm (Y/N)', 'Y hoặc N\n(không phân biệt hoa thường)', 'Y\nn', 'yes\nno\n1'],
        ],
        col_widths=[3, 4, 4, 4]
    )

    doc.add_paragraph()
    doc.add_paragraph('Quy tắc tính phí đăng ký:')
    add_styled_table(doc,
        ['Nhà mạng', 'Đầu số', 'Phí đăng ký'],
        [
            ['Viettel', '032, 033, 034, 035, 036, 037, 038, 039, 096, 097, 098, 086', '3,900,000 VND (giảm 35%)'],
            ['VNPT', '081, 082, 083, 084, 085, 088, 091, 094', '3,900,000 VND (giảm 35%)'],
            ['Khác (Mobifone, ...)', 'Các đầu số khác', '6,000,000 VND (không giảm)'],
        ],
        col_widths=[3.5, 8, 3.5]
    )

    doc.add_page_break()

    # ===================== CHƯƠNG 8: KẾT LUẬN =====================
    doc.add_heading('8. Kết luận', level=1)

    doc.add_heading('8.1. Kết quả đạt được', level=2)
    results = [
        'Hoàn thành đầy đủ 9 chức năng theo yêu cầu đề bài LAB211.',
        'Áp dụng lập trình hướng đối tượng (OOP) với cấu trúc 9 class/interface rõ ràng.',
        'Validate dữ liệu đầu vào chặt chẽ bằng Regex, chương trình không bị crash khi nhập sai.',
        'Tính phí đăng ký tự động dựa trên nhà mạng SĐT (Viettel/VNPT giảm 35%).',
        'Đọc dữ liệu núi từ file CSV, lưu đăng ký ra file nhị phân (.dat) và xuất CSV.',
        'Xử lý thoát chương trình thông minh: cảnh báo khi có dữ liệu chưa lưu.',
        'Hiển thị dữ liệu dạng bảng có header, đường kẻ, sắp xếp theo ID.',
        'Code đạt trên 200 LOC, đặt tên biến/phương thức rõ ràng theo convention.',
    ]
    for r in results:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading('8.2. Kiến thức áp dụng', level=2)
    knowledge = [
        'OOP: Class, Interface, Kế thừa, Đa hình, Đóng gói (Encapsulation).',
        'Java Collections: ArrayList, HashMap, Collections.sort().',
        'Java I/O: BufferedReader, FileReader, ObjectInputStream/ObjectOutputStream, BufferedWriter.',
        'Serialization: Lưu trữ đối tượng Java ra file nhị phân.',
        'Comparable: Sắp xếp đối tượng theo tiêu chí tuỳ chỉnh.',
        'Regex: Kiểm tra định dạng dữ liệu đầu vào.',
        'Design Pattern: Menu-driven architecture, Separation of Concerns.',
    ]
    for k in knowledge:
        doc.add_paragraph(k, style='List Bullet')

    doc.add_heading('8.3. Hạn chế và hướng phát triển', level=2)
    doc.add_paragraph('Hạn chế:')
    limitations = [
        'Chưa có giao diện đồ hoạ (GUI), chỉ là ứng dụng console.',
        'Chưa hỗ trợ đa ngôn ngữ (Tiếng Việt có dấu trong tên có thể gặp vấn đề encoding).',
        'Chưa có unit test tự động.',
    ]
    for l in limitations:
        doc.add_paragraph(l, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Hướng phát triển:')
    improvements = [
        'Xây dựng giao diện đồ hoạ bằng JavaFX hoặc Swing.',
        'Kết nối cơ sở dữ liệu (MySQL, SQLite) thay thế file.',
        'Thêm chức năng xuất báo cáo PDF.',
        'Viết unit test bằng JUnit.',
        'Hỗ trợ đa ngôn ngữ và encoding UTF-8 đầy đủ.',
    ]
    for imp in improvements:
        doc.add_paragraph(imp, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('--- Hết báo cáo ---')
    run.bold = True
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # ===================== SAVE =====================
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'BaoCao_MountainHiking_LAB211.docx'
    )
    doc.save(output_path)
    print(f"[OK] Report created: {output_path}")
    return output_path

if __name__ == '__main__':
    create_report()
