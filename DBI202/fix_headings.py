"""
Script to fix missing Heading 2 and Heading 3 styles in final_report.docx.
Uses text-based search within ranges to find correct paragraph indices.
"""

import docx
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import sys, io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

INPUT_FILE = r"final_report.docx"
OUTPUT_FILE = r"final_report.docx"

doc = docx.Document(INPUT_FILE)

def set_heading_style(para, level):
    """Set paragraph style to Heading 1/2/3."""
    style_id = f"Heading{level}"
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml('<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        para._element.insert(0, pPr)
    
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = parse_xml(f'<w:pStyle xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{style_id}"/>')
        pPr.insert(0, pStyle)
    else:
        pStyle.set(qn('w:val'), style_id)

def find_and_fix(text_prefix, search_start, search_end, heading_level):
    """Find a Normal-styled paragraph by text prefix within a range and set heading."""
    for i in range(search_start, min(search_end, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if t.startswith(text_prefix) and p.style.name == 'Normal':
            set_heading_style(p, heading_level)
            print(f"[OK] [{i}] -> Heading {heading_level} | {t[:80]}")
            return True
    print(f"[SKIP] Not found: '{text_prefix}' in range [{search_start}-{search_end})")
    return False

count = 0

# ================================================================
# LAB 1 (I. Lab 1 heading is around para 29, content starts ~30)
# Main sections: Heading 2 | Sub-sections (2.x): Heading 3
# ================================================================
print("=== LAB 1 ===")
count += find_and_fix("1. Objective", 29, 35, 2)
count += find_and_fix("2. Study of Data Models", 29, 40, 2)
count += find_and_fix("2.1. Relational Data Model", 29, 50, 3)
count += find_and_fix("2.2. Semi-structured Data Model", 40, 65, 3)
count += find_and_fix("2.3. NoSQL Models", 55, 70, 3)
count += find_and_fix("3. Model Selection", 60, 70, 2)
count += find_and_fix("4. Conclusion", 65, 75, 2)

# ================================================================
# LAB 2 (II. Lab 2 heading is around para 70, content starts ~73)
# Main sections: Heading 2 | Sub-sections (2.x): Heading 3
# ================================================================
print("\n=== LAB 2 ===")
count += find_and_fix("1. Objective", 70, 78, 2)
count += find_and_fix("2. Entity Analysis", 73, 80, 2)
count += find_and_fix("2.1. User Management", 75, 85, 3)
count += find_and_fix("2.2. Product", 88, 100, 3)
count += find_and_fix("2.3. Sales", 100, 120, 3)
count += find_and_fix("2.4. Maintenance", 130, 155, 3)
count += find_and_fix("3. Conclusion", 160, 175, 2)

# ================================================================
# LAB 3 (III. Lab 3 heading is around para 176-186)
# Most headings already correct, just "6. Conclusion" missing
# ================================================================
print("\n=== LAB 3 ===")
count += find_and_fix("6. Conclusion", 435, 450, 2)

# ================================================================
# LAB 4 - Already has proper headings
# ================================================================
print("\n=== LAB 4 ===")
print("[OK] Lab 4 headings already correct")

# ================================================================
# LAB 5 (V. Lab 5 heading is around para 588)
# "3.3. Advanced Queries" is missing Heading 2
# ================================================================
print("\n=== LAB 5 ===")
count += find_and_fix("3.3. Advanced Queries", 895, 910, 2)

# ================================================================
print(f"\nTotal heading fixes applied: {count}")

# Save
doc.save(OUTPUT_FILE)
print(f"[OK] Saved to {OUTPUT_FILE}")
print("\nRemember to update the Table of Contents in Word:")
print("  Right-click TOC -> Update Field -> Update entire table")
