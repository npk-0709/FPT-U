"""
Generate Lab 3 Report: Anomaly Detection and Normalization
for Household Cleaning Robot Sales & Maintenance Management System
FIXED: proper line breaks, removed Section 2, clean formatting
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.54)

FONT = 'Times New Roman'
SZ_NORMAL = 12
SZ_SMALL = 11

# ── Helper Functions ──
def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def add_para(doc, text, bold=False, italic=False, font_size=SZ_NORMAL,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, space_before=0,
             first_line_indent=None):
    """Add a single-line paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = FONT
    return p

def add_multiline_para(doc, lines, bold_first=False, font_size=SZ_SMALL,
                       space_after=6, indent=None):
    """Add a paragraph with multiple lines using proper Word line breaks."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    
    for i, line in enumerate(lines):
        if i > 0:
            # Add a line break before each subsequent line
            run_br = p.add_run()
            run_br.font.size = Pt(font_size)
            run_br.font.name = FONT
            run_br.add_break()
        
        run = p.add_run(line)
        run.font.size = Pt(font_size)
        run.font.name = FONT
        if bold_first and i == 0:
            run.bold = True
    return p

def add_bullet(doc, text, level=0, bold=False, font_size=SZ_SMALL):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = FONT
    return p

def add_bullet_multiline(doc, lines, level=0, font_size=SZ_SMALL):
    """Bullet with multiple lines using line breaks."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for i, line in enumerate(lines):
        if i > 0:
            run_br = p.add_run()
            run_br.font.size = Pt(font_size)
            run_br.font.name = FONT
            run_br.add_break()
        run = p.add_run(line)
        run.font.size = Pt(font_size)
        run.font.name = FONT
    return p

def heading1(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.name = FONT
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def heading2(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.name = FONT
        run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    return h

def create_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = FONT
        set_cell_shading(cell, 'D9E2F3')
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = FONT
    return table


# ══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

add_para(doc, 'FPT UNIVERSITY', bold=True, font_size=18,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para(doc, 'DEPARTMENT OF INFORMATION TECHNOLOGY', bold=True, font_size=14,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

doc.add_paragraph()

add_para(doc, 'DBI202 – Database Systems', bold=True, font_size=16,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para(doc, 'Lab 3: Anomaly Detection and Normalization', bold=True, font_size=20,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para(doc, 'Household Cleaning Robot Sales & Maintenance Management System',
         italic=True, font_size=14,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

for _ in range(2):
    doc.add_paragraph()

add_para(doc, 'Group Members:', bold=True, font_size=13,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
for i in range(1, 6):
    add_para(doc, f'[Member {i} – Student ID]', font_size=12,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)

add_para(doc, '', font_size=12, space_after=12)
add_para(doc, 'Date: June 2026', bold=True, font_size=12,
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════
heading1(doc, 'Table of Contents')
toc = [
    '1. Objective',
    '2. Analysis of Anomalies',
    '    2.1. Redundancy Analysis',
    '    2.2. Update Anomalies',
    '    2.3. Insertion Anomalies',
    '    2.4. Deletion Anomalies',
    '3. Identification of Normal Forms',
    '    3.1. First Normal Form (1NF)',
    '    3.2. Second Normal Form (2NF)',
    '    3.3. Third Normal Form (3NF)',
    '    3.4. Boyce–Codd Normal Form (BCNF)',
    '4. Step-by-Step Decomposition',
    '5. Final Normalized Schema',
    '6. Conclusion',
]
for item in toc:
    add_para(doc, item, font_size=12, space_after=3, alignment=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 1. OBJECTIVE
# ══════════════════════════════════════════════════════════════════════
heading1(doc, '1. Objective')

add_para(doc,
    'The objective of this lab is to identify anomalies in the relational database schema '
    'designed in Lab 2 for the "Household Cleaning Robot Sales & Maintenance Management System" '
    'and to apply normalization techniques through systematic decomposition. This report will:')

add_bullet(doc, 'Detect redundancy, insertion, update, and deletion anomalies in the existing schema.')
add_bullet(doc, 'Determine whether each relation satisfies 1NF, 2NF, 3NF, and BCNF.')
add_bullet(doc, 'Decompose relations step by step into higher normal forms (up to BCNF).')
add_bullet(doc, 'Ensure lossless join and dependency preservation properties.')
add_bullet(doc, 'Present the final normalized schema and demonstrate how anomalies are eliminated.')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
# 2. ANALYSIS OF ANOMALIES
# ══════════════════════════════════════════════════════════════════════
heading1(doc, '2. Analysis of Anomalies')

add_para(doc,
    'This section systematically examines the Lab 2 schema for four types of anomalies: '
    'redundancy, update anomalies, insertion anomalies, and deletion anomalies. The analysis '
    'focuses on relations that exhibit potential issues due to their functional dependency structures.')

# ── 2.1 Redundancy ──
heading2(doc, '2.1. Redundancy Analysis')

add_para(doc,
    'Redundancy occurs when the same piece of information is stored multiple times across tuples. '
    'We examine each relation for data redundancy stemming from functional dependencies.')

add_para(doc, 'Relation: OrderDetail', bold=True, space_after=3)
add_para(doc,
    'The OrderDetail relation has the composite primary key {OrderID, RobotID} with the FD: '
    '{OrderID, RobotID} → SellingPrice. However, there also exists the FD: RobotID → OrderID, '
    'meaning each robot unit can only be sold once. This implies that RobotID alone is a candidate key. '
    'The inclusion of OrderID in the key creates a situation where OrderID is redundantly stored.',
    font_size=SZ_SMALL)

add_para(doc, 'Example:', bold=True, italic=True, font_size=SZ_SMALL, space_after=3)
create_table(doc,
    ['OrderID', 'RobotID', 'SellingPrice'],
    [
        ['ORD001', 'RBT101', '15,000,000 VND'],
        ['ORD001', 'RBT102', '18,500,000 VND'],
        ['ORD002', 'RBT103', '22,000,000 VND'],
    ])
doc.add_paragraph()

add_para(doc,
    'Since RobotID → OrderID, the OrderID value "ORD001" is derivable from each RobotID. '
    'The composite key carries redundant information because a single-attribute key (RobotID) suffices.',
    font_size=SZ_SMALL)

add_para(doc, 'Relation: Payment', bold=True, space_after=3, space_before=6)
add_para(doc,
    'The Payment relation stores both OrderID and ServiceRecordID as nullable foreign keys. '
    'For a sales payment, ServiceRecordID is NULL; for a service payment, OrderID is NULL. '
    'This dual-purpose design introduces NULL-value redundancy.',
    font_size=SZ_SMALL)

add_para(doc, 'Example:', bold=True, italic=True, font_size=SZ_SMALL, space_after=3)
create_table(doc,
    ['PaymentID', 'OrderID', 'ServiceRecordID', 'Amount', 'PaymentDate', 'Method'],
    [
        ['PAY001', 'ORD001', 'NULL', '33,500,000', '2026-01-15', 'Credit Card'],
        ['PAY002', 'NULL', 'REC005', '500,000', '2026-02-20', 'Cash'],
        ['PAY003', 'ORD002', 'NULL', '22,000,000', '2026-03-01', 'Bank Transfer'],
    ])
doc.add_paragraph()

add_para(doc,
    'Most other relations (Customer, Employee, RobotModel, RobotUnit, SalesOrder, '
    'WarrantyRegistration, ServiceRequest, MaintenanceRecord, DeviceLog) have FDs where '
    'every non-key attribute is fully and directly determined by the primary key alone, '
    'resulting in minimal redundancy.',
    font_size=SZ_SMALL)

# ── 2.2 Update Anomalies ──
heading2(doc, '2.2. Update Anomalies')

add_para(doc,
    'Update anomalies occur when modifying data in one tuple requires changes in multiple '
    'tuples to maintain consistency.')

add_para(doc, 'Relation: OrderDetail', bold=True, space_after=3)
add_para(doc,
    'Since RobotID → OrderID, if the OrderID of a particular robot sale needs to be corrected '
    '(e.g., due to an order merge), updating it in the OrderDetail table must be done carefully. '
    'Any inconsistency between the stored OrderID and the value implied by RobotID violates '
    'the functional dependency, resulting in an update anomaly.',
    font_size=SZ_SMALL)

add_para(doc, 'Example:', bold=True, italic=True, font_size=SZ_SMALL, space_after=3)
add_para(doc,
    'If RBT101 is reassigned from ORD001 to ORD003, the update must ensure that the FD '
    'RobotID → OrderID remains consistent. Failing to update all references would create '
    'contradictory data.',
    font_size=SZ_SMALL)

add_para(doc, 'Relation: Payment (structural update anomaly)', bold=True, space_after=3, space_before=6)
add_para(doc,
    'If a payment is initially recorded as a sales payment (OrderID = ORD001, ServiceRecordID = NULL) '
    'but later needs to be reclassified as a service payment, both OrderID and ServiceRecordID must '
    'be updated simultaneously. A partial update would leave the record in an inconsistent state.',
    font_size=SZ_SMALL)

# ── 2.3 Insertion Anomalies ──
heading2(doc, '2.3. Insertion Anomalies')

add_para(doc,
    'Insertion anomalies occur when certain data cannot be inserted into the database without '
    'the presence of other unrelated data.')

add_para(doc, 'Relation: OrderDetail', bold=True, space_after=3)
add_para(doc,
    'With the primary key {OrderID, RobotID}, it is impossible to record a robot\'s selling price '
    'without first creating a sales order. If a business wants to pre-set selling prices for '
    'specific robot units before any order is placed, this cannot be done within the OrderDetail '
    'relation without an associated OrderID.',
    font_size=SZ_SMALL)

add_para(doc, 'Relation: MaintenanceRecord', bold=True, space_after=3, space_before=6)
add_para(doc,
    'The FD RequestID → RecordID implies a 1:1 relationship between ServiceRequest and '
    'MaintenanceRecord. Standalone maintenance activities (e.g., proactive scheduled maintenance '
    'not initiated by a customer complaint) cannot be recorded without first fabricating a '
    'ServiceRequest entry.',
    font_size=SZ_SMALL)

add_para(doc, 'Relation: Payment', bold=True, space_after=3, space_before=6)
add_para(doc,
    'A payment cannot be recorded unless it is linked to either an OrderID or a ServiceRecordID. '
    'If the business receives a deposit or advance payment not yet tied to a specific order '
    'or service, the current schema cannot accommodate it.',
    font_size=SZ_SMALL)

# ── 2.4 Deletion Anomalies ──
heading2(doc, '2.4. Deletion Anomalies')

add_para(doc,
    'Deletion anomalies occur when removing a tuple causes unintended loss of other useful information.')

add_para(doc, 'Relation: OrderDetail', bold=True, space_after=3)
add_para(doc,
    'If all OrderDetail rows for a specific OrderID are deleted (e.g., an order is cancelled), '
    'the selling price information for those specific robot units is permanently lost.',
    font_size=SZ_SMALL)

add_para(doc, 'Relation: WarrantyRegistration', bold=True, space_after=3, space_before=6)
add_para(doc,
    'Since RobotID → WarrantyID (each robot has exactly one warranty), deleting a warranty '
    'registration entry removes the warranty period information. There is no status flag to mark '
    'it as expired — deletion causes loss of historical warranty data.',
    font_size=SZ_SMALL)

add_para(doc, 'Relation: ServiceRequest + MaintenanceRecord', bold=True, space_after=3, space_before=6)
add_para(doc,
    'Due to the 1:1 relationship (RequestID → RecordID), deleting a service request also '
    'necessitates deleting its associated maintenance record. The detailed technical resolution data '
    '(actions taken, parts replaced, service fee) would be lost along with the request.',
    font_size=SZ_SMALL)

# Summary table
add_para(doc, 'Summary of Anomalies by Relation', bold=True, space_after=6, space_before=12)
create_table(doc,
    ['Relation', 'Redundancy', 'Update', 'Insertion', 'Deletion'],
    [
        ['Customer', 'None', 'None', 'None', 'None'],
        ['Employee', 'None', 'None', 'None', 'None'],
        ['RobotModel', 'None', 'None', 'None', 'None'],
        ['RobotUnit', 'None', 'None', 'None', 'None'],
        ['SalesOrder', 'None', 'None', 'None', 'None'],
        ['OrderDetail', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Payment', 'Structural', 'Yes', 'Yes', 'Minor'],
        ['WarrantyReg.', 'None', 'None', 'None', 'Yes'],
        ['ServiceRequest', 'None', 'None', 'None', 'Yes (cascade)'],
        ['Maintenance Rec.', 'None', 'None', 'Yes', 'Yes (cascade)'],
        ['DeviceLog', 'None', 'None', 'None', 'None'],
    ])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 3. IDENTIFICATION OF NORMAL FORMS
# ══════════════════════════════════════════════════════════════════════
heading1(doc, '3. Identification of Normal Forms')

add_para(doc,
    'This section examines each relation against the criteria for 1NF, 2NF, 3NF, and BCNF. '
    'We use the functional dependencies established in Lab 2 to rigorously justify the '
    'normal form classification of each relation.')

# ── 3.1 1NF ──
heading2(doc, '3.1. First Normal Form (1NF)')

add_para(doc,
    'A relation is in 1NF if and only if every attribute contains only atomic (indivisible) '
    'values and there are no repeating groups.',
    font_size=SZ_SMALL)

add_para(doc, 'Analysis:', bold=True, space_after=6)

create_table(doc,
    ['Relation', '1NF?', 'Justification'],
    [
        ['Customer', 'Yes', 'All attributes are atomic. No multi-valued attributes.'],
        ['Employee', 'Yes', 'All attributes are atomic. Role is a single value.'],
        ['RobotModel', 'Issue*', 'Features and Specifications may be multi-valued.'],
        ['RobotUnit', 'Yes', 'All attributes are atomic. Status is single-valued.'],
        ['SalesOrder', 'Yes', 'All attributes are atomic single values.'],
        ['OrderDetail', 'Yes', 'All attributes are atomic.'],
        ['Payment', 'Yes', 'All attributes are atomic single values.'],
        ['WarrantyReg.', 'Yes', 'All attributes are atomic date/ID values.'],
        ['ServiceRequest', 'Yes', 'IssueDescription is a single text field.'],
        ['Maintenance Rec.', 'Issue*', 'PartsReplaced may contain multiple part names.'],
        ['DeviceLog', 'Issue*', 'UsageStatistics may store structured/multiple data.'],
    ])
doc.add_paragraph()

add_para(doc,
    '* Note: For this analysis, we treat Features, Specifications, PartsReplaced, and '
    'UsageStatistics as atomic text fields. If these store multiple independent values, '
    'decomposition is required (addressed in Section 4). Under current design, all relations satisfy 1NF.',
    font_size=10, italic=True)

add_para(doc, 'Conclusion: All 11 relations satisfy 1NF under current design assumptions.',
         bold=True, font_size=SZ_SMALL, space_after=6)

# ── 3.2 2NF ──
heading2(doc, '3.2. Second Normal Form (2NF)')

add_para(doc,
    'A relation is in 2NF if it is in 1NF and every non-prime attribute is fully functionally '
    'dependent on every candidate key (no partial dependencies). Partial dependencies can only '
    'occur when a candidate key is composite.',
    font_size=SZ_SMALL)

add_para(doc, 'Relations with single-attribute candidate keys only:', bold=True, space_after=3)
add_para(doc,
    'Customer, Employee, RobotModel, RobotUnit, SalesOrder, Payment, WarrantyRegistration, '
    'ServiceRequest, and MaintenanceRecord all have single-attribute candidate keys. '
    'Since partial dependency requires a composite key, these are automatically in 2NF.',
    font_size=SZ_SMALL)

add_para(doc, 'Relations with composite candidate keys — detailed analysis:',
         bold=True, space_after=6, space_before=8)

# OrderDetail 2NF
add_para(doc, 'Relation: OrderDetail ({OrderID, RobotID}, SellingPrice)', bold=True, space_after=3)

add_multiline_para(doc, [
    'Candidate Keys: RobotID and {OrderID, RobotID}',
    'FDs: {OrderID, RobotID} → SellingPrice;  RobotID → OrderID',
], font_size=SZ_SMALL, space_after=4, indent=0.5)

add_para(doc,
    'Since RobotID is a candidate key by itself, every non-prime attribute (SellingPrice) is '
    'determined by RobotID alone. The composite key {OrderID, RobotID} contains a proper subset '
    '(RobotID) that is itself a candidate key.',
    font_size=SZ_SMALL)

add_para(doc, 'Checking partial dependency on {OrderID, RobotID}:', bold=True, font_size=SZ_SMALL, space_after=3)
add_bullet(doc, 'RobotID → OrderID: OrderID is part of the composite CK, hence a prime attribute. Not a partial dependency of a non-prime attribute.')
add_bullet(doc, 'RobotID → SellingPrice: Since RobotID itself is a CK, SellingPrice is fully dependent on this key.')

add_para(doc,
    'Since the only non-prime attribute SellingPrice is fully dependent on the candidate key '
    'RobotID, there is no partial dependency. OrderDetail is in 2NF. ✓',
    font_size=SZ_SMALL, space_after=8)

# DeviceLog 2NF
add_para(doc, 'Relation: DeviceLog (LogID, RobotID, LogTime, ErrorCode, UsageStatistics)',
         bold=True, space_after=3)

add_multiline_para(doc, [
    'Candidate Keys: LogID and {RobotID, LogTime}',
    'FDs: LogID → RobotID, LogTime, ErrorCode, UsageStatistics',
    '     {RobotID, LogTime} → LogID, ErrorCode, UsageStatistics',
], font_size=SZ_SMALL, space_after=4, indent=0.5)

add_para(doc, 'Checking partial dependency on {RobotID, LogTime}:', bold=True, font_size=SZ_SMALL, space_after=3)
add_bullet(doc, 'RobotID alone does NOT determine ErrorCode or UsageStatistics (a robot has many logs over time). No partial dependency.')
add_bullet(doc, 'LogTime alone does NOT determine anything (many robots log at the same time). No partial dependency.')

add_para(doc,
    'All non-prime attributes require the full composite key {RobotID, LogTime}. '
    'DeviceLog is in 2NF. ✓',
    font_size=SZ_SMALL)

add_para(doc, 'Conclusion: All 11 relations satisfy 2NF.',
         bold=True, font_size=SZ_SMALL, space_after=6, space_before=8)

doc.add_page_break()

# ── 3.3 3NF ──
heading2(doc, '3.3. Third Normal Form (3NF)')

add_para(doc,
    'A relation is in 3NF if it is in 2NF and no non-prime attribute is transitively dependent '
    'on any candidate key. Formally: for every non-trivial FD X → A, either X is a superkey, '
    'or A is a prime attribute.',
    font_size=SZ_SMALL)

add_para(doc, 'Systematic check for each relation:', bold=True, space_after=6)

# 3NF analysis - each relation as separate structured paragraphs
nf3_data = [
    ('Customer', [
        'FDs: CustomerID → {FullName, PhoneNumber, Email, Address, Password}',
        '     Email → CustomerID;  PhoneNumber → CustomerID',
    ], [
        'CustomerID → FullName: CustomerID is a superkey. ✓',
        'Email → CustomerID: CustomerID is a prime attribute (part of CK). ✓',
        'PhoneNumber → CustomerID: CustomerID is a prime attribute. ✓',
    ], 'No transitive dependencies. In 3NF. ✓'),

    ('Employee', [
        'FDs: EmployeeID → {FullName, Role, PhoneNumber, Email, Password}',
        '     Email → EmployeeID',
    ], [
        'EmployeeID → Role: EmployeeID is a superkey. ✓',
        'Email → EmployeeID: EmployeeID is a prime attribute. ✓',
    ], 'No transitive dependencies. In 3NF. ✓'),

    ('RobotModel', [
        'FDs: ModelID → {Brand, ModelName, Features, Specs, UnitPrice, WarrantyDuration}',
        '     ModelName → ModelID',
    ], [
        'ModelID → Brand: ModelID is a superkey. ✓',
        'ModelName → ModelID: ModelID is a prime attribute. ✓',
    ], 'No transitive dependencies. In 3NF. ✓'),

    ('RobotUnit', [
        'FDs: RobotID → {ModelID, SerialNumber, Status}',
        '     SerialNumber → {RobotID, ModelID, Status}',
    ], [
        'Both RobotID and SerialNumber are superkeys.',
        'All FDs have superkey LHS. ✓',
    ], 'In 3NF. ✓'),

    ('SalesOrder', [
        'FD: OrderID → {CustomerID, EmployeeID, OrderDate, TotalAmount, OrderStatus}',
    ], [
        'OrderID is the only CK and is a superkey.',
        'No transitive dependencies possible with a single FD. ✓',
    ], 'In 3NF. ✓'),

    ('OrderDetail', [
        'FDs: {OrderID, RobotID} → SellingPrice;  RobotID → OrderID',
    ], [
        'RobotID → OrderID: RobotID is a superkey (it is a CK). ✓',
        '{OrderID, RobotID} → SellingPrice: {OrderID, RobotID} is a superkey. ✓',
    ], 'All non-prime attributes directly determined by superkeys. In 3NF. ✓'),

    ('Payment', [
        'FD: PaymentID → {OrderID, ServiceRecordID, Amount, PaymentDate, PaymentMethod}',
    ], [
        'PaymentID is the only CK and superkey.',
        'Single determinant. ✓',
    ], 'In 3NF. ✓'),

    ('WarrantyRegistration', [
        'FDs: WarrantyID → {RobotID, CustomerID, StartDate, EndDate}',
        '     RobotID → WarrantyID',
    ], [
        'WarrantyID → CustomerID: WarrantyID is a superkey. ✓',
        'RobotID → WarrantyID: WarrantyID is a prime attribute. ✓',
    ], 'No transitive dependencies. In 3NF. ✓'),

    ('ServiceRequest', [
        'FD: RequestID → {RobotID, CustomerID, IssueDescription, RequestDate, Status}',
    ], [
        'RequestID is the only CK. Single determinant is a superkey. ✓',
    ], 'In 3NF. ✓'),

    ('MaintenanceRecord', [
        'FDs: RecordID → {RequestID, TechnicianID, ActionsTaken, PartsReplaced, ServiceFee, CompletionDate}',
        '     RequestID → RecordID',
    ], [
        'RecordID → TechnicianID: RecordID is a superkey. ✓',
        'RequestID → RecordID: RecordID is a prime attribute. ✓',
    ], 'No transitive dependencies. In 3NF. ✓'),

    ('DeviceLog', [
        'FDs: LogID → {RobotID, LogTime, ErrorCode, UsageStatistics}',
        '     {RobotID, LogTime} → {LogID, ErrorCode, UsageStatistics}',
    ], [
        'LogID → ErrorCode: LogID is a superkey. ✓',
        '{RobotID, LogTime} → ErrorCode: {RobotID, LogTime} is a superkey. ✓',
    ], 'No transitive dependencies. In 3NF. ✓'),
]

for rel_name, fds, checks, conclusion in nf3_data:
    add_para(doc, f'Relation: {rel_name}', bold=True, font_size=SZ_SMALL, space_after=2)
    add_multiline_para(doc, fds, font_size=10, space_after=3, indent=0.5)
    for check in checks:
        add_bullet(doc, check, font_size=10)
    add_para(doc, conclusion, italic=True, font_size=10, space_after=10)

add_para(doc, 'Conclusion: All 11 relations satisfy 3NF.',
         bold=True, font_size=SZ_SMALL, space_after=6)

doc.add_page_break()

# ── 3.4 BCNF ──
heading2(doc, '3.4. Boyce–Codd Normal Form (BCNF)')

add_para(doc,
    'A relation is in BCNF if for every non-trivial FD X → A, X is a superkey. '
    'BCNF is stricter than 3NF: it does not allow the exception where A is a prime attribute.',
    font_size=SZ_SMALL)

add_para(doc, 'Systematic BCNF check:', bold=True, space_after=6)

bcnf_data = [
    ('Customer', [
        'CKs: CustomerID, Email, PhoneNumber',
    ], [
        'CustomerID → ... : CustomerID is a superkey (CK). ✓',
        'Email → CustomerID: Email⁺ = all attributes → Email is a superkey. ✓',
        'PhoneNumber → CustomerID: PhoneNumber⁺ = all attributes → superkey. ✓',
    ], 'All determinants are superkeys. In BCNF. ✓'),

    ('Employee', [
        'CKs: EmployeeID, Email',
    ], [
        'EmployeeID is a superkey. ✓',
        'Email → EmployeeID: Email⁺ = all attributes → superkey. ✓',
    ], 'In BCNF. ✓'),

    ('RobotModel', [
        'CKs: ModelID, ModelName',
    ], [
        'ModelID is a superkey. ✓',
        'ModelName → ModelID: ModelName⁺ = all attributes → superkey. ✓',
    ], 'In BCNF. ✓'),

    ('RobotUnit', [
        'CKs: RobotID, SerialNumber',
    ], [
        'RobotID is a superkey. ✓',
        'SerialNumber⁺ = all attributes → superkey. ✓',
    ], 'In BCNF. ✓'),

    ('SalesOrder', [
        'CK: OrderID (single CK)',
    ], [
        'OrderID is the only determinant and is a superkey. ✓',
    ], 'In BCNF. ✓'),

    ('OrderDetail', [
        'CKs: RobotID, {OrderID, RobotID}',
    ], [
        '{OrderID, RobotID} → SellingPrice: superkey (superset of CK). ✓',
        'RobotID → OrderID: RobotID⁺ = {RobotID, OrderID, SellingPrice} = all → superkey. ✓',
    ], 'In BCNF. ✓'),

    ('Payment', [
        'CK: PaymentID',
    ], [
        'Single determinant is a superkey. ✓',
    ], 'In BCNF. ✓'),

    ('WarrantyRegistration', [
        'CKs: WarrantyID, RobotID',
    ], [
        'WarrantyID is a superkey. ✓',
        'RobotID → WarrantyID: RobotID⁺ = all attributes → superkey. ✓',
    ], 'In BCNF. ✓'),

    ('ServiceRequest', [
        'CK: RequestID',
    ], [
        'Single determinant is a superkey. ✓',
    ], 'In BCNF. ✓'),

    ('MaintenanceRecord', [
        'CKs: RecordID, RequestID',
    ], [
        'RecordID is a superkey. ✓',
        'RequestID → RecordID: RequestID⁺ = all attributes → superkey. ✓',
    ], 'In BCNF. ✓'),

    ('DeviceLog', [
        'CKs: LogID, {RobotID, LogTime}',
    ], [
        'LogID is a superkey. ✓',
        '{RobotID, LogTime}⁺ = all attributes → superkey. ✓',
    ], 'In BCNF. ✓'),
]

for rel_name, cks, checks, conclusion in bcnf_data:
    add_para(doc, f'Relation: {rel_name}', bold=True, font_size=SZ_SMALL, space_after=2)
    add_multiline_para(doc, cks, font_size=10, space_after=3, indent=0.5)
    for check in checks:
        add_bullet(doc, check, font_size=10)
    add_para(doc, conclusion, italic=True, font_size=10, space_after=10)

# NF Summary table
add_para(doc, 'Summary: Normal Form Classification of All Relations',
         bold=True, space_after=6, space_before=8)

create_table(doc,
    ['Relation', '1NF', '2NF', '3NF', 'BCNF'],
    [
        ['Customer', '✓', '✓', '✓', '✓'],
        ['Employee', '✓', '✓', '✓', '✓'],
        ['RobotModel', '✓*', '✓', '✓', '✓'],
        ['RobotUnit', '✓', '✓', '✓', '✓'],
        ['SalesOrder', '✓', '✓', '✓', '✓'],
        ['OrderDetail', '✓', '✓', '✓', '✓'],
        ['Payment', '✓', '✓', '✓', '✓'],
        ['WarrantyReg.', '✓', '✓', '✓', '✓'],
        ['ServiceRequest', '✓', '✓', '✓', '✓'],
        ['Maintenance Rec.', '✓*', '✓', '✓', '✓'],
        ['DeviceLog', '✓*', '✓', '✓', '✓'],
    ])
doc.add_paragraph()

add_para(doc,
    '* ✓* indicates 1NF depends on treating multi-valued fields (Features, PartsReplaced, '
    'UsageStatistics) as single atomic text values.',
    font_size=10, italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 4. STEP-BY-STEP DECOMPOSITION
# ══════════════════════════════════════════════════════════════════════
heading1(doc, '4. Step-by-Step Decomposition')

add_para(doc,
    'Based on the analysis in Section 3, all 11 relations already satisfy BCNF with respect to '
    'the functional dependencies identified in Lab 2. Therefore, no decomposition is required '
    'to achieve higher normal forms from a strict FD-based normalization perspective.')

add_para(doc,
    'However, there are practical design refinements that address the anomalies identified in '
    'Section 2 and the potential 1NF concerns. These refinements follow normalization principles '
    'and improve data integrity.')

# ── Step 1 ──
heading2(doc, 'Step 1: Decompose Payment to Eliminate Structural Redundancy')

add_para(doc, 'Problem:', bold=True, font_size=SZ_SMALL, space_after=3)
add_para(doc,
    'The Payment relation uses nullable foreign keys (OrderID, ServiceRecordID) to serve '
    'dual purposes — tracking sales payments and service payments. This causes NULL-value '
    'redundancy and insertion/update anomalies.',
    font_size=SZ_SMALL)

add_para(doc, 'Original Relation:', bold=True, font_size=SZ_SMALL, space_after=3, space_before=6)
add_multiline_para(doc, [
    'Payment(PaymentID, OrderID, ServiceRecordID, Amount, PaymentDate, PaymentMethod)',
    'FD: PaymentID → OrderID, ServiceRecordID, Amount, PaymentDate, PaymentMethod',
], font_size=SZ_SMALL, space_after=6, indent=0.5)

add_para(doc, 'Decomposed Relations:', bold=True, font_size=SZ_SMALL, space_after=3, space_before=6)

add_para(doc, 'R1: Payment (core)', bold=True, font_size=SZ_SMALL, space_after=2, space_before=4)
add_multiline_para(doc, [
    'Payment(PaymentID, Amount, PaymentDate, PaymentMethod)',
    'FD: PaymentID → Amount, PaymentDate, PaymentMethod',
    'PK: PaymentID',
], font_size=SZ_SMALL, space_after=4, indent=0.5)

add_para(doc, 'R2: OrderPayment (links payment to sales order)', bold=True, font_size=SZ_SMALL, space_after=2, space_before=4)
add_multiline_para(doc, [
    'OrderPayment(PaymentID, OrderID)',
    'FD: PaymentID → OrderID',
    'PK: PaymentID',
    'FK: PaymentID → Payment, OrderID → SalesOrder',
], font_size=SZ_SMALL, space_after=4, indent=0.5)

add_para(doc, 'R3: ServicePayment (links payment to service record)', bold=True, font_size=SZ_SMALL, space_after=2, space_before=4)
add_multiline_para(doc, [
    'ServicePayment(PaymentID, ServiceRecordID)',
    'FD: PaymentID → ServiceRecordID',
    'PK: PaymentID',
    'FK: PaymentID → Payment, ServiceRecordID → MaintenanceRecord',
], font_size=SZ_SMALL, space_after=6, indent=0.5)

add_para(doc, 'Verification:', bold=True, font_size=SZ_SMALL, space_after=3)
add_bullet(doc, 'Lossless Join: Payment = R1 ⋈ (R2 ∪ R3) on PaymentID. Since PaymentID is the key in all decomposed relations, the join is lossless.')
add_bullet(doc, 'Dependency Preservation: All original FDs are preserved across the decomposed relations.')
add_bullet(doc, 'Anomaly Resolution: NULLs are eliminated — a sales payment only exists in OrderPayment, a service payment only in ServicePayment.')

# ── Step 2 ──
heading2(doc, 'Step 2: Simplify OrderDetail Primary Key')

add_para(doc, 'Problem:', bold=True, font_size=SZ_SMALL, space_after=3)
add_para(doc,
    'OrderDetail has composite PK {OrderID, RobotID}, but RobotID → OrderID means RobotID '
    'alone is a candidate key.',
    font_size=SZ_SMALL)

add_para(doc, 'Original Relation:', bold=True, font_size=SZ_SMALL, space_after=3, space_before=6)
add_multiline_para(doc, [
    'OrderDetail(OrderID, RobotID, SellingPrice)',
    'FDs: {OrderID, RobotID} → SellingPrice;  RobotID → OrderID',
    'PK: {OrderID, RobotID}',
], font_size=SZ_SMALL, space_after=6, indent=0.5)

add_para(doc, 'Refined Relation:', bold=True, font_size=SZ_SMALL, space_after=3, space_before=6)
add_multiline_para(doc, [
    'OrderDetail(RobotID, OrderID, SellingPrice)',
    'FD: RobotID → OrderID, SellingPrice',
    'PK: RobotID',
    'FK: RobotID → RobotUnit, OrderID → SalesOrder',
], font_size=SZ_SMALL, space_after=6, indent=0.5)

add_para(doc, 'Verification:', bold=True, font_size=SZ_SMALL, space_after=3)
add_bullet(doc, 'This is a PK refinement, not a decomposition. The schema remains the same; only the PK changes to the minimal candidate key.')
add_bullet(doc, 'All FDs are directly determined by the single PK (RobotID).')
add_bullet(doc, 'Anomaly Resolution: Insertion anomaly mitigated — selling price can be associated with a RobotID independently.')

# ── Step 3 ──
heading2(doc, 'Step 3: Decompose Multi-Valued Attributes (1NF Enforcement)')

add_para(doc, 'Problem:', bold=True, font_size=SZ_SMALL, space_after=3)
add_para(doc,
    'Three relations contain attributes that may store multiple values: RobotModel (Features, '
    'Specifications), MaintenanceRecord (PartsReplaced), and DeviceLog (UsageStatistics).',
    font_size=SZ_SMALL)

# 3a
add_para(doc, 'Decomposition 3a: RobotModel — Extract Features',
         bold=True, font_size=SZ_SMALL, space_after=3, space_before=8)

add_para(doc, 'If Features stores multiple values (e.g., "WiFi, Auto-charging, Mopping"), decompose:',
         font_size=SZ_SMALL, space_after=4)

add_multiline_para(doc, [
    'RobotModel(ModelID, Brand, ModelName, Specifications, UnitPrice, WarrantyDuration)',
    '    PK: ModelID  |  Unique: ModelName',
    '',
    'ModelFeature(ModelID, Feature)   ← NEW',
    '    PK: {ModelID, Feature}',
    '    FK: ModelID → RobotModel',
], font_size=SZ_SMALL, space_after=6, indent=0.5)

# 3b
add_para(doc, 'Decomposition 3b: MaintenanceRecord — Extract PartsReplaced',
         bold=True, font_size=SZ_SMALL, space_after=3, space_before=8)

add_para(doc, 'If PartsReplaced stores multiple parts (e.g., "brush roller, battery, filter"):',
         font_size=SZ_SMALL, space_after=4)

add_multiline_para(doc, [
    'MaintenanceRecord(RecordID, RequestID, TechnicianID, ActionsTaken, ServiceFee, CompletionDate)',
    '    PK: RecordID  |  Unique: RequestID',
    '',
    'ReplacedPart(RecordID, PartName)   ← NEW',
    '    PK: {RecordID, PartName}',
    '    FK: RecordID → MaintenanceRecord',
], font_size=SZ_SMALL, space_after=6, indent=0.5)

# 3c
add_para(doc, 'Decomposition 3c: DeviceLog — Extract UsageStatistics',
         bold=True, font_size=SZ_SMALL, space_after=3, space_before=8)

add_para(doc, 'If UsageStatistics contains multiple metrics:',
         font_size=SZ_SMALL, space_after=4)

add_multiline_para(doc, [
    'DeviceLog(LogID, RobotID, LogTime, ErrorCode)',
    '    PK: LogID  |  Unique: {RobotID, LogTime}',
    '',
    'LogStatistic(LogID, MetricName, MetricValue)   ← NEW',
    '    PK: {LogID, MetricName}',
    '    FK: LogID → DeviceLog',
], font_size=SZ_SMALL, space_after=6, indent=0.5)

add_para(doc, 'Verification for all Step 3 decompositions:', bold=True, font_size=SZ_SMALL, space_after=3)
add_bullet(doc, 'Lossless Join: Each decomposition splits a multi-valued attribute into a separate relation linked by the original PK. The natural join reconstructs the data.')
add_bullet(doc, 'Dependency Preservation: All original FDs involving non-decomposed attributes remain in the parent relation.')
add_bullet(doc, 'Normal Form: All new relations are in BCNF — each has a composite PK that determines all attributes.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 5. FINAL NORMALIZED SCHEMA
# ══════════════════════════════════════════════════════════════════════
heading1(doc, '5. Final Normalized Schema')

add_para(doc,
    'The following presents the complete, normalized database schema after all decomposition steps. '
    'All relations are in BCNF. Anomalies identified in Section 2 have been addressed.')

# ── 5.1 User Management ──
heading2(doc, '5.1. User Management')

add_multiline_para(doc, [
    'Customer(CustomerID, FullName, PhoneNumber, Email, Address, Password)',
    '    PK: CustomerID  |  Unique: Email, PhoneNumber',
    '    Normal Form: BCNF ✓ — No changes required.',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

add_multiline_para(doc, [
    'Employee(EmployeeID, FullName, Role, PhoneNumber, Email, Password)',
    '    PK: EmployeeID  |  Unique: Email',
    '    Normal Form: BCNF ✓ — No changes required.',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

# ── 5.2 Product & Inventory ──
heading2(doc, '5.2. Product & Inventory Management')

add_multiline_para(doc, [
    'RobotModel(ModelID, Brand, ModelName, Specifications, UnitPrice, WarrantyDuration)',
    '    PK: ModelID  |  Unique: ModelName',
    '    Normal Form: BCNF ✓ — Features attribute extracted.',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

add_multiline_para(doc, [
    'ModelFeature(ModelID, Feature)   [NEW]',
    '    PK: {ModelID, Feature}  |  FK: ModelID → RobotModel',
    '    Normal Form: BCNF ✓',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

add_multiline_para(doc, [
    'RobotUnit(RobotID, ModelID, SerialNumber, Status)',
    '    PK: RobotID  |  Unique: SerialNumber  |  FK: ModelID → RobotModel',
    '    Normal Form: BCNF ✓ — No changes required.',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

# ── 5.3 Sales & Transactions ──
heading2(doc, '5.3. Sales & Transactions Management')

add_multiline_para(doc, [
    'SalesOrder(OrderID, CustomerID, EmployeeID, OrderDate, TotalAmount, OrderStatus)',
    '    PK: OrderID  |  FK: CustomerID → Customer, EmployeeID → Employee',
    '    Normal Form: BCNF ✓ — No changes required.',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

add_multiline_para(doc, [
    'OrderDetail(RobotID, OrderID, SellingPrice)   [MODIFIED — PK changed]',
    '    PK: RobotID (changed from composite {OrderID, RobotID})',
    '    FK: RobotID → RobotUnit, OrderID → SalesOrder',
    '    Normal Form: BCNF ✓',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

add_multiline_para(doc, [
    'Payment(PaymentID, Amount, PaymentDate, PaymentMethod)   [MODIFIED — decomposed]',
    '    PK: PaymentID',
    '    Normal Form: BCNF ✓',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

add_multiline_para(doc, [
    'OrderPayment(PaymentID, OrderID)   [NEW]',
    '    PK: PaymentID  |  FK: PaymentID → Payment, OrderID → SalesOrder',
    '    Normal Form: BCNF ✓',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

add_multiline_para(doc, [
    'ServicePayment(PaymentID, ServiceRecordID)   [NEW]',
    '    PK: PaymentID  |  FK: PaymentID → Payment, ServiceRecordID → MaintenanceRecord',
    '    Normal Form: BCNF ✓',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

add_multiline_para(doc, [
    'WarrantyRegistration(WarrantyID, RobotID, CustomerID, StartDate, EndDate)',
    '    PK: WarrantyID  |  Unique: RobotID',
    '    FK: RobotID → RobotUnit, CustomerID → Customer',
    '    Normal Form: BCNF ✓ — No changes required.',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

# ── 5.4 Maintenance & IoT ──
heading2(doc, '5.4. Maintenance, Repair & IoT Extension')

add_multiline_para(doc, [
    'ServiceRequest(RequestID, RobotID, CustomerID, IssueDescription, RequestDate, Status)',
    '    PK: RequestID  |  FK: RobotID → RobotUnit, CustomerID → Customer',
    '    Normal Form: BCNF ✓ — No changes required.',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

add_multiline_para(doc, [
    'MaintenanceRecord(RecordID, RequestID, TechnicianID, ActionsTaken, ServiceFee, CompletionDate)   [MODIFIED]',
    '    PK: RecordID  |  Unique: RequestID',
    '    FK: RequestID → ServiceRequest, TechnicianID → Employee',
    '    Normal Form: BCNF ✓ — PartsReplaced extracted.',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

add_multiline_para(doc, [
    'ReplacedPart(RecordID, PartName)   [NEW]',
    '    PK: {RecordID, PartName}  |  FK: RecordID → MaintenanceRecord',
    '    Normal Form: BCNF ✓',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

add_multiline_para(doc, [
    'DeviceLog(LogID, RobotID, LogTime, ErrorCode)   [MODIFIED]',
    '    PK: LogID  |  Unique: {RobotID, LogTime}  |  FK: RobotID → RobotUnit',
    '    Normal Form: BCNF ✓ — UsageStatistics extracted.',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

add_multiline_para(doc, [
    'LogStatistic(LogID, MetricName, MetricValue)   [NEW]',
    '    PK: {LogID, MetricName}  |  FK: LogID → DeviceLog',
    '    Normal Form: BCNF ✓',
], font_size=SZ_SMALL, space_after=8, indent=0.5)

# ── 5.5 Anomaly Resolution ──
heading2(doc, '5.5. How Anomalies Are Eliminated')

create_table(doc,
    ['Original Anomaly', 'Resolution in New Schema'],
    [
        ['OrderDetail: Redundant composite PK',
         'PK changed to RobotID (single CK). OrderID remains as FK.'],
        ['Payment: NULL-value redundancy',
         'Decomposed into Payment + OrderPayment + ServicePayment. No NULLs.'],
        ['Payment: Insertion anomaly',
         'Payments can exist independently without requiring an order/service link.'],
        ['Payment: Update anomaly',
         'Reclassification = delete from OrderPayment, insert into ServicePayment.'],
        ['OrderDetail: Insertion anomaly',
         'Selling price keyed by RobotID alone, allowing pre-assignment.'],
        ['MaintenanceRecord: 1NF (PartsReplaced)',
         'Decomposed into ReplacedPart relation with atomic values.'],
        ['DeviceLog: 1NF (UsageStatistics)',
         'Decomposed into LogStatistic relation with atomic metrics.'],
        ['RobotModel: 1NF (Features)',
         'Decomposed into ModelFeature relation with atomic values.'],
    ])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 6. CONCLUSION
# ══════════════════════════════════════════════════════════════════════
heading1(doc, '6. Conclusion')

add_para(doc,
    'Through this lab, the group has performed a comprehensive anomaly detection and normalization '
    'analysis on the relational database schema designed in Lab 2 for the "Household Cleaning Robot '
    'Sales & Maintenance Management System."')

add_para(doc, 'Key Findings:', bold=True, space_after=3, space_before=6)

add_bullet(doc,
    'The Lab 2 schema was already well-designed — all 11 original relations satisfy BCNF '
    'with respect to the identified functional dependencies. This validates the careful entity '
    'analysis and FD identification performed in Lab 2.')

add_bullet(doc,
    'Anomalies were identified primarily in three areas: (1) the OrderDetail relation\'s '
    'redundant composite primary key, (2) the Payment relation\'s dual-purpose nullable foreign '
    'keys, and (3) potential 1NF violations in multi-valued text fields.')

add_bullet(doc,
    'Three decomposition steps were performed: Payment was split into three relations to '
    'eliminate NULL redundancy; OrderDetail\'s primary key was refined to use the minimal '
    'candidate key; and multi-valued attributes (Features, PartsReplaced, UsageStatistics) were '
    'extracted into separate relations to enforce strict 1NF compliance.')

add_bullet(doc,
    'All decompositions maintain lossless join and dependency preservation properties, ensuring '
    'no information is lost and all original functional dependencies can be verified.')

add_bullet(doc,
    'The final normalized schema consists of 16 relations (up from 11), all in BCNF, with '
    'anomalies eliminated or mitigated. The schema is now more robust, maintainable, and '
    'resistant to data inconsistencies.')

doc.add_paragraph()

add_para(doc,
    'This normalization exercise reinforced the group\'s understanding that good initial design '
    '(as achieved in Lab 2) significantly reduces the need for extensive decomposition. The primary '
    'refinements focused on eliminating structural redundancy and enforcing strict atomicity — '
    'practical improvements that enhance the long-term integrity and flexibility of the database system.')

# ── Save ──
output_path = r'c:\Users\Khuong\Desktop\FPTU\DBI202\LAB3\Lab3_report.docx'
doc.save(output_path)
print(f'Lab 3 report saved to: {output_path}')
print('Done!')
