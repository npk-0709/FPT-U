"""
Lab 4 Report Generator
Generates a complete Lab 4: Relational Database Design Process report
for the Household Cleaning Robot Sales & Maintenance Management System.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ==============================================================================
# HELPER FUNCTIONS (same as Lab 3)
# ==============================================================================

def set_cell_shading(cell, color):
    """Set cell background shading color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_para(doc, text, bold=False, italic=False, font_size=12,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
             space_before=0, first_line_indent=None):
    """Add a formatted paragraph to the document."""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = Pt(12 * (font_size / 12))
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return para


def add_multiline_para(doc, lines, bold_first=False, font_size=11,
                       space_after=6, indent=None):
    """Add a multi-line paragraph (lines joined by newline)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(space_after)
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    for i, line in enumerate(lines):
        run = para.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        if bold_first and i == 0:
            run.bold = True
        if i < len(lines) - 1:
            run.add_break()
    return para


def add_bullet(doc, text, level=0, bold=False, font_size=12):
    """Add a bullet point."""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return para


def add_bullet_multiline(doc, lines, level=0, font_size=12):
    """Add a bullet with multiple lines."""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    para.paragraph_format.space_after = Pt(2)
    for i, line in enumerate(lines):
        run = para.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        if i < len(lines) - 1:
            run.add_break()
    return para


def heading1(doc, text):
    """Add Heading 1."""
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return h


def heading2(doc, text):
    """Add Heading 2."""
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return h


def create_table(doc, headers, rows):
    """Create a formatted table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    return table


# ==============================================================================
# DOCUMENT SETUP
# ==============================================================================

def setup_document():
    """Create and configure the document with proper margins and styles."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = Pt(12)

    return doc


# ==============================================================================
# COVER PAGE
# ==============================================================================

def add_cover_page(doc):
    """Add the cover page."""
    # 4 blank paragraphs
    for _ in range(4):
        doc.add_paragraph()

    add_para(doc, 'FPT UNIVERSITY', bold=True, font_size=18,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, 'DEPARTMENT OF INFORMATION TECHNOLOGY', bold=True,
             font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=24)

    # Blank paragraph
    doc.add_paragraph()

    add_para(doc, 'DBI202 – Database Systems', bold=True, font_size=16,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, 'Lab 4: Relational Database Design Process', bold=True,
             font_size=20, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, 'Household Cleaning Robot Sales & Maintenance Management System',
             italic=True, font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=36)

    # 2 blank paragraphs
    for _ in range(2):
        doc.add_paragraph()

    add_para(doc, 'Group Members:', bold=True, font_size=13,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    for i in range(1, 6):
        add_para(doc, f'[Member {i} – Student ID]', font_size=12,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    add_para(doc, 'Date: June 2026', bold=True, font_size=12,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    doc.add_page_break()


# ==============================================================================
# TABLE OF CONTENTS
# ==============================================================================

def add_table_of_contents(doc):
    """Add the table of contents."""
    heading1(doc, 'Table of Contents')

    toc_items = [
        ('1.', 'Objective', 0),
        ('2.', 'Conceptual Model (ERD)', 0),
        ('', '2.1. Entities and Attributes', 1),
        ('', '2.2. Relationships and Cardinalities', 1),
        ('3.', 'Logical Model', 0),
        ('', '3.1. Relational Schema', 1),
        ('', '3.2. Primary Keys and Foreign Keys', 1),
        ('', '3.3. M:N Relationship Resolution', 1),
        ('4.', 'Physical Model', 0),
        ('', '4.1. Table Definitions with Data Types', 1),
        ('5.', 'Constraints', 0),
        ('', '5.1. Column-Level Constraints', 1),
        ('', '5.2. Table-Level Constraints', 1),
        ('6.', 'Conclusion and Reflection', 0),
    ]

    for num, title, level in toc_items:
        text = f'{num} {title}'.strip() if num else title
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        if level == 1:
            para.paragraph_format.left_indent = Cm(1.27)
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        if level == 0:
            run.bold = True

    doc.add_page_break()


# ==============================================================================
# SECTION 1: OBJECTIVE
# ==============================================================================

def add_section_1(doc):
    """Add Section 1: Objective."""
    heading1(doc, '1. Objective')

    add_para(doc, (
        'The purpose of Lab 4 is to guide students through the structured process of '
        'relational database design — translating business requirements into a conceptual model '
        '(Entity-Relationship Diagram), refining it into a logical relational model, implementing '
        'it as a physical database schema for Microsoft SQL Server, and specifying comprehensive '
        'constraints to ensure data integrity. This lab integrates the work from previous labs '
        '(requirements analysis, ERD design, and normalization) into a complete, deployable '
        'database design for the Household Cleaning Robot Sales & Maintenance Management System.'
    ))

    add_bullet(doc, 'Develop an Entity-Relationship Diagram (ERD) for the Household Cleaning '
               'Robot Sales & Maintenance Management System')
    add_bullet(doc, 'Convert the ERD into a logical relational schema with primary keys (PK) '
               'and foreign keys (FK)')
    add_bullet(doc, 'Map the logical schema to a physical model with SQL Server data types')
    add_bullet(doc, 'Specify comprehensive constraints (NOT NULL, UNIQUE, PK, FK, CHECK, DEFAULT) '
               'to enforce data integrity')
    add_bullet(doc, 'Document the complete design process and reflect on data integrity considerations')

    doc.add_page_break()


# ==============================================================================
# SECTION 2: CONCEPTUAL MODEL (ERD)
# ==============================================================================

def add_section_2(doc):
    """Add Section 2: Conceptual Model (ERD)."""
    heading1(doc, '2. Conceptual Model (ERD)')

    add_para(doc, (
        'This section presents the Entity-Relationship Diagram for the Household Cleaning Robot '
        'Sales & Maintenance Management System. The ERD was developed based on the project '
        'requirements and business rules identified in previous labs, identifying all entities, '
        'their attributes, and the relationships between them with appropriate cardinalities.'
    ))

    add_para(doc, '[Insert Conceptual ERD (Chen notation) here — file: Lab4_Conceptual_ERD.drawio]',
             bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12, space_before=12)

    # 2.1. Entities and Attributes
    heading2(doc, '2.1. Entities and Attributes')

    add_para(doc, (
        'The following table summarizes all 16 entities identified for the system, organized '
        'by functional domain. Each entity is listed with its attributes and a brief description '
        'of its purpose within the system.'
    ))

    entities_rows = [
        ['Customer', 'CustomerID (PK), FullName, PhoneNumber, Email, Address, Password',
         'Stores customer personal information and account credentials'],
        ['Employee', 'EmployeeID (PK), FullName, Role, PhoneNumber, Email, Password',
         'Manages staff information including sales staff, technicians, and administrators'],
        ['RobotModel', 'ModelID (PK), Brand, ModelName, Specifications, UnitPrice, WarrantyDuration',
         'Defines robot product models with pricing and warranty info'],
        ['ModelFeature', 'ModelID (PK, FK), Feature (PK)',
         'Stores individual features for each robot model (multi-valued attribute decomposition)'],
        ['RobotUnit', 'RobotID (PK), ModelID (FK), SerialNumber, Status',
         'Tracks individual robot units in inventory'],
        ['SalesOrder', 'OrderID (PK), CustomerID (FK), EmployeeID (FK), OrderDate, TotalAmount, OrderStatus',
         'Records sales transactions'],
        ['OrderDetail', 'RobotID (PK, FK), OrderID (FK), SellingPrice',
         'Links specific robot units to sales orders'],
        ['Payment', 'PaymentID (PK), Amount, PaymentDate, PaymentMethod',
         'Records payment transactions (base entity)'],
        ['OrderPayment', 'PaymentID (PK, FK), OrderID (FK)',
         'Links payments to sales orders'],
        ['ServicePayment', 'PaymentID (PK, FK), ServiceRecordID (FK)',
         'Links payments to maintenance services'],
        ['WarrantyRegistration', 'WarrantyID (PK), RobotID (FK), CustomerID (FK), StartDate, EndDate',
         'Manages warranty coverage for purchased robots'],
        ['ServiceRequest', 'RequestID (PK), RobotID (FK), CustomerID (FK), IssueDescription, RequestDate, Status',
         'Tracks customer service/repair requests'],
        ['MaintenanceRecord', 'RecordID (PK), RequestID (FK), TechnicianID (FK), ActionsTaken, ServiceFee, CompletionDate',
         'Documents maintenance/repair work performed'],
        ['ReplacedPart', 'RecordID (PK, FK), PartName (PK)',
         'Tracks individual parts replaced during maintenance'],
        ['DeviceLog', 'LogID (PK), RobotID (FK), LogTime, ErrorCode',
         'Stores IoT data from connected robots'],
        ['LogStatistic', 'LogID (PK, FK), MetricName (PK), MetricValue',
         'Stores individual usage metrics from device logs'],
    ]

    create_table(doc, ['Entity', 'Attributes', 'Description'], entities_rows)

    # 2.2. Relationships and Cardinalities
    heading2(doc, '2.2. Relationships and Cardinalities')

    add_para(doc, (
        'The following table documents all relationships between entities in the system, '
        'including the cardinality type and a description of each relationship.'
    ))

    rel_rows = [
        ['Customer → SalesOrder', '1:N', 'A customer can place multiple sales orders'],
        ['Employee → SalesOrder', '1:N', 'An employee can process multiple sales orders'],
        ['SalesOrder → OrderDetail', '1:N', 'A sales order contains multiple order detail lines'],
        ['RobotUnit → OrderDetail', '1:1', 'Each robot unit can only be sold once'],
        ['RobotModel → RobotUnit', '1:N', 'A robot model can have multiple individual units'],
        ['RobotModel → ModelFeature', '1:N', 'A robot model can have multiple features'],
        ['Payment → OrderPayment', '1:1', 'Each order payment links to exactly one payment'],
        ['Payment → ServicePayment', '1:1', 'Each service payment links to exactly one payment'],
        ['SalesOrder → OrderPayment', '1:N', 'A sales order can have multiple payments'],
        ['MaintenanceRecord → ServicePayment', '1:N', 'A maintenance record can have multiple payments'],
        ['Customer → WarrantyRegistration', '1:N', 'A customer can have multiple warranty registrations'],
        ['RobotUnit → WarrantyRegistration', '1:1', 'Each robot unit has exactly one warranty'],
        ['Customer → ServiceRequest', '1:N', 'A customer can submit multiple service requests'],
        ['RobotUnit → ServiceRequest', '1:N', 'A robot unit can have multiple service requests'],
        ['ServiceRequest → MaintenanceRecord', '1:1', 'Each service request has one maintenance record'],
        ['Employee → MaintenanceRecord', '1:N', 'A technician can handle multiple maintenance records'],
        ['RobotUnit → DeviceLog', '1:N', 'A robot unit generates multiple device logs'],
        ['DeviceLog → LogStatistic', '1:N', 'A device log contains multiple usage statistics'],
        ['MaintenanceRecord → ReplacedPart', '1:N', 'A maintenance record can involve multiple replaced parts'],
    ]

    create_table(doc, ['Relationship', 'Type', 'Description'], rel_rows)

    doc.add_page_break()


# ==============================================================================
# SECTION 3: LOGICAL MODEL
# ==============================================================================

def add_section_3(doc):
    """Add Section 3: Logical Model."""
    heading1(doc, '3. Logical Model')

    add_para(doc, (
        'This section converts the conceptual ERD into a relational schema. The logical model '
        'defines the structure of each relation, identifies primary keys and foreign keys, and '
        'resolves any M:N relationships. All relations are normalized to BCNF as established '
        'in Lab 3.'
    ))

    add_para(doc, '[Insert Logical/Relational Model Diagram here — file: Lab4_Relational_Model.drawio]',
             bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12, space_before=12)

    # 3.1. Relational Schema
    heading2(doc, '3.1. Relational Schema')

    add_para(doc, (
        'The following presents the complete relational schema for the system, organized by '
        'functional domain. The schema consists of 16 relations, all normalized to BCNF '
        '(as established in Lab 3).'
    ))

    # User Management
    add_para(doc, 'User Management', bold=True, font_size=12, space_before=6)
    add_multiline_para(doc, [
        'Customer(CustomerID, FullName, PhoneNumber, Email, Address, Password)',
        '    PK: CustomerID  |  Unique: Email, PhoneNumber',
    ], bold_first=True, font_size=11, space_after=4)

    add_multiline_para(doc, [
        'Employee(EmployeeID, FullName, Role, PhoneNumber, Email, Password)',
        '    PK: EmployeeID  |  Unique: Email',
    ], bold_first=True, font_size=11, space_after=8)

    # Product & Inventory
    add_para(doc, 'Product & Inventory', bold=True, font_size=12, space_before=6)
    add_multiline_para(doc, [
        'RobotModel(ModelID, Brand, ModelName, Specifications, UnitPrice, WarrantyDuration)',
        '    PK: ModelID  |  Unique: ModelName',
    ], bold_first=True, font_size=11, space_after=4)

    add_multiline_para(doc, [
        'ModelFeature(ModelID, Feature)',
        '    PK: {ModelID, Feature}  |  FK: ModelID → RobotModel',
    ], bold_first=True, font_size=11, space_after=4)

    add_multiline_para(doc, [
        'RobotUnit(RobotID, ModelID, SerialNumber, Status)',
        '    PK: RobotID  |  Unique: SerialNumber  |  FK: ModelID → RobotModel',
    ], bold_first=True, font_size=11, space_after=8)

    # Sales & Transactions
    add_para(doc, 'Sales & Transactions', bold=True, font_size=12, space_before=6)
    add_multiline_para(doc, [
        'SalesOrder(OrderID, CustomerID, EmployeeID, OrderDate, TotalAmount, OrderStatus)',
        '    PK: OrderID  |  FK: CustomerID → Customer, EmployeeID → Employee',
    ], bold_first=True, font_size=11, space_after=4)

    add_multiline_para(doc, [
        'OrderDetail(RobotID, OrderID, SellingPrice)',
        '    PK: RobotID  |  FK: RobotID → RobotUnit, OrderID → SalesOrder',
    ], bold_first=True, font_size=11, space_after=4)

    add_multiline_para(doc, [
        'Payment(PaymentID, Amount, PaymentDate, PaymentMethod)',
        '    PK: PaymentID',
    ], bold_first=True, font_size=11, space_after=4)

    add_multiline_para(doc, [
        'OrderPayment(PaymentID, OrderID)',
        '    PK: PaymentID  |  FK: PaymentID → Payment, OrderID → SalesOrder',
    ], bold_first=True, font_size=11, space_after=4)

    add_multiline_para(doc, [
        'ServicePayment(PaymentID, ServiceRecordID)',
        '    PK: PaymentID  |  FK: PaymentID → Payment, ServiceRecordID → MaintenanceRecord',
    ], bold_first=True, font_size=11, space_after=4)

    add_multiline_para(doc, [
        'WarrantyRegistration(WarrantyID, RobotID, CustomerID, StartDate, EndDate)',
        '    PK: WarrantyID  |  Unique: RobotID  |  FK: RobotID → RobotUnit, CustomerID → Customer',
    ], bold_first=True, font_size=11, space_after=8)

    # Maintenance & IoT
    add_para(doc, 'Maintenance & IoT', bold=True, font_size=12, space_before=6)
    add_multiline_para(doc, [
        'ServiceRequest(RequestID, RobotID, CustomerID, IssueDescription, RequestDate, Status)',
        '    PK: RequestID  |  FK: RobotID → RobotUnit, CustomerID → Customer',
    ], bold_first=True, font_size=11, space_after=4)

    add_multiline_para(doc, [
        'MaintenanceRecord(RecordID, RequestID, TechnicianID, ActionsTaken, ServiceFee, CompletionDate)',
        '    PK: RecordID  |  Unique: RequestID  |  FK: RequestID → ServiceRequest, TechnicianID → Employee',
    ], bold_first=True, font_size=11, space_after=4)

    add_multiline_para(doc, [
        'ReplacedPart(RecordID, PartName)',
        '    PK: {RecordID, PartName}  |  FK: RecordID → MaintenanceRecord',
    ], bold_first=True, font_size=11, space_after=4)

    add_multiline_para(doc, [
        'DeviceLog(LogID, RobotID, LogTime, ErrorCode)',
        '    PK: LogID  |  Unique: {RobotID, LogTime}  |  FK: RobotID → RobotUnit',
    ], bold_first=True, font_size=11, space_after=4)

    add_multiline_para(doc, [
        'LogStatistic(LogID, MetricName, MetricValue)',
        '    PK: {LogID, MetricName}  |  FK: LogID → DeviceLog',
    ], bold_first=True, font_size=11, space_after=8)

    # 3.2. Primary Keys and Foreign Keys
    heading2(doc, '3.2. Primary Keys and Foreign Keys')

    add_para(doc, (
        'The following table provides a comprehensive summary of all primary keys, foreign keys, '
        'and unique keys across all 16 relations in the system.'
    ))

    pk_fk_rows = [
        ['Customer', 'CustomerID', '—', 'Email, PhoneNumber'],
        ['Employee', 'EmployeeID', '—', 'Email'],
        ['RobotModel', 'ModelID', '—', 'ModelName'],
        ['ModelFeature', '{ModelID, Feature}', 'ModelID → RobotModel', '—'],
        ['RobotUnit', 'RobotID', 'ModelID → RobotModel', 'SerialNumber'],
        ['SalesOrder', 'OrderID', 'CustomerID → Customer, EmployeeID → Employee', '—'],
        ['OrderDetail', 'RobotID', 'RobotID → RobotUnit, OrderID → SalesOrder', '—'],
        ['Payment', 'PaymentID', '—', '—'],
        ['OrderPayment', 'PaymentID', 'PaymentID → Payment, OrderID → SalesOrder', '—'],
        ['ServicePayment', 'PaymentID', 'PaymentID → Payment, ServiceRecordID → MaintenanceRecord', '—'],
        ['WarrantyRegistration', 'WarrantyID', 'RobotID → RobotUnit, CustomerID → Customer', 'RobotID'],
        ['ServiceRequest', 'RequestID', 'RobotID → RobotUnit, CustomerID → Customer', '—'],
        ['MaintenanceRecord', 'RecordID', 'RequestID → ServiceRequest, TechnicianID → Employee', 'RequestID'],
        ['ReplacedPart', '{RecordID, PartName}', 'RecordID → MaintenanceRecord', '—'],
        ['DeviceLog', 'LogID', 'RobotID → RobotUnit', '{RobotID, LogTime}'],
        ['LogStatistic', '{LogID, MetricName}', 'LogID → DeviceLog', '—'],
    ]

    create_table(doc, ['Relation', 'Primary Key', 'Foreign Key(s)', 'Unique Key(s)'], pk_fk_rows)

    # 3.3. M:N Relationship Resolution
    heading2(doc, '3.3. M:N Relationship Resolution')

    add_para(doc, (
        'In the original Lab 2 design, there were no explicit many-to-many (M:N) relationships. '
        'However, the conceptual model contained implicit multi-valued attributes which were '
        'resolved through decomposition during the normalization process in Lab 3. The following '
        'decompositions were performed:'
    ))

    add_bullet(doc, 'RobotModel.Features → ModelFeature (junction table): The multi-valued '
               'Features attribute was decomposed into a separate relation with a composite '
               'primary key {ModelID, Feature}.')
    add_bullet(doc, 'MaintenanceRecord.PartsReplaced → ReplacedPart (junction table): The '
               'multi-valued PartsReplaced attribute was decomposed into a separate relation '
               'with a composite primary key {RecordID, PartName}.')
    add_bullet(doc, 'DeviceLog.UsageStatistics → LogStatistic (junction table): The multi-valued '
               'UsageStatistics attribute was decomposed into a separate relation with a composite '
               'primary key {LogID, MetricName}.')

    add_para(doc, (
        'Additionally, the Payment entity was decomposed into three relations — Payment, '
        'OrderPayment, and ServicePayment — using a supertype/subtype pattern. This eliminates '
        'nullable foreign keys (OrderID and ServiceRecordID) that would have existed in a single '
        'Payment table, as each payment is either for an order or for a service, but not both.'
    ), space_before=6)

    doc.add_page_break()


# ==============================================================================
# SECTION 4: PHYSICAL MODEL
# ==============================================================================

def add_section_4(doc):
    """Add Section 4: Physical Model."""
    heading1(doc, '4. Physical Model')

    add_para(doc, (
        'This section maps the logical schema to a physical database design for Microsoft SQL '
        'Server. It includes detailed table definitions with specific data types, identity columns, '
        'default values, and storage considerations appropriate for a production environment.'
    ))

    add_para(doc, '[Insert Physical Model Diagram here — based on Lab4_Relational_Model.drawio '
             'with the SQL Server data types listed in the tables below]',
             bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=12, space_before=12)

    heading2(doc, '4.1. Table Definitions with Data Types')

    add_para(doc, (
        'The following tables define the complete physical schema for all 16 relations. '
        'Each table specifies column names, SQL Server data types, nullability, keys, default '
        'values, and descriptions.'
    ))

    # Table definitions data
    table_defs = {
        'Table 1: Customer': [
            ['CustomerID', 'INT', 'NO', 'PK (IDENTITY)', 'Auto-increment', 'Unique customer identifier'],
            ['FullName', 'NVARCHAR(100)', 'NO', '', '', 'Customer full name'],
            ['PhoneNumber', 'VARCHAR(15)', 'NO', 'UNIQUE', '', 'Contact phone number'],
            ['Email', 'NVARCHAR(100)', 'NO', 'UNIQUE', '', 'Email address'],
            ['Address', 'NVARCHAR(255)', 'YES', '', 'NULL', 'Residential address'],
            ['Password', 'NVARCHAR(255)', 'NO', '', '', 'Hashed account password'],
        ],
        'Table 2: Employee': [
            ['EmployeeID', 'INT', 'NO', 'PK (IDENTITY)', 'Auto-increment', 'Unique employee identifier'],
            ['FullName', 'NVARCHAR(100)', 'NO', '', '', 'Employee full name'],
            ['Role', 'NVARCHAR(50)', 'NO', '', '', 'Job role (Sales Staff/Technician/Administrator)'],
            ['PhoneNumber', 'VARCHAR(15)', 'NO', '', '', 'Contact phone number'],
            ['Email', 'NVARCHAR(100)', 'NO', 'UNIQUE', '', 'Email address'],
            ['Password', 'NVARCHAR(255)', 'NO', '', '', 'Hashed account password'],
        ],
        'Table 3: RobotModel': [
            ['ModelID', 'INT', 'NO', 'PK (IDENTITY)', 'Auto-increment', 'Unique model identifier'],
            ['Brand', 'NVARCHAR(100)', 'NO', '', '', 'Manufacturer brand name'],
            ['ModelName', 'NVARCHAR(100)', 'NO', 'UNIQUE', '', 'Model designation'],
            ['Specifications', 'NVARCHAR(MAX)', 'YES', '', 'NULL', 'Technical specifications'],
            ['UnitPrice', 'DECIMAL(18,2)', 'NO', '', '', 'Retail price per unit'],
            ['WarrantyDuration', 'INT', 'NO', '', '', 'Warranty period in months'],
        ],
        'Table 4: ModelFeature': [
            ['ModelID', 'INT', 'NO', 'PK, FK → RobotModel', '', 'Reference to robot model'],
            ['Feature', 'NVARCHAR(100)', 'NO', 'PK', '', 'Individual feature name'],
        ],
        'Table 5: RobotUnit': [
            ['RobotID', 'INT', 'NO', 'PK (IDENTITY)', 'Auto-increment', 'Unique robot unit identifier'],
            ['ModelID', 'INT', 'NO', 'FK → RobotModel', '', 'Reference to robot model'],
            ['SerialNumber', 'VARCHAR(50)', 'NO', 'UNIQUE', '', 'Manufacturer serial number'],
            ['Status', 'NVARCHAR(30)', 'NO', '', "'Available'", 'Current unit status'],
        ],
        'Table 6: SalesOrder': [
            ['OrderID', 'INT', 'NO', 'PK (IDENTITY)', 'Auto-increment', 'Unique order identifier'],
            ['CustomerID', 'INT', 'NO', 'FK → Customer', '', 'Purchasing customer'],
            ['EmployeeID', 'INT', 'NO', 'FK → Employee', '', 'Processing staff member'],
            ['OrderDate', 'DATETIME', 'NO', '', 'GETDATE()', 'Date/time order was placed'],
            ['TotalAmount', 'DECIMAL(18,2)', 'NO', '', '', 'Total order value'],
            ['OrderStatus', 'NVARCHAR(30)', 'NO', '', "'Pending'", 'Current order status'],
        ],
        'Table 7: OrderDetail': [
            ['RobotID', 'INT', 'NO', 'PK, FK → RobotUnit', '', 'Sold robot unit'],
            ['OrderID', 'INT', 'NO', 'FK → SalesOrder', '', 'Associated sales order'],
            ['SellingPrice', 'DECIMAL(18,2)', 'NO', '', '', 'Actual selling price'],
        ],
        'Table 8: Payment': [
            ['PaymentID', 'INT', 'NO', 'PK (IDENTITY)', 'Auto-increment', 'Unique payment identifier'],
            ['Amount', 'DECIMAL(18,2)', 'NO', '', '', 'Payment amount'],
            ['PaymentDate', 'DATETIME', 'NO', '', 'GETDATE()', 'Date/time of payment'],
            ['PaymentMethod', 'NVARCHAR(30)', 'NO', '', '', 'Method of payment'],
        ],
        'Table 9: OrderPayment': [
            ['PaymentID', 'INT', 'NO', 'PK, FK → Payment', '', 'Reference to payment'],
            ['OrderID', 'INT', 'NO', 'FK → SalesOrder', '', 'Reference to sales order'],
        ],
        'Table 10: ServicePayment': [
            ['PaymentID', 'INT', 'NO', 'PK, FK → Payment', '', 'Reference to payment'],
            ['ServiceRecordID', 'INT', 'NO', 'FK → MaintenanceRecord', '', 'Reference to maintenance record'],
        ],
        'Table 11: WarrantyRegistration': [
            ['WarrantyID', 'INT', 'NO', 'PK (IDENTITY)', 'Auto-increment', 'Unique warranty identifier'],
            ['RobotID', 'INT', 'NO', 'FK → RobotUnit, UNIQUE', '', 'Registered robot unit'],
            ['CustomerID', 'INT', 'NO', 'FK → Customer', '', 'Warranty holder'],
            ['StartDate', 'DATE', 'NO', '', '', 'Warranty start date'],
            ['EndDate', 'DATE', 'NO', '', '', 'Warranty expiration date'],
        ],
        'Table 12: ServiceRequest': [
            ['RequestID', 'INT', 'NO', 'PK (IDENTITY)', 'Auto-increment', 'Unique request identifier'],
            ['RobotID', 'INT', 'NO', 'FK → RobotUnit', '', 'Robot requiring service'],
            ['CustomerID', 'INT', 'NO', 'FK → Customer', '', 'Requesting customer'],
            ['IssueDescription', 'NVARCHAR(MAX)', 'NO', '', '', 'Description of the issue'],
            ['RequestDate', 'DATETIME', 'NO', '', 'GETDATE()', 'Date/time request was submitted'],
            ['Status', 'NVARCHAR(30)', 'NO', '', "'Pending'", 'Current request status'],
        ],
        'Table 13: MaintenanceRecord': [
            ['RecordID', 'INT', 'NO', 'PK (IDENTITY)', 'Auto-increment', 'Unique record identifier'],
            ['RequestID', 'INT', 'NO', 'FK → ServiceRequest, UNIQUE', '', 'Associated service request'],
            ['TechnicianID', 'INT', 'NO', 'FK → Employee', '', 'Assigned technician'],
            ['ActionsTaken', 'NVARCHAR(MAX)', 'YES', '', 'NULL', 'Description of repair actions'],
            ['ServiceFee', 'DECIMAL(18,2)', 'NO', '', '0', 'Service charge amount'],
            ['CompletionDate', 'DATETIME', 'YES', '', 'NULL', 'Date/time of completion'],
        ],
        'Table 14: ReplacedPart': [
            ['RecordID', 'INT', 'NO', 'PK, FK → MaintenanceRecord', '', 'Reference to maintenance record'],
            ['PartName', 'NVARCHAR(100)', 'NO', 'PK', '', 'Name of replaced part'],
        ],
        'Table 15: DeviceLog': [
            ['LogID', 'INT', 'NO', 'PK (IDENTITY)', 'Auto-increment', 'Unique log identifier'],
            ['RobotID', 'INT', 'NO', 'FK → RobotUnit', '', 'Source robot unit'],
            ['LogTime', 'DATETIME', 'NO', '', 'GETDATE()', 'Timestamp of log entry'],
            ['ErrorCode', 'VARCHAR(20)', 'YES', '', 'NULL', 'Error code (if any)'],
        ],
        'Table 16: LogStatistic': [
            ['LogID', 'INT', 'NO', 'PK, FK → DeviceLog', '', 'Reference to device log'],
            ['MetricName', 'NVARCHAR(100)', 'NO', 'PK', '', 'Name of usage metric'],
            ['MetricValue', 'NVARCHAR(255)', 'YES', '', 'NULL', 'Value of the metric'],
        ],
    }

    headers = ['Column', 'Data Type', 'Nullable', 'Key', 'Default', 'Description']

    for title, rows in table_defs.items():
        add_para(doc, title, bold=True, font_size=12, space_before=10, space_after=4,
                 alignment=WD_ALIGN_PARAGRAPH.LEFT)
        create_table(doc, headers, rows)

    doc.add_page_break()


# ==============================================================================
# SECTION 5: CONSTRAINTS
# ==============================================================================

def add_section_5(doc):
    """Add Section 5: Constraints."""
    heading1(doc, '5. Constraints')

    add_para(doc, (
        'This section provides a comprehensive description of all constraints applied to the '
        'database schema. Constraints ensure data integrity, consistency, and reliability by '
        'enforcing business rules at the database level. The constraints are organized into '
        'column-level and table-level categories.'
    ))

    # 5.1. Column-Level Constraints
    heading2(doc, '5.1. Column-Level Constraints')

    # NOT NULL
    add_para(doc, 'NOT NULL Constraints', bold=True, font_size=12, space_before=6)
    add_para(doc, (
        'NOT NULL constraints are applied to all mandatory columns to prevent incomplete records. '
        'The following table highlights key NOT NULL columns across the schema. All primary key '
        'columns and most business-critical columns are defined as NOT NULL.'
    ))

    not_null_rows = [
        ['Customer', 'FullName, PhoneNumber, Email, Password', 'Essential customer information'],
        ['Employee', 'FullName, Role, PhoneNumber, Email, Password', 'Essential employee information'],
        ['RobotModel', 'Brand, ModelName, UnitPrice, WarrantyDuration', 'Core product information'],
        ['RobotUnit', 'ModelID, SerialNumber, Status', 'Inventory tracking fields'],
        ['SalesOrder', 'CustomerID, EmployeeID, OrderDate, TotalAmount, OrderStatus', 'Transaction details'],
        ['OrderDetail', 'OrderID, SellingPrice', 'Order line item details'],
        ['Payment', 'Amount, PaymentDate, PaymentMethod', 'Payment transaction details'],
        ['ServiceRequest', 'RobotID, CustomerID, IssueDescription, RequestDate, Status', 'Service request details'],
        ['MaintenanceRecord', 'RequestID, TechnicianID, ServiceFee', 'Maintenance record details'],
        ['WarrantyRegistration', 'RobotID, CustomerID, StartDate, EndDate', 'Warranty coverage details'],
    ]

    create_table(doc, ['Table', 'NOT NULL Columns', 'Purpose'], not_null_rows)

    # UNIQUE
    add_para(doc, 'UNIQUE Constraints', bold=True, font_size=12, space_before=10)
    add_para(doc, (
        'UNIQUE constraints prevent duplicate values in columns that must contain distinct data. '
        'These are applied in addition to the uniqueness guaranteed by primary keys.'
    ))

    unique_rows = [
        ['Customer', 'Email', 'Prevent duplicate email addresses'],
        ['Customer', 'PhoneNumber', 'Prevent duplicate phone numbers'],
        ['Employee', 'Email', 'Prevent duplicate employee emails'],
        ['RobotModel', 'ModelName', 'Ensure unique model names'],
        ['RobotUnit', 'SerialNumber', 'Ensure unique serial numbers'],
        ['WarrantyRegistration', 'RobotID', 'Each robot has at most one active warranty'],
        ['MaintenanceRecord', 'RequestID', 'Each service request has at most one maintenance record'],
    ]

    create_table(doc, ['Table', 'Column(s)', 'Purpose'], unique_rows)

    # DEFAULT
    add_para(doc, 'DEFAULT Values', bold=True, font_size=12, space_before=10)
    add_para(doc, (
        'DEFAULT constraints provide automatic values for columns when no explicit value is '
        'supplied during insertion. These improve usability and reduce the risk of insertion errors.'
    ))

    default_rows = [
        ['RobotUnit', 'Status', "'Available'", 'New units default to available'],
        ['SalesOrder', 'OrderDate', 'GETDATE()', 'Auto-set order timestamp'],
        ['SalesOrder', 'OrderStatus', "'Pending'", 'New orders start as pending'],
        ['Payment', 'PaymentDate', 'GETDATE()', 'Auto-set payment timestamp'],
        ['ServiceRequest', 'RequestDate', 'GETDATE()', 'Auto-set request timestamp'],
        ['ServiceRequest', 'Status', "'Pending'", 'New requests start as pending'],
        ['MaintenanceRecord', 'ServiceFee', '0', 'Default no charge (warranty coverage)'],
    ]

    create_table(doc, ['Table', 'Column', 'Default Value', 'Purpose'], default_rows)

    # CHECK
    add_para(doc, 'CHECK Constraints', bold=True, font_size=12, space_before=10)
    add_para(doc, (
        'CHECK constraints enforce domain-specific validation rules at the database level, '
        'ensuring that column values fall within acceptable ranges or match predefined lists.'
    ))

    check_rows = [
        ['Employee', 'Role', "IN ('Sales Staff', 'Technician', 'Administrator')", 'Restrict valid roles'],
        ['RobotModel', 'UnitPrice', '> 0', 'Price must be positive'],
        ['RobotModel', 'WarrantyDuration', '> 0', 'Duration must be positive'],
        ['RobotUnit', 'Status', "IN ('Available', 'Sold', 'Under Maintenance', 'Retired')", 'Valid statuses only'],
        ['SalesOrder', 'TotalAmount', '>= 0', 'Non-negative order total'],
        ['SalesOrder', 'OrderStatus', "IN ('Pending', 'Confirmed', 'Shipped', 'Delivered', 'Cancelled')", 'Valid statuses'],
        ['OrderDetail', 'SellingPrice', '> 0', 'Price must be positive'],
        ['Payment', 'Amount', '> 0', 'Payment must be positive'],
        ['Payment', 'PaymentMethod', "IN ('Cash', 'Credit Card', 'Bank Transfer', 'E-Wallet')", 'Valid methods'],
        ['WarrantyRegistration', 'EndDate', 'EndDate > StartDate', 'End must be after start'],
        ['ServiceRequest', 'Status', "IN ('Pending', 'Assigned', 'In Progress', 'Completed', 'Cancelled')", 'Valid statuses'],
        ['MaintenanceRecord', 'ServiceFee', '>= 0', 'Non-negative fee'],
    ]

    create_table(doc, ['Table', 'Column', 'CHECK Expression', 'Purpose'], check_rows)

    # 5.2. Table-Level Constraints
    heading2(doc, '5.2. Table-Level Constraints')

    # PRIMARY KEY
    add_para(doc, 'PRIMARY KEY Constraints', bold=True, font_size=12, space_before=6)
    add_para(doc, (
        'Every table has a primary key constraint that uniquely identifies each row. The following '
        'table lists all primary key definitions, including whether they use IDENTITY auto-increment '
        'or are composite keys.'
    ))

    pk_rows = [
        ['Customer', 'CustomerID', 'IDENTITY (1,1)'],
        ['Employee', 'EmployeeID', 'IDENTITY (1,1)'],
        ['RobotModel', 'ModelID', 'IDENTITY (1,1)'],
        ['ModelFeature', '{ModelID, Feature}', 'Composite'],
        ['RobotUnit', 'RobotID', 'IDENTITY (1,1)'],
        ['SalesOrder', 'OrderID', 'IDENTITY (1,1)'],
        ['OrderDetail', 'RobotID', 'Single column (no IDENTITY)'],
        ['Payment', 'PaymentID', 'IDENTITY (1,1)'],
        ['OrderPayment', 'PaymentID', 'Single column (no IDENTITY)'],
        ['ServicePayment', 'PaymentID', 'Single column (no IDENTITY)'],
        ['WarrantyRegistration', 'WarrantyID', 'IDENTITY (1,1)'],
        ['ServiceRequest', 'RequestID', 'IDENTITY (1,1)'],
        ['MaintenanceRecord', 'RecordID', 'IDENTITY (1,1)'],
        ['ReplacedPart', '{RecordID, PartName}', 'Composite'],
        ['DeviceLog', 'LogID', 'IDENTITY (1,1)'],
        ['LogStatistic', '{LogID, MetricName}', 'Composite'],
    ]

    create_table(doc, ['Table', 'Primary Key Column(s)', 'Type'], pk_rows)

    # FOREIGN KEY
    add_para(doc, 'FOREIGN KEY Constraints', bold=True, font_size=12, space_before=10)
    add_para(doc, (
        'Foreign key constraints maintain referential integrity between related tables. The following '
        'table lists all foreign key relationships with their referential actions (ON DELETE and ON UPDATE).'
    ))

    fk_rows = [
        ['ModelFeature', 'ModelID', 'RobotModel.ModelID', 'CASCADE', 'CASCADE'],
        ['RobotUnit', 'ModelID', 'RobotModel.ModelID', 'NO ACTION', 'CASCADE'],
        ['SalesOrder', 'CustomerID', 'Customer.CustomerID', 'NO ACTION', 'CASCADE'],
        ['SalesOrder', 'EmployeeID', 'Employee.EmployeeID', 'NO ACTION', 'CASCADE'],
        ['OrderDetail', 'RobotID', 'RobotUnit.RobotID', 'NO ACTION', 'NO ACTION'],
        ['OrderDetail', 'OrderID', 'SalesOrder.OrderID', 'CASCADE', 'CASCADE'],
        ['OrderPayment', 'PaymentID', 'Payment.PaymentID', 'CASCADE', 'CASCADE'],
        ['OrderPayment', 'OrderID', 'SalesOrder.OrderID', 'NO ACTION', 'NO ACTION'],
        ['ServicePayment', 'PaymentID', 'Payment.PaymentID', 'CASCADE', 'CASCADE'],
        ['ServicePayment', 'ServiceRecordID', 'MaintenanceRecord.RecordID', 'NO ACTION', 'NO ACTION'],
        ['WarrantyRegistration', 'RobotID', 'RobotUnit.RobotID', 'NO ACTION', 'CASCADE'],
        ['WarrantyRegistration', 'CustomerID', 'Customer.CustomerID', 'NO ACTION', 'CASCADE'],
        ['ServiceRequest', 'RobotID', 'RobotUnit.RobotID', 'NO ACTION', 'CASCADE'],
        ['ServiceRequest', 'CustomerID', 'Customer.CustomerID', 'NO ACTION', 'CASCADE'],
        ['MaintenanceRecord', 'RequestID', 'ServiceRequest.RequestID', 'NO ACTION', 'CASCADE'],
        ['MaintenanceRecord', 'TechnicianID', 'Employee.EmployeeID', 'NO ACTION', 'CASCADE'],
        ['ReplacedPart', 'RecordID', 'MaintenanceRecord.RecordID', 'CASCADE', 'CASCADE'],
        ['DeviceLog', 'RobotID', 'RobotUnit.RobotID', 'NO ACTION', 'CASCADE'],
        ['LogStatistic', 'LogID', 'DeviceLog.LogID', 'CASCADE', 'CASCADE'],
    ]

    create_table(doc, ['Table', 'FK Column', 'References', 'ON DELETE', 'ON UPDATE'], fk_rows)

    # Composite Unique Constraints
    add_para(doc, 'Composite Unique Constraints', bold=True, font_size=12, space_before=10)
    add_para(doc, (
        'Composite unique constraints ensure that combinations of column values are unique, '
        'preventing logical duplicate records.'
    ))

    comp_unique_rows = [
        ['DeviceLog', '{RobotID, LogTime}', 'Prevent duplicate logs for same robot at same time'],
    ]

    create_table(doc, ['Table', 'Columns', 'Purpose'], comp_unique_rows)

    doc.add_page_break()


# ==============================================================================
# SECTION 6: CONCLUSION AND REFLECTION
# ==============================================================================

def add_section_6(doc):
    """Add Section 6: Conclusion and Reflection."""
    heading1(doc, '6. Conclusion and Reflection')

    add_para(doc, (
        'This lab successfully designed a complete relational database for the Household Cleaning '
        'Robot Sales & Maintenance Management System through three structured design phases: '
        'conceptual modeling (Entity-Relationship Diagram), logical design (relational schema with '
        'primary and foreign keys), and physical implementation (SQL Server table definitions with '
        'specific data types and constraints). The design builds upon the requirements analysis from '
        'Lab 1, the ERD from Lab 2, and the normalization work from Lab 3.'
    ))

    # Key Accomplishments
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.space_before = Pt(6)
    run = para.add_run('Key Accomplishments:')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    add_bullet(doc, 'Identified 16 entities capturing all aspects of the business domain — from '
               'customer management and product inventory to sales transactions, warranty tracking, '
               'maintenance services, and IoT device monitoring')
    add_bullet(doc, 'Established 19 relationships with proper cardinalities (1:1, 1:N) ensuring '
               'accurate representation of business rules')
    add_bullet(doc, 'Defined comprehensive data types optimized for Microsoft SQL Server, including '
               'appropriate use of NVARCHAR for Unicode support, DECIMAL for financial precision, '
               'and IDENTITY for auto-increment primary keys')
    add_bullet(doc, 'Specified 40+ constraints ensuring data integrity at multiple levels — '
               'including NOT NULL, UNIQUE, CHECK, DEFAULT, PRIMARY KEY, and FOREIGN KEY constraints')

    # Reflection on Constraints
    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para2.paragraph_format.space_after = Pt(6)
    para2.paragraph_format.space_before = Pt(10)
    run2 = para2.add_run('Reflection on Constraints:')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    run2.bold = True
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    add_para(doc, (
        'The constraints defined in this lab serve as the primary mechanism for enforcing data '
        'integrity at the database level. Each type of constraint addresses a specific aspect of '
        'data quality:'
    ), space_before=4)

    add_bullet(doc, 'NOT NULL constraints prevent incomplete records by requiring values for '
               'all essential fields, ensuring that critical business data is always captured')
    add_bullet(doc, 'UNIQUE constraints eliminate duplicate data in columns such as Email, '
               'PhoneNumber, SerialNumber, and ModelName, maintaining data distinctness')
    add_bullet(doc, 'CHECK constraints enforce business rules directly at the database level — '
               'for example, ensuring prices are positive, dates are logically ordered, and '
               'status values conform to predefined enumerations')
    add_bullet(doc, 'FOREIGN KEY constraints maintain referential integrity between related tables, '
               'with appropriate ON DELETE and ON UPDATE actions (CASCADE for dependent child records, '
               'NO ACTION for important references that should not be implicitly removed)')
    add_bullet(doc, 'DEFAULT values improve usability by automatically populating timestamps '
               '(GETDATE()), initial statuses (\'Pending\', \'Available\'), and zero-value fees, '
               'reducing the burden on application code and minimizing insertion errors')

    add_para(doc, (
        'The combination of these constraints creates a robust, self-documenting database that '
        'enforces business rules independently of the application layer, providing a defense-in-depth '
        'approach to data integrity.'
    ), space_before=6)

    add_para(doc, (
        'Finally, the normalization work completed in Lab 3 (achieving BCNF for all relations) '
        'ensures minimal redundancy and eliminates update anomalies. The constraints defined in '
        'Lab 4 add an additional layer of protection by enforcing domain rules, referential '
        'integrity, and business logic directly within the database schema. Together, normalization '
        'and constraints create a database design that is both structurally sound and operationally '
        'reliable for the Household Cleaning Robot Sales & Maintenance Management System.'
    ), space_before=6)


# ==============================================================================
# MAIN
# ==============================================================================

def main():
    """Generate the complete Lab 4 report."""
    doc = setup_document()

    add_cover_page(doc)
    add_table_of_contents(doc)
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)
    add_section_4(doc)
    add_section_5(doc)
    add_section_6(doc)

    # Save
    output_dir = r'c:\Users\Khuong\Desktop\FPTU\DBI202\LAB4'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'Lab4_report.docx')
    doc.save(output_path)
    print(f'✅ Lab 4 report generated successfully: {output_path}')


if __name__ == '__main__':
    main()
