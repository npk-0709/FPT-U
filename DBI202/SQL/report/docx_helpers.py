"""Cac ham tro giup dinh dang tai lieu Word cho bao cao dbCOMPANY."""

import re

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BODY_FONT = "Times New Roman"
CODE_FONT = "Consolas"
BODY_SIZE = Pt(12)
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
ACCENT_HEX = "1F4E79"
ROW_ALT_HEX = "EDF3F9"
CODE_BG_HEX = "F5F5F0"
NOTE_BG_HEX = "FFF4E5"
GREY = RGBColor(0x59, 0x59, 0x59)
DANGER = RGBColor(0xC0, 0x00, 0x00)

TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|__.+?__)")


# ------------------------------------------------------------ khung tai lieu
def new_document():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.4
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    sizes = {1: 15, 2: 13.5, 3: 12.5}
    for lvl in (1, 2, 3):
        st = doc.styles["Heading %d" % lvl]
        st.font.name = BODY_FONT
        st.font.size = Pt(sizes[lvl])
        st.font.bold = True
        st.font.color.rgb = ACCENT
        st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        st.paragraph_format.space_before = Pt(14 if lvl == 1 else 10)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing = 1.25
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin, sec.right_margin = Cm(3.0), Cm(2.0)
    sec.top_margin, sec.bottom_margin = Cm(2.2), Cm(2.2)
    return doc


def _field(paragraph, instruction, placeholder=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(sep)
    if placeholder:
        t = OxmlElement("w:t")
        t.text = placeholder
        run._r.append(t)
    run._r.append(end)
    return run


def add_footer_page_numbers(section, label="Trang "):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    fr = _field(p, "PAGE", "1")
    fr.font.name = BODY_FONT
    fr.font.size = Pt(10)
    fr.font.color.rgb = GREY
    r2 = p.add_run(" / ")
    r2.font.name = BODY_FONT
    r2.font.size = Pt(10)
    r2.font.color.rgb = GREY
    fr2 = _field(p, "NUMPAGES", "1")
    fr2.font.name = BODY_FONT
    fr2.font.size = Pt(10)
    fr2.font.color.rgb = GREY


def add_toc(doc, levels="1-3"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    _field(p, 'TOC \\o "%s" \\h \\z \\u' % levels,
           "Nhấn Ctrl+A rồi F9 (hoặc bấm chuột phải > Update Field) để hiện mục lục.")


def page_break(doc):
    doc.add_page_break()


def landscape_section(doc):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    sec.left_margin, sec.right_margin = Cm(1.6), Cm(1.6)
    sec.top_margin, sec.bottom_margin = Cm(1.3), Cm(1.3)
    return sec


def portrait_section(doc):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin, sec.right_margin = Cm(3.0), Cm(2.0)
    sec.top_margin, sec.bottom_margin = Cm(2.2), Cm(2.2)
    return sec


# ------------------------------------------------------------------ van ban
def add_runs(paragraph, text, size=None, color=None, font=BODY_FONT):
    """Ho tro cu phap rut gon: **dam**, `ma nguon`, __in nghieng__."""
    for token in TOKEN_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            # Ben trong doan dam van co the con `ma nguon`.
            for i, piece in enumerate(token[2:-2].split("`")):
                if not piece:
                    continue
                run = paragraph.add_run(piece)
                run.bold = True
                if i % 2:
                    run.font.name = CODE_FONT
                    run.font.size = Pt((size or 12) - 1.5)
                else:
                    run.font.name = font
                    if size:
                        run.font.size = Pt(size)
                if color is not None:
                    run.font.color.rgb = color
            continue
        elif token.startswith("__") and token.endswith("__"):
            run = paragraph.add_run(token[2:-2])
            run.italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt((size or 12) - 1.5)
            continue
        else:
            run = paragraph.add_run(token)
        run.font.name = font
        if size:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
    return paragraph


def para(doc, text="", size=None, align=None, space_after=6, first_line=0.8,
         color=None, italic=False, keep_with_next=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    if first_line:
        p.paragraph_format.first_line_indent = Cm(first_line)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if text:
        add_runs(p, text, size=size, color=color)
        if italic:
            for r in p.runs:
                r.italic = True
    return p


def heading(doc, text, level=1, new_page=False):
    h = doc.add_heading(level=level)
    if new_page:
        # Dung page_break_before thay cho add_page_break() de tranh sinh ra
        # doan van rong -> Word co the day thanh mot trang trang.
        h.paragraph_format.page_break_before = True
    run = h.add_run(text)
    run.font.name = BODY_FONT
    run.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    return h


def bullets(doc, items, size=None, style="List Bullet", space_after=3):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.left_indent = Cm(0.9)
        add_runs(p, it, size=size)


def caption(doc, text, above=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2 if above else 4)
    p.paragraph_format.space_after = Pt(10 if above else 12)
    p.paragraph_format.line_spacing = 1.15
    add_runs(p, text, size=10.5)
    for r in p.runs:
        r.italic = True
        r.font.color.rgb = GREY
    return p


# -------------------------------------------------------------------- bang
def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_borders(cell, color="9CB7CF", size=6):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:%s" % edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def table(doc, headers, rows, widths=None, aligns=None, size=10.0,
          header_size=10.0, zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    def fill(cell, text, bold=False, fsize=size, align=None, color=None):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.line_spacing = 1.1
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_runs(p, "" if text is None else str(text), size=fsize, color=color)
        if bold:
            for r in p.runs:
                r.bold = True

    hdr = t.rows[0]
    _repeat_header(hdr)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        fill(c, h, bold=True, fsize=header_size, align="center",
             color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade(c, ACCENT_HEX)
        _cell_borders(c, color=ACCENT_HEX)

    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            al = aligns[ci] if aligns else None
            fill(cells[ci], val, align=al)
            _cell_borders(cells[ci])
            if zebra and ri % 2 == 1:
                _shade(cells[ci], ROW_ALT_HEX)

    if widths:
        for r in t.rows:
            for ci, w in enumerate(widths):
                r.cells[ci].width = Cm(w)

    # Neo hang cuoi voi doan ke tiep de chu thich bang khong bi day sang
    # trang khac mot minh. Voi bang ngan, neo moi hang de ca bang khong bi cat.
    keep_all = len(rows) <= 14
    for row in (t.rows if keep_all else [t.rows[-1]]):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_with_next = True
    return t


# -------------------------------------------------------------- ma nguon SQL
def code_block(doc, code, size=9.0, bg=CODE_BG_HEX, no_split=None):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.text = ""
    _shade(cell, bg)
    _cell_borders(cell, color="C8C8BE", size=4)
    lines = code.strip("\n").split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line if line.strip() else " ")
        run.font.name = CODE_FONT
        run.font.size = Pt(size)
        run.element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT)
    # Khoi ma du ngan thi khong cho Word cat giua hai trang.
    if no_split or (no_split is None and len(lines) <= 26):
        trPr = t.rows[0]._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def note_box(doc, text, title="Ghi chú", bg=NOTE_BG_HEX, size=10.5):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.text = ""
    _shade(cell, bg)
    _cell_borders(cell, color="E0B060", size=6)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    add_runs(p, "**%s.** %s" % (title, text), size=size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def figure(doc, path, width_cm, cap=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(path, width=Cm(width_cm))
    if cap:
        caption(doc, cap)
    return p
