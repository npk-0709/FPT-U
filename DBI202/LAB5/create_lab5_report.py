# -*- coding: utf-8 -*-
"""
Lab 5 Report Generator
=======================
Reads Lab5_HCRS.sql (the single source of truth) and generates the Word report
"Lab5_report.docx" following the Lab 5 submission structure (Objective -> SQL
Queries -> Functions -> Procedures -> Triggers -> Views/Indexes -> Conclusion).

Usage:
    python create_lab5_report.py

Requirement: pip install python-docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

HERE = os.path.dirname(os.path.abspath(__file__))
SQL_FILE = os.path.join(HERE, 'Lab5_HCRS.sql')
OUT_FILE = os.path.join(HERE, 'Lab5_report.docx')

# Map section title in the .sql file -> (number, display title in the report)
SECTION_MAP = {
    'Basic SQL Queries':        ('3.1', 'Basic Queries'),
    'Intermediate SQL Queries': ('3.2', 'Intermediate Queries'),
    'Advanced SQL Queries':     ('3.3', 'Advanced Queries'),
    'User-defined Functions':   ('4',   'Functions'),
    'Stored Procedures':        ('5',   'Stored Procedures'),
    'Triggers':                 ('6',   'Triggers'),
    'Views and Indexes':        ('7',   'Views and Indexes'),
}
QUERY_SECTIONS = ('Basic SQL Queries', 'Intermediate SQL Queries', 'Advanced SQL Queries')


# ============================================================================
# PARSER: read the .sql file via the --## markers
# ============================================================================
def parse_sql(path):
    meta, sections = {}, []
    cur_section = cur_group = cur_item = None

    def finalize_item():
        nonlocal cur_item
        if cur_item is not None:
            cur_item['sql'] = '\n'.join(cur_item.pop('sql_lines')).strip('\n').strip()
            cur_group['items'].append(cur_item)
            cur_item = None

    def ensure_group():
        nonlocal cur_group
        if cur_group is None:
            cur_group = {'title': None, 'items': []}
            cur_section['groups'].append(cur_group)

    with open(path, encoding='utf-8') as f:
        for raw in f:
            line = raw.rstrip('\n')
            s = line.strip()
            if s.startswith('--##'):
                body = s[4:].strip()
                key = body.split('|', 1)[0].strip() if '|' in body else body
                val = body.split('|', 1)[1].strip() if '|' in body else ''
                if body.startswith('META'):
                    rest = body[4:].strip()
                    if '|' in rest:
                        k, v = rest.split('|', 1)
                        meta[k.strip()] = v.strip()
                elif body.startswith('SECTION'):
                    finalize_item()
                    cur_section = {'title': val, 'groups': []}
                    cur_group = None
                    sections.append(cur_section)
                elif body.startswith('GROUP'):
                    finalize_item()
                    cur_group = {'title': val, 'items': []}
                    cur_section['groups'].append(cur_group)
                elif body.startswith('ITEM'):
                    finalize_item()
                    ensure_group()
                    cur_item = {'title': val, 'desc': [], 'sql_lines': []}
                elif body.startswith('DESC'):
                    if cur_item is not None:
                        cur_item['desc'].append(val)
                continue
            # non-marker line
            if cur_item is not None:
                if s.startswith('/*='):       # PART banner -> end of this item's code
                    finalize_item()
                else:
                    cur_item['sql_lines'].append(line)
        finalize_item()
    return meta, sections


# ============================================================================
# Word formatting HELPERS
# ============================================================================
def set_cell_shading(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml('<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color)))


def _set_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_para(doc, text, bold=False, italic=False, size=12,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, space_before=0, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    _set_font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(text), size=size)
    return p


def _heading(doc, text, level, color):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return h


def heading1(doc, text):
    return _heading(doc, text, 1, RGBColor(0, 0, 0))


def heading2(doc, text):
    return _heading(doc, text, 2, RGBColor(0x1F, 0x3B, 0x73))


def heading3(doc, text):
    return _heading(doc, text, 3, RGBColor(0x2E, 0x55, 0x97))


def add_code_block(doc, code):
    """SQL code block: Consolas font, very light gray background, keep line breaks."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.line_spacing = 1.0
    p._p.get_or_add_pPr().append(parse_xml('<w:shd {} w:fill="F3F3F3"/>'.format(nsdecls('w'))))
    lines = code.split('\n')
    for i, ln in enumerate(lines):
        run = p.add_run(ln if ln else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        if i < len(lines) - 1:
            run.add_break()
    return p


def add_screenshot(doc, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    _set_font(p.add_run('[Insert execution result screenshot here - %s]' % label),
              size=10, italic=True, color=RGBColor(0x80, 0x80, 0x80))
    return p


def create_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        _set_font(run, size=10, bold=True)
        set_cell_shading(cell, 'D9E2F3')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            _set_font(cell.paragraphs[0].add_run(str(val)), size=10)
    return table


def add_item(doc, n, item):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run('%d. %s' % (n, item['title'])), size=12, bold=True,
              color=RGBColor(0x33, 0x33, 0x33))
    if item['desc']:
        add_para(doc, ' '.join(item['desc']), size=11, space_after=4)
    add_code_block(doc, item['sql'])
    add_screenshot(doc, item['title'])


# ============================================================================
# FIXED PARTS (Cover, TOC, Objective, DB overview, Conclusion)
# ============================================================================
def setup_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    return doc


def add_cover(doc, meta):
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, 'FPT UNIVERSITY', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'DEPARTMENT OF INFORMATION TECHNOLOGY', bold=True, size=13,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)
    doc.add_paragraph()
    add_para(doc, 'DBI202 - Database Systems', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, meta.get('TITLE', 'Lab 5'), bold=True, size=19,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, meta.get('SYSTEM', ''), italic=True, size=13,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=34)
    doc.add_paragraph()
    add_para(doc, 'Group Members:', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    for i in range(1, 6):
        add_para(doc, '[Member %d - Student ID]' % i, size=12,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_para(doc, 'Date: June 2026', bold=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)
    doc.add_page_break()


def add_toc(doc):
    heading1(doc, 'Table of Contents')
    items = [
        ('1.', 'Objective', 0),
        ('2.', 'Database Design (from Lab 4)', 0),
        ('3.', 'SQL Queries', 0),
        ('', '3.1. Basic Queries', 1),
        ('', '3.2. Intermediate Queries', 1),
        ('', '3.3. Advanced Queries', 1),
        ('4.', 'Functions', 0),
        ('5.', 'Stored Procedures', 0),
        ('6.', 'Triggers', 0),
        ('7.', 'Views and Indexes', 0),
        ('8.', 'Conclusion and Reflection', 0),
    ]
    for num, title, lvl in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        if lvl == 1:
            p.paragraph_format.left_indent = Cm(1.0)
        _set_font(p.add_run(('%s %s' % (num, title)).strip()), size=12, bold=(lvl == 0))
    doc.add_page_break()


def add_objective(doc):
    heading1(doc, '1. Objective')
    add_para(doc, (
        'Lab 5 focuses on practicing SQL programming on the database designed in Lab 4 for the '
        'Household Cleaning Robot Sales & Maintenance Management System (HCRS&MMS). Students write '
        'queries ranging from basic to advanced, and build Functions, Stored Procedures, and '
        'Triggers to enforce business rules and automate tasks, together with Views and Indexes.'))
    add_para(doc, 'Specific objectives:', bold=True, space_before=4, space_after=2)
    add_bullet(doc, 'Write SELECT, WHERE, ORDER BY queries and aggregate functions (COUNT/SUM/AVG/MAX/MIN).')
    add_bullet(doc, 'Write intermediate queries: multi-table JOINs, GROUP BY/HAVING and subqueries.')
    add_bullet(doc, 'Write advanced queries: nested subqueries, EXISTS/IN/ANY/ALL, and set operations UNION/INTERSECT/EXCEPT.')
    add_bullet(doc, 'Build at least 4 Functions, 4 Stored Procedures and 4 Triggers.')
    add_bullet(doc, 'Create Views to simplify queries and Indexes to improve performance.')
    add_para(doc, (
        'All SQL source code is included in the accompanying file "Lab5_HCRS.sql" and can be '
        'executed end-to-end in SQL Server Management Studio (SSMS).'), space_before=6, italic=True)
    doc.add_page_break()


DB_TABLES = [
    ['Customer', 'CustomerID (PK), FullName, PhoneNumber, Email, Address, Password', 'Customers'],
    ['Employee', 'EmployeeID (PK), FullName, Role, PhoneNumber, Email, Password', 'Employees (sales / technical / admin)'],
    ['RobotModel', 'ModelID (PK), Brand, ModelName, Specifications, UnitPrice, WarrantyDuration', 'Robot models (catalog)'],
    ['ModelFeature', 'ModelID (PK,FK), Feature (PK)', 'Features of each robot model'],
    ['RobotUnit', 'RobotID (PK), ModelID (FK), SerialNumber, Status', 'Individual robot units in stock'],
    ['SalesOrder', 'OrderID (PK), CustomerID (FK), EmployeeID (FK), OrderDate, TotalAmount, OrderStatus', 'Sales orders'],
    ['OrderDetail', 'RobotID (PK,FK), OrderID (FK), SellingPrice', 'Order line items'],
    ['Payment', 'PaymentID (PK), Amount, PaymentDate, PaymentMethod', 'Payments (super-type)'],
    ['OrderPayment', 'PaymentID (PK,FK), OrderID (FK)', 'Payments for sales orders'],
    ['ServicePayment', 'PaymentID (PK,FK), ServiceRecordID (FK)', 'Payments for maintenance services'],
    ['WarrantyRegistration', 'WarrantyID (PK), RobotID (FK,UQ), CustomerID (FK), StartDate, EndDate', 'Warranty registrations'],
    ['ServiceRequest', 'RequestID (PK), RobotID (FK), CustomerID (FK), IssueDescription, RequestDate, Status', 'Service / repair requests'],
    ['MaintenanceRecord', 'RecordID (PK), RequestID (FK,UQ), TechnicianID (FK), ActionsTaken, ServiceFee, CompletionDate', 'Maintenance results'],
    ['ReplacedPart', 'RecordID (PK,FK), PartName (PK)', 'Replaced parts'],
    ['DeviceLog', 'LogID (PK), RobotID (FK), LogTime, ErrorCode', 'IoT data from robots'],
    ['LogStatistic', 'LogID (PK,FK), MetricName (PK), MetricValue', 'Metrics derived from logs'],
]


def add_db_overview(doc):
    heading1(doc, '2. Database Design (from Lab 4)')
    add_para(doc, (
        'The HCRS_DB database consists of 16 relations (normalized to BCNF in Lab 4) covering the '
        'main business areas: user management, product & inventory, sales & payment, warranty, '
        'maintenance, and IoT data. The table below summarizes the main entities; the full DDL '
        '(data types, PK/FK/UNIQUE/CHECK/DEFAULT) and sample data are in PART 1 and PART 2 of the '
        'file Lab5_HCRS.sql.'))
    create_table(doc, ['Relation', 'Key Attributes', 'Meaning'], DB_TABLES)
    add_para(doc, (
        'In addition to the 16 tables above, Lab 5 adds a supporting table RobotStatusAudit to '
        'record the history of robot status changes (written automatically by a trigger).'), space_before=6)
    add_para(doc, (
        'Sample data: 8 customers, 8 employees, 6 robot models, 15 robots, 8 sales orders, 8 '
        'payments, 8 warranty registrations, 7 service requests, 6 maintenance records and 10 IoT '
        'log records - enough for every query to return meaningful results.'), space_before=4)
    doc.add_page_break()


def add_conclusion(doc):
    heading1(doc, '8. Conclusion and Reflection')
    add_para(doc, (
        'Through Lab 5, the team has fully practiced SQL programming techniques on a realistic '
        'database for the household cleaning robot sales & maintenance system. From basic queries '
        '(SELECT, WHERE, ORDER BY, aggregate functions), to intermediate queries (JOIN, GROUP BY/'
        'HAVING, subqueries) and advanced queries (nested subqueries, EXISTS/IN/ANY/ALL, UNION/'
        'INTERSECT/EXCEPT), the team has explored the data from many business perspectives.'))
    add_para(doc, 'Key takeaways:', bold=True, space_before=4, space_after=2)
    add_bullet(doc, 'Functions enable logic reuse (e.g., checking warranty status, computing a customer total spending) and can be reused inside Views and Procedures.')
    add_bullet(doc, 'Stored Procedures encapsulate multi-step operations (creating a sales order, registering a warranty, completing maintenance) within transactions, ensuring data integrity when errors occur.')
    add_bullet(doc, 'Triggers enforce business rules at the database level: automatically updating inventory after a sale, logging status changes, preventing deletion of referenced parent records, and waiving service fees while under warranty.')
    add_bullet(doc, 'Views simplify complex queries, while Indexes (single-column and composite) improve the performance of filtering/sorting queries.')
    add_para(doc, (
        'Combining constraints (Lab 4) with Functions/Procedures/Triggers (Lab 5) produces a '
        'database that is both structurally robust and capable of automating business logic, '
        'reducing dependence on the application layer and limiting data errors.'), space_before=6)


# ============================================================================
# BUILD
# ============================================================================
def build_report(meta, sections):
    doc = setup_document()
    add_cover(doc, meta)
    add_toc(doc)
    add_objective(doc)
    add_db_overview(doc)

    emitted_sql_heading = False
    for sec in sections:
        title = sec['title']
        num, disp = SECTION_MAP.get(title, ('', title))
        if title in QUERY_SECTIONS:
            if not emitted_sql_heading:
                heading1(doc, '3. SQL Queries')
                add_para(doc, (
                    'This section presents SQL queries at three increasing levels: basic, '
                    'intermediate and advanced. Each query includes an explanation and a placeholder '
                    'for the execution result screenshot.'))
                emitted_sql_heading = True
            heading2(doc, '%s. %s' % (num, disp))
        else:
            heading1(doc, '%s. %s' % (num, disp))

        for g in sec['groups']:
            n = 0
            if g['title']:
                heading3(doc, g['title'])
            for it in g['items']:
                n += 1
                add_item(doc, n, it)

        doc.add_page_break()

    add_conclusion(doc)
    doc.save(OUT_FILE)
    return doc


def main():
    if not os.path.exists(SQL_FILE):
        raise SystemExit('File not found: %s' % SQL_FILE)
    meta, sections = parse_sql(SQL_FILE)

    total_items = sum(len(g['items']) for s in sections for g in s['groups'])
    print('Parsed %s' % SQL_FILE)
    for s in sections:
        cnt = sum(len(g['items']) for g in s['groups'])
        print('  - Section "%s": %d items' % (s['title'], cnt))
    print('Total: %d items across %d sections.' % (total_items, len(sections)))

    build_report(meta, sections)
    print('Report generated: %s' % OUT_FILE)


if __name__ == '__main__':
    main()
