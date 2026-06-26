# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_img_placeholder(doc, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[INSERT IMAGE: {caption}]\n")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 0, 0)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"(Please insert image: {caption})")
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(128, 128, 128)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    header_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2E74B5')
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(text)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(row.cells[c_idx], 'D6E4F0')
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

def create_report_extension():
    doc = Document()

    # PAGE SETUP
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # STYLES
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.5

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        if level == 1:
            heading_style.font.size = Pt(16)
        elif level == 2:
            heading_style.font.size = Pt(14)
        elif level == 3:
            heading_style.font.size = Pt(13)

    # DOC START
    doc.add_heading('ADDITIONAL REPORT CONTENT - FEATURE EXTENSION', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run("This file contains the content that needs to be appended to the current report file. The locations to insert images are clearly marked.")

    # 9.
    doc.add_heading('9. Feature Extension – Volunteer Management', level=1)
    
    # 9.1
    doc.add_heading('9.1. Introduction to the New Feature', level=2)
    doc.add_paragraph('In this extended version, the Mountain Hiking Challenge Registration system is supplemented with the Volunteer Management module. The objectives of this extension are:')
    
    reqs = [
        'Refactoring: Apply the Inheritance principle by creating an abstract class named Person as the common superclass for both Student and Volunteer, which helps to reuse source code and comply with OOP principles.',
        'Adding Volunteer module: Allows managing the list of volunteers supporting the mountain hiking trips, including CRUD operations and shift assignment functions (Assign to Shift).',
        'Extending Acceptable: Adding new regex patterns to validate the input data for Volunteer.',
        'Extending Main Menu: Add menu option 9 (Volunteer Management) with its own sub-menu.'
    ]
    for r in reqs:
        doc.add_paragraph(r, style='List Bullet')
        
    doc.add_paragraph('New files added:')
    add_styled_table(doc, ['No.', 'File Name', 'Description'], [
        ['1', 'Person.java', 'Abstract class – common superclass'],
        ['2', 'Skill.java', 'Enum – list of volunteer skills'],
        ['3', 'Volunteer.java', 'Class – Volunteer entity'],
        ['4', 'Volunteers.java', 'Class – manages the list of Volunteers']
    ], col_widths=[1.5, 4, 9])
    
    doc.add_paragraph('Modified files:')
    add_styled_table(doc, ['No.', 'File Name', 'Modifications'], [
        ['1', 'Student.java', 'Refactored: inherits from Person instead of declaring id and name directly'],
        ['2', 'Acceptable.java', 'Added 3 new regex patterns for Volunteer validation'],
        ['3', 'Main.java', 'Added Volunteer Management menu and its corresponding handling methods']
    ], col_widths=[1.5, 4, 9])
    
    # 9.2
    doc.add_heading('9.2. Class Hierarchy Design (Inheritance – Abstract Class)', level=2)
    doc.add_paragraph('Prior to the extension, the Student class had the id and name attributes declared directly. When adding the Volunteer object (which also requires id and name), code duplication would be inevitable without refactoring.')
    doc.add_paragraph('Solution: Create an abstract class Person containing common attributes and methods, then have both Student and Volunteer inherit from Person.')
    
    add_img_placeholder(doc, "Overall System Class Diagram after refactoring – including all classes")
    
    # 9.3
    doc.add_heading('9.3. Description of New Classes', level=2)
    
    doc.add_heading('9.3.1. Abstract Class – Person', level=3)
    doc.add_paragraph('File: Person.java | Lines of code: 36 lines')
    add_styled_table(doc, ['Component', 'Description'], [
        ['abstract class Person', 'Abstract class, cannot be instantiated directly'],
        ['implements Serializable', 'Allows serialization/deserialization of objects to save to the .dat file'],
        ['protected String id, name', 'Common attributes, accessible from subclasses'],
        ['abstract getDisplayInfo()', 'Abstract method – forces subclasses to override it']
    ], col_widths=[5, 9])
    
    doc.add_heading('9.3.2. Enum – Skill', level=3)
    doc.add_paragraph('File: Skill.java | Lines of code: 20 lines')
    add_styled_table(doc, ['Value', 'Meaning'], [
        ['MEDIC', 'Volunteer with medical/first aid skills'],
        ['LOGISTIC', 'Volunteer responsible for logistics and transportation'],
        ['GUIDE_ASSIST', 'Volunteer assisting the tour guides']
    ], col_widths=[4, 10])
    
    doc.add_heading('9.3.3. Class – Volunteer', level=3)
    doc.add_paragraph('File: Volunteer.java | Lines of code: 115 lines')
    doc.add_paragraph('Declared to inherit from Person (public class Volunteer extends Person implements Comparable<Volunteer>).')
    add_styled_table(doc, ['Attribute', 'Type', 'Description'], [
        ['skill', 'Skill', 'Specialized skill'],
        ['maxShiftsPerDay', 'int', 'Maximum number of shifts per day (1–3)'],
        ['shiftsToday', 'int', 'Number of shifts accepted for the current day']
    ], col_widths=[4, 3, 7])
    
    add_styled_table(doc, ['Method', 'Description'], [
        ['assign()', 'Assign shift: increments shiftsToday by 1 if max is not reached. Returns true if successful.'],
        ['hasSkillFor(Skill)', 'Checks if the volunteer has the required skill. GENERAL slot accepts anyone.'],
        ['getDisplayInfo()', 'Overridden from Person – displays information in a formatted table.'],
        ['toCsv()', 'Converts the information to CSV format.']
    ], col_widths=[4, 10])

    doc.add_heading('9.3.4. Class – Volunteers', level=3)
    doc.add_paragraph('File: Volunteers.java | Lines of code: 150 lines')
    doc.add_paragraph('Inherits from ArrayList<Volunteer> and manages the Volunteer list. Supports methods such as add, searchById, delete, showAll, readFromFile, saveToFile (saves to volunteers.dat and exports CSV).')
    
    # 9.4
    doc.add_heading('9.4. Modifications to Existing Classes', level=2)
    
    doc.add_heading('9.4.1. Student (refactored)', level=3)
    add_styled_table(doc, ['Before (Old)', 'After (New)'], [
        ['implements Serializable, Comparable<Student>', 'extends Person implements Comparable<Student>'],
        ['Declared private String id; private String name;', 'Inherits id and name from Person (protected)'],
        ['Constructor: this.id = id; this.name = name;', 'Constructor: super(id, name);'],
        ['No getDisplayInfo() method', 'Overrides getDisplayInfo() returning toString()']
    ], col_widths=[7, 7])
    
    doc.add_heading('9.4.2. Acceptable (updated)', level=3)
    add_styled_table(doc, ['Pattern', 'Explanation'], [
        ['VOLUNTEER_ID', 'Starts with "VL" (case-insensitive) + 3 digits'],
        ['VOLUNTEER_NAME_VALID', '3–30 alphabetic characters (including Vietnamese) + spaces'],
        ['SHIFT_VALID', 'Accepts values 1, 2, or 3 only']
    ], col_widths=[5, 9])
    
    doc.add_heading('9.4.3. Main (updated)', level=3)
    doc.add_paragraph('Main Menu modifications: Expanded from 9 options to 10 options, adding "9. Volunteer Management".')
    add_img_placeholder(doc, "Screenshot of Main Menu showing all 10 options")
    
    doc.add_paragraph('Added Volunteer Management sub-menu with 6 features (Add, Display, Update, Assign, Delete, Back).')
    add_img_placeholder(doc, "Screenshot of Volunteer Management sub-menu")
    
    # 9.5
    doc.add_heading('9.5. Updated Class Diagram', level=2)
    add_img_placeholder(doc, "Complete Class Diagram (PlantUML or Draw.io) – including all 13 classes/interfaces/enums")
    
    # 9.6
    doc.add_heading('9.6. Volunteer Management Features', level=2)
    
    doc.add_heading('9.6.1. Add New Volunteer', level=3)
    add_img_placeholder(doc, "Screenshot of successfully adding a new Volunteer")
    add_img_placeholder(doc, "Screenshot of duplicate ID entry – displaying error message")
    
    doc.add_heading('9.6.2. Display Volunteer List', level=3)
    add_img_placeholder(doc, "Screenshot of the full Volunteer list")
    add_img_placeholder(doc, "Screenshot of an empty Volunteer list")
    
    doc.add_heading('9.6.3. Update Volunteer', level=3)
    add_img_placeholder(doc, "Screenshot of Update Volunteer process – showing old and new information")
    
    doc.add_heading('9.6.4. Assign Volunteer to Shift', level=3)
    add_img_placeholder(doc, "Screenshot of successful Assignment")
    add_img_placeholder(doc, "Screenshot of Over shift limit error")
    add_img_placeholder(doc, "Screenshot of insufficient skill for MEDIC slot error")
    
    doc.add_heading('9.6.5. Delete Volunteer', level=3)
    add_img_placeholder(doc, "Screenshot of deleting Volunteer – confirming with Y")
    add_img_placeholder(doc, "Screenshot of cancelling deletion – pressing N")
    
    # 9.7
    doc.add_heading('9.7. Volunteer Data Storage Integration', level=2)
    doc.add_paragraph('The saveDataToFile() method was updated to save both Student (registrations.dat) and Volunteer (volunteers.dat) data.')
    add_img_placeholder(doc, "Screenshot of Save Data result – showing both Student and Volunteer saved successfully")
    doc.add_paragraph('Exit Program: Checks both data sources. Warns if any data is unsaved.')
    add_img_placeholder(doc, "Screenshot of Exit prompt when there are unsaved changes")
    
    # 9.8
    doc.add_heading('9.8. Conclusion of the Extension', level=2)
    doc.add_paragraph('The Volunteer Management extension has successfully:')
    kl = [
        'Applied Inheritance: Created abstract class Person as the superclass.',
        'Applied Polymorphism: getDisplayInfo() is overridden in each subclass.',
        'Applied Encapsulation: Utilized Skill enum, and Volunteers class to manage the list.',
        'Integrated fully and operated smoothly with the existing Student system.'
    ]
    for k in kl:
        doc.add_paragraph(k, style='List Bullet')
        
    doc.add_paragraph('The total lines of code after the extension reached 1,518 lines with 13 class/enum/interface files.')

    # ============================================================
    # 10. AUTHORIZATION (RBAC)
    # ============================================================
    doc.add_page_break()
    doc.add_heading('10. Feature Extension – Role-Based Access Control (RBAC)', level=1)

    # 10.1
    doc.add_heading('10.1. Introduction to the Authorization Feature', level=2)
    doc.add_paragraph('In this next extended version, the system is supplemented with Authentication (Login) and Role-Based Access Control (RBAC). Previously, anyone running the program could use every function; now each user must log in and only sees/uses the functions appropriate to their role.')

    reqs = [
        'Authentication: Users must log in with a username/password before entering the system (up to 3 attempts).',
        'Authorization (RBAC): Each account is bound to a Role; each Role owns a set of Permissions. The menu is generated dynamically – only functions the user is permitted to use are displayed.',
        'Applying OOP: Separating Permission – Role – Account into independent components, each class responsible for its own task (Single Responsibility).',
        'Account management: Added an Account Management menu for administrators (ADMIN) to view, add, and delete accounts.',
        'Clean code refactor: Replaced the bulky switch-case menu with a "menu engine" using a list of permission-aware MenuItem objects.'
    ]
    for r in reqs:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_paragraph('New files added:')
    add_styled_table(doc, ['No.', 'File Name', 'Lines', 'Description'], [
        ['1', 'Permission.java', '20', 'Enum – lists the functional permissions of the system'],
        ['2', 'Role.java', '44', 'Enum – roles, each bound to a set of Permissions'],
        ['3', 'Account.java', '70', 'Class – user account, inherits from Person'],
        ['4', 'Accounts.java', '128', 'Class – manages the account list and handles login'],
        ['5', 'MenuItem.java', '24', 'Class – a menu entry bound to a permission and an action'],
    ], col_widths=[1.5, 3.5, 2, 8])

    doc.add_paragraph('Modified files:')
    add_styled_table(doc, ['No.', 'File Name', 'Modifications'], [
        ['1', 'Main.java', 'Added login flow, permission-based dynamic menu, and Account Management; replaced switch-case with a menu engine'],
        ['2', 'Acceptable.java', 'Added 2 regex patterns: USERNAME_VALID and PASSWORD_VALID'],
        ['3', 'Person.java', 'Reused as the superclass for Account (id = username, name = display name)']
    ], col_widths=[1.5, 3.5, 10])

    # 10.2
    doc.add_heading('10.2. The RBAC Model', level=2)
    doc.add_paragraph('RBAC (Role-Based Access Control) is an access-control model based on roles. Instead of assigning permissions directly to each user, the system maps: User → Role → Set of Permissions. This centralizes permission management, makes it easy to extend, and aligns well with OOP principles.')
    doc.add_paragraph('Relationship between components:')
    add_styled_table(doc, ['Component', 'Role in the Model'], [
        ['Account', 'Represents a logged-in user; holds a reference to one Role.'],
        ['Role', 'Represents a role; holds a Set<Permission> (using EnumSet).'],
        ['Permission', 'Represents an access right to a specific group of functions.'],
        ['MenuItem', 'Each menu entry declares its required permission; it checks whether the account is allowed.']
    ], col_widths=[4, 10])
    add_img_placeholder(doc, "RBAC model diagram: Account -> Role -> Set<Permission>")

    # 10.3
    doc.add_heading('10.3. Description of New Components', level=2)

    doc.add_heading('10.3.1. Enum – Permission', level=3)
    doc.add_paragraph('File: Permission.java | Lines of code: 20 lines')
    doc.add_paragraph('Lists 8 permissions corresponding to functional groups, each with a description:')
    add_styled_table(doc, ['Permission', 'Allows'], [
        ['CREATE_REGISTRATION', 'Create a new registration'],
        ['UPDATE_REGISTRATION', 'Update registration information'],
        ['VIEW_REGISTRATION', 'View / search / filter the registration list'],
        ['DELETE_REGISTRATION', 'Delete a registration'],
        ['VIEW_STATISTICS', 'View statistics by location'],
        ['SAVE_DATA', 'Save data to file'],
        ['MANAGE_VOLUNTEER', 'Manage volunteers'],
        ['MANAGE_ACCOUNT', 'Manage accounts (ADMIN only)']
    ], col_widths=[5, 9])

    doc.add_heading('10.3.2. Enum – Role', level=3)
    doc.add_paragraph('File: Role.java | Lines of code: 44 lines')
    doc.add_paragraph('Each role constant is initialized with an EnumSet<Permission>. The method has(Permission) checks whether the role contains that permission.')
    add_styled_table(doc, ['Role', 'Permission Scope'], [
        ['ADMIN', 'Full access (EnumSet.allOf) – every function including account management'],
        ['STAFF', 'Create, update, view, statistics, save and volunteer management (no delete, no account management)'],
        ['VIEWER', 'View list and statistics only (read-only)']
    ], col_widths=[3, 11])

    doc.add_heading('10.3.3. Class – Account', level=3)
    doc.add_paragraph('File: Account.java | Lines of code: 70 lines')
    doc.add_paragraph('Inherits from Person (public class Account extends Person), reusing id as the username and name as the display name.')
    add_styled_table(doc, ['Component', 'Description'], [
        ['password (private)', 'The password is fully encapsulated with NO getter to prevent leakage.'],
        ['role (Role)', 'The role of the account.'],
        ['authenticate(pwd)', 'The account validates the entered password itself (returns true/false).'],
        ['can(Permission)', 'Delegates to Role for checking: role.has(permission).'],
        ['getDisplayInfo()', 'Overridden from Person – displays username | name | role.']
    ], col_widths=[5, 9])

    doc.add_heading('10.3.4. Class – Accounts', level=3)
    doc.add_paragraph('File: Accounts.java | Lines of code: 128 lines')
    doc.add_paragraph('Inherits from ArrayList<Account>, manages the account list and saves/loads the accounts.dat file (Serialization). If the file is empty, the system automatically seeds 3 default accounts. It provides login(username, password) returning an Account if valid, plus searchByUsername, add, delete, and showAll.')

    doc.add_heading('10.3.5. Class – MenuItem', level=3)
    doc.add_paragraph('File: MenuItem.java | Lines of code: 24 lines')
    doc.add_paragraph('Encapsulates a menu entry consisting of: a display label, a required Permission, and an action (Runnable). The method isAllowedFor(account) returns true when the entry requires no permission or the account has the matching permission – so the menu only shows entries the user is allowed to access.')

    # 10.4
    doc.add_heading('10.4. Modifications to Existing Classes', level=2)

    doc.add_heading('10.4.1. Main (refactored – menu engine + login)', level=3)
    add_styled_table(doc, ['Before (Old)', 'After (New)'], [
        ['Enters the menu directly, no login', 'Requires login() first, up to 3 attempts'],
        ['Large switch-case (case 1..10)', 'runMenu() iterates a list of MenuItem, generating the menu dynamically'],
        ['Fixed menu for everyone', 'Only shows entries the currentUser is permitted to use (visibleItems)'],
        ['No account management', 'Added Account Management menu (ADMIN)']
    ], col_widths=[7, 7])
    add_img_placeholder(doc, "Screenshot of the LOGIN screen")

    doc.add_heading('10.4.2. Acceptable (updated)', level=3)
    add_styled_table(doc, ['Pattern', 'Explanation'], [
        ['USERNAME_VALID', '3–20 characters consisting of letters, digits, or underscore'],
        ['PASSWORD_VALID', '6–20 characters, no whitespace']
    ], col_widths=[5, 9])

    # 10.5
    doc.add_heading('10.5. Permission Matrix (Role × Permission)', level=2)
    doc.add_paragraph('The table below summarizes the permissions of each role (Yes = allowed, No = not allowed):')
    add_styled_table(doc, ['Function', 'ADMIN', 'STAFF', 'VIEWER'], [
        ['New Registration', 'Yes', 'Yes', 'No'],
        ['Update Registration', 'Yes', 'Yes', 'No'],
        ['View / Search / Filter', 'Yes', 'Yes', 'Yes'],
        ['Delete Registration', 'Yes', 'No', 'No'],
        ['Statistics', 'Yes', 'Yes', 'Yes'],
        ['Save Data', 'Yes', 'Yes', 'No'],
        ['Volunteer Management', 'Yes', 'Yes', 'No'],
        ['Account Management', 'Yes', 'No', 'No'],
    ], col_widths=[7, 2.3, 2.3, 2.3])

    # 10.6
    doc.add_heading('10.6. Default Accounts', level=2)
    doc.add_paragraph('On the first run (when accounts.dat does not yet exist), the system creates 3 default accounts to demonstrate authorization:')
    add_styled_table(doc, ['Username', 'Password', 'Role'], [
        ['admin', '123456', 'ADMIN'],
        ['staff', '123456', 'STAFF'],
        ['viewer', '123456', 'VIEWER'],
    ], col_widths=[4, 4, 4])

    # 10.7
    doc.add_heading('10.7. Authorization Feature Demonstrations', level=2)

    doc.add_heading('10.7.1. Login', level=3)
    add_img_placeholder(doc, "Screenshot of a successful login with the admin account")
    add_img_placeholder(doc, "Screenshot of a wrong password – showing the remaining attempts")

    doc.add_heading('10.7.2. Menu Displayed by Role', level=3)
    add_img_placeholder(doc, "Screenshot of the ADMIN menu – showing all 10 functions + Exit")
    add_img_placeholder(doc, "Screenshot of the VIEWER menu – showing only view/statistics functions")

    doc.add_heading('10.7.3. Account Management', level=3)
    add_img_placeholder(doc, "Screenshot of the account list")
    add_img_placeholder(doc, "Screenshot of adding a new account and selecting a role")
    add_img_placeholder(doc, "Screenshot of deleting an account (and the case of not allowing deletion of the currently logged-in account)")

    # 10.8
    doc.add_heading('10.8. Conclusion of the Authorization Extension', level=2)
    doc.add_paragraph('The authorization extension has achieved:')
    kl2 = [
        'Applied standard RBAC: clearly separating Permission – Role – Account, making it easy to add new roles/permissions without modifying existing logic (Open/Closed Principle).',
        'Encapsulation: the password is hidden, and authentication and permission checks are handled by the object itself.',
        'Inheritance & Polymorphism: Account inherits from Person and overrides getDisplayInfo().',
        'Cleaner code: the MenuItem-based menu engine eliminates the long switch-case while automatically hiding functions the user is not authorized to use.'
    ]
    for k in kl2:
        doc.add_paragraph(k, style='List Bullet')

    doc.add_paragraph('After adding authorization, the project has a total of 18 class/enum/interface files with approximately 1,844 lines of code.')

    # SAVE
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'Report_Extension_MountainHiking_EN.docx'
    )
    doc.save(output_path)
    print(f"[OK] Extension report created at: {output_path}")

if __name__ == '__main__':
    create_report_extension()
