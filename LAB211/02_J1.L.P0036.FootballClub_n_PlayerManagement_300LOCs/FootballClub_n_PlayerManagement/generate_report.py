# -*- coding: utf-8 -*-
"""Generate BAO_CAO Word document for Football Club & Player Management project."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "BAO_CAO_FootballClub_PlayerManagement.docx")


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return h


def add_para(doc, text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    return p


def add_image_placeholder(doc, caption):
    """Leave empty box for user to insert screenshot."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n[ CHÈN ẢNH: {caption} ]\n\n")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Hình: {caption}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.italic = True


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "D9E2F3")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    doc.add_paragraph()
    return table


def build_report():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("FPT UNIVERSITY\n")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Times New Roman"

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("BÁO CÁO DỰ ÁN LAB211\n")
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.name = "Times New Roman"

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run(
        "HỆ THỐNG QUẢN LÝ CÂU LẠC BỘ VÀ CẦU THỦ\n"
        "European Elite League (EEL) Management System\n\n"
        "Mã đề: J1.L.P0036 | LOC: 300\n"
    )
    r3.font.size = Pt(14)
    r3.bold = True
    r3.font.name = "Times New Roman"

    add_para(doc, "Môn: LAB211 – Java Programming", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Học kỳ: SU26", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "GVHD: Lê Võ Minh Thư", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Sinh viên thực hiện: ....................................", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "MSSV: ....................................", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Lớp: ....................................", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Mục lục placeholder
    add_heading(doc, "MỤC LỤC", 1)
    toc_items = [
        "1. Giới thiệu (Introduction)",
        "2. Phân tích bài toán (Problem Analysis)",
        "3. Computational Thinking",
        "   3.1 Decomposition – Phân rã",
        "   3.2 Pattern Recognition – Nhận diện mẫu",
        "   3.3 Abstraction – Trừu tượng hóa",
        "   3.4 Algorithm Design – Thiết kế thuật toán",
        "4. Thiết kế hướng đối tượng (OOP Design)",
        "   4.1 Use Case Diagram",
        "   4.2 Class Diagram",
        "   4.3 Sequence Diagram",
        "   4.4 Bốn nguyên lý OOP trong dự án",
        "5. Cấu trúc dự án & Triển khai",
        "6. Sơ đồ luồng (Flowchart)",
        "7. Kiểm thử & Kết quả chạy chương trình",
        "8. Kết luận",
    ]
    for item in toc_items:
        add_para(doc, item)
    doc.add_page_break()

    # 1. Introduction
    add_heading(doc, "1. Giới thiệu (Introduction)", 1)
    add_para(doc,
        "Giải đấu European Elite League (EEL) cần một hệ thống số hóa để quản lý thông tin "
        "các câu lạc bộ bóng đá, cầu thủ chính thức và cầu thủ trẻ (academy). Công ty phần mềm "
        "FTC Software giao nhiệm vụ xây dựng ứng dụng console bằng Java theo hướng lập trình "
        "hướng đối tượng (OOP), đảm bảo validation chặt chẽ và lưu trữ dữ liệu trên file văn bản.")
    add_para(doc,
        "Mục tiêu dự án: Xây dựng hệ thống quản lý hoàn chỉnh gồm đăng nhập phân quyền, "
        "CRUD câu lạc bộ/cầu thủ/cầu thủ trẻ, tìm kiếm – lọc – sắp xếp, đọc/ghi file và thoát "
        "có tự động lưu khi có thay đổi.")
    add_para(doc,
        "Phạm vi: Ứng dụng chạy trên console (JDK), dữ liệu lưu tại clubs.txt, players.txt "
        "và youth_players.txt. Hệ thống hỗ trợ ba vai trò: ADMIN (toàn quyền), MANAGER "
        "(quản lý cầu thủ & academy), VIEWER (chỉ xem).")
    add_para(doc, "Đối tượng sử dụng: Quản lý giải đấu EEL, nhân viên vận hành dữ liệu câu lạc bộ.")

    # 2. Problem Analysis
    add_heading(doc, "2. Phân tích bài toán (Problem Analysis)", 1)
    add_heading(doc, "2.1 Bối cảnh và thực thể dữ liệu", 2)
    add_para(doc,
        "Hệ thống quản lý ba nhóm thực thể chính: Club (câu lạc bộ), Player (cầu thủ đội một) "
        "và YouthPlayer (cầu thủ academy). Quan hệ: một Club có nhiều Player và nhiều YouthPlayer; "
        "Player/YouthPlayer tham chiếu Club qua clubID (khóa ngoại).")
    add_table(doc,
        ["Thực thể", "Thuộc tính", "Định dạng file"],
        [
            ["Club", "clubID, name, sponsor, budget", "clubs.txt"],
            ["Player", "playerID, clubID, name, position, shirtNumber", "players.txt"],
            ["YouthPlayer", "youthID, clubID, name, age", "youth_players.txt"],
        ])

    add_heading(doc, "2.2 Ràng buộc dữ liệu (Validation Rules)", 2)
    add_table(doc,
        ["Trường", "Quy tắc", "Thông báo lỗi"],
        [
            ["Club ID", "Duy nhất; định dạng CL-xxxx", "This club ID already exists!"],
            ["Club name / Sponsor", "Không rỗng", "Input cannot be empty"],
            ["Budget", "Số thực dương (triệu EUR)", "Value must be a positive number"],
            ["Player ID", "Duy nhất; định dạng Pxxxx", "This player ID already exists!"],
            ["Club ID (player)", "Phải tồn tại trong danh sách club", "This club does not exist!"],
            ["Position", "Goalkeeper/Defender/Midfielder/Forward/Winger", "Invalid position!"],
            ["Shirt number", "1–99, duy nhất trong cùng club", "This shirt number already exists in this club!"],
            ["Youth ID", "Duy nhất; định dạng AC-xxxx", "This youth player ID already exists!"],
            ["Age (youth)", "8–21; đủ 18 tuổi → đủ điều kiện lên đội một", "Age must be between 8 and 21"],
        ])

    add_heading(doc, "2.3 Danh sách chức năng hệ thống", 2)
    add_table(doc,
        ["STT", "Chức năng", "Quyền yêu cầu"],
        [
            ["1", "List all clubs", "CLUB_VIEW"],
            ["2", "Add a new club", "CLUB_MANAGE"],
            ["3", "Search for a club by ID", "CLUB_VIEW"],
            ["4", "Update a club by ID", "CLUB_MANAGE"],
            ["5", "List clubs with budget ≤ input", "CLUB_VIEW"],
            ["6", "List players sorted by club name, shirt number", "PLAYER_VIEW"],
            ["7", "Search players by partial name", "PLAYER_VIEW"],
            ["8", "Add a new player", "PLAYER_MANAGE"],
            ["9", "Remove a player by ID", "PLAYER_MANAGE"],
            ["10", "Update a player by ID", "PLAYER_MANAGE"],
            ["11", "List players by position", "PLAYER_VIEW"],
            ["12", "List all youth players", "YOUTH_VIEW"],
            ["13", "Add a new youth player", "YOUTH_MANAGE"],
            ["14", "Update a youth player by ID", "YOUTH_MANAGE"],
            ["15", "Remove a youth player by ID", "YOUTH_MANAGE"],
            ["16", "List youth players eligible for first team", "YOUTH_VIEW"],
            ["17", "Save data to files", "DATA_PERSIST"],
            ["18", "Load data from files", "DATA_PERSIST"],
            ["19", "Quit", "Tất cả vai trò đã đăng nhập"],
        ])

    add_heading(doc, "2.4 Trường hợp biên (Edge Cases)", 2)
    add_table(doc,
        ["Nhóm", "Trường hợp", "Xử lý trong code"],
        [
            ["Danh sách rỗng", "Chưa có club/player/youth", "In thông báo \"The ... list is empty.\""],
            ["Trùng khóa", "Thêm ID đã tồn tại", "Chặn và báo lỗi"],
            ["Khóa ngoại sai", "clubID không tồn tại khi thêm player", "Báo \"This club does not exist!\""],
            ["Nhập sai kiểu", "Nhập chữ khi cần số", "Vòng lặp nhập lại qua Inputter"],
            ["Load file lỗi", "Dòng sai format hoặc vi phạm ràng buộc", "loadStrict() trả false → \"Load data failed!\""],
            ["Thoát chưa lưu", "dirty = true", "quit() tự gọi saveAll()"],
            ["Đăng nhập sai", "Quá 3 lần", "Thoát chương trình"],
        ])

    # 3. CT
    add_heading(doc, "3. Computational Thinking", 1)

    add_heading(doc, "3.1 Decomposition – Phân rã", 2)
    add_para(doc, "Phân rã theo dữ liệu:")
    add_para(doc, "• Tầng Model: Club, Person → Player, YouthPlayer, User, Role, Permission, Validatable")
    add_para(doc, "• Tầng Business: ClubsManager, PlayersManager, YouthPlayersManager, AuthManager")
    add_para(doc, "• Tầng Dispatcher: Menu, MenuItem")
    add_para(doc, "• Tầng Tools: Inputter (nhập liệu & validation cơ bản)")
    add_para(doc, "• Tầng Persistence: clubs.txt, players.txt, youth_players.txt")
    add_para(doc, "Phân rã theo chức năng: Nhóm Club (1–5), Player (6–11), Youth (12–16), System (17–19).")

    add_heading(doc, "3.2 Pattern Recognition – Nhận diện mẫu", 2)
    add_para(doc,
        "• Mẫu CRUD lặp lại: add → validate → kiểm tra trùng ID → thêm vào List → đánh dấu dirty.")
    add_para(doc,
        "• Mẫu validation REGEX tập trung tại interface Validatable (CL-xxxx, Pxxxx, AC-xxxx).")
    add_para(doc,
        "• Mẫu loadStrict: đọc file → parse từng dòng → validate → build list tạm → gán nếu toàn bộ hợp lệ.")
    add_para(doc,
        "• Mẫu phân quyền: MenuItem gắn Permission → User.can() → lọc menu động.")

    add_heading(doc, "3.3 Abstraction – Trừu tượng hóa", 2)
    add_table(doc,
        ["Lớp trừu tượng", "Thuộc tính/hành vi cốt lõi", "Bỏ qua"],
        [
            ["Person", "id, name, getDisplayInfo()", "Chi tiết vị trí, số áo, tuổi"],
            ["Validatable", "REGEX, isValid(), isPosition()", "Logic nhập liệu cụ thể"],
            ["MenuItem", "label, permission, action", "Chi tiết từng chức năng"],
            ["ClubsManager", "list, dirty, CRUD, I/O", "Giao diện console"],
        ])

    add_heading(doc, "3.4 Algorithm Design – Thiết kế thuật toán", 2)
    add_para(doc, "Pseudocode – Thêm cầu thủ mới (Function 8):", bold=True)
    add_para(doc,
        "INPUT playerID\n"
        "IF NOT match(playerID, Pxxxx) THEN retry\n"
        "IF findById(playerID) != null THEN print error; RETURN\n"
        "DISPLAY all clubs\n"
        "INPUT clubID\n"
        "IF clubs.findById(clubID) == null THEN print error; RETURN\n"
        "INPUT name, position, shirtNumber\n"
        "IF NOT isPosition(position) THEN retry\n"
        "IF shirtTakenInClub(clubID, shirt) THEN print error; RETURN\n"
        "list.add(new Player(...)); dirty = true\n"
        "PRINT success message")

    add_para(doc, "Pseudocode – Sắp xếp cầu thủ (Function 6):", bold=True)
    add_para(doc,
        "copy = new ArrayList(list)\n"
        "byClubName = Comparator.comparing(p -> clubs.getName(p.clubID), nullsLast)\n"
        "copy.sort(byClubName.thenComparingInt(Player::shirtNumber))\n"
        "DISPLAY copy")

    add_para(doc, "Pseudocode – Load dữ liệu chặt (Function 18):", bold=True)
    add_para(doc,
        "IF NOT clubs.loadStrict(\"clubs.txt\") THEN FAIL\n"
        "IF NOT players.loadStrict(\"players.txt\", clubs) THEN FAIL\n"
        "IF NOT youthPlayers.loadStrict(\"youth_players.txt\", clubs) THEN FAIL\n"
        "PRINT \"Load data successfully!\"")

    # 4. OOP Design
    add_heading(doc, "4. Thiết kế hướng đối tượng (OOP Design)", 1)

    add_heading(doc, "4.1 Use Case Diagram", 2)
    add_para(doc,
        "Use Case Diagram mô tả 19 use case được nhóm theo: Club Management, Player Management, "
        "Youth Player Management và System. Actor chính là User (Manager/Admin/Viewer). "
        "File draw.io: diagrams/UseCase_Diagram.drawio (tham chiếu từ thư mục all/).")
    add_image_placeholder(doc, "Use Case Diagram – mở file diagrams/UseCase_Diagram.drawio và export PNG")

    add_heading(doc, "4.2 Class Diagram", 2)
    add_para(doc,
        "Class Diagram thể hiện đầy đủ các package model, business, dispatcher, tools. "
        "Quan hệ kế thừa: Player, YouthPlayer extends Person. "
        "Quan hệ composition: các Manager quản lý List thực thể tương ứng. "
        "Quan hệ association: Player/YouthPlayer tham chiếu Club qua clubID. "
        "Hệ thống phân quyền: User → Role → Permission.")
    add_image_placeholder(doc, "Class Diagram – export từ diagrams/Class_Diagram.drawio")

    add_heading(doc, "4.3 Sequence Diagram", 2)
    add_para(doc, "Hai sequence diagram chính:")
    add_para(doc, "• Đăng nhập & nạp dữ liệu khởi động: diagrams/Sequence_Diagram_Login_LoadData.drawio")
    add_para(doc, "• Thêm cầu thủ mới: diagrams/Sequence_Diagram_AddPlayer.drawio")
    add_image_placeholder(doc, "Sequence Diagram – Đăng nhập & Load Data")
    add_image_placeholder(doc, "Sequence Diagram – Thêm cầu thủ mới (Add Player)")

    add_heading(doc, "4.4 Bốn nguyên lý OOP trong dự án", 2)
    add_para(doc, "Encapsulation: Các field trong Club, Player, User đều private; truy cập qua getter/setter. "
             "setBudget() chỉ cập nhật khi giá trị > 0.")
    add_para(doc, "Inheritance: Person là lớp abstract; Player và YouthPlayer kế thừa id, name "
             "và override getDisplayInfo().")
    add_para(doc, "Polymorphism: Comparator chaining trong listSortedByClubThenShirt(); "
             "MenuItem dùng Runnable cho action đa hình.")
    add_para(doc, "Abstraction: Person định nghĩa getDisplayInfo() abstract; Validatable tập trung "
             "quy tắc validation.")

    add_heading(doc, "4.5 Cấu trúc package", 2)
    add_table(doc,
        ["Package", "Class", "Vai trò"],
        [
            ["model", "Club, Person, Player, YouthPlayer", "Đại diện dữ liệu nghiệp vụ"],
            ["model", "User, Role, Permission", "Xác thực & phân quyền"],
            ["model", "Validatable", "Interface chứa REGEX & hàm validate"],
            ["business", "ClubsManager, PlayersManager, YouthPlayersManager", "Logic nghiệp vụ & I/O"],
            ["business", "AuthManager", "Đăng nhập, quản lý tài khoản"],
            ["dispatcher", "Menu, MenuItem", "Điều phối menu & luồng chương trình"],
            ["tools", "Inputter", "Nhập liệu console có kiểm tra"],
        ])

    # 5. Implementation
    add_heading(doc, "5. Cấu trúc dự án & Triển khai", 1)
    add_para(doc,
        "Điểm vào chương trình: Menu.main() → start() → auth.login() → reloadAll() → "
        "buildItems() → run(). Menu sử dụng danh sách MenuItem, mỗi item gắn một Permission. "
        "Phương thức allowedItems() lọc menu theo quyền của currentUser.")
    add_para(doc, "Tài khoản mặc định:")
    add_table(doc,
        ["Username", "Password", "Vai trò", "Quyền"],
        [
            ["admin", "admin123", "ADMIN", "Toàn quyền (7 permission)"],
            ["manager", "manager123", "MANAGER", "Xem club; quản lý player & youth; save/load"],
            ["viewer", "viewer123", "VIEWER", "Chỉ xem club, player, youth"],
        ])
    add_image_placeholder(doc, "Cấu trúc thư mục dự án trên IntelliJ IDEA / VS Code")

    # 6. Flowchart
    add_heading(doc, "6. Sơ đồ luồng (Flowchart)", 1)
    add_para(doc,
        "Ba flowchart chính cần trình bày: (1) Luồng chương trình tổng thể từ khởi động đến thoát; "
        "(2) Luồng thêm cầu thủ; (3) Luồng loadStrict – đọc file có validation chặt.")
    add_image_placeholder(doc, "Flowchart – Luồng chương trình tổng thể (Start → Login → Menu → Quit)")
    add_image_placeholder(doc, "Flowchart – Thêm cầu thủ mới (Function 8)")
    add_image_placeholder(doc, "Flowchart – Load dữ liệu chặt (Function 18)")

    # 7. Testing
    add_heading(doc, "7. Kiểm thử & Kết quả chạy chương trình", 1)
    add_heading(doc, "7.1 Kịch bản kiểm thử", 2)
    add_table(doc,
        ["TC", "Mô tả", "Input", "Kết quả mong đợi"],
        [
            ["TC01", "Đăng nhập thành công", "admin/admin123", "Welcome, admin [ADMIN]"],
            ["TC02", "Đăng nhập sai 3 lần", "Sai password x3", "Too many failed attempts. Exiting."],
            ["TC03", "Thêm club trùng ID", "CL-0001", "This club ID already exists!"],
            ["TC04", "Thêm player club không tồn tại", "CL-9999", "This club does not exist!"],
            ["TC05", "Thêm player trùng số áo", "Cùng club, cùng shirt", "This shirt number already exists in this club!"],
            ["TC06", "Viewer không thấy Add club", "viewer login", "Menu không có mục Add a new club"],
            ["TC07", "Load file lỗi", "Sửa 1 dòng sai format", "Load data failed!"],
            ["TC08", "Thoát có thay đổi", "Sửa data, chọn Quit", "Changes detected. Saving data..."],
            ["TC09", "Youth đủ 18 tuổi", "Thêm/update age ≥ 18", "Hiện PROMOTION SUGGESTION"],
            ["TC10", "Xóa youth – hủy", "Chọn N khi confirm", "Deletion cancelled."],
        ])

    add_heading(doc, "7.2 Kết quả chạy chương trình", 2)
    add_image_placeholder(doc, "Màn hình đăng nhập (Login screen)")
    add_image_placeholder(doc, "Menu chính khi đăng nhập ADMIN (đủ 19 chức năng)")
    add_image_placeholder(doc, "Menu khi đăng nhập VIEWER (chỉ chức năng xem)")
    add_image_placeholder(doc, "Chức năng 1 – List all clubs (bảng có header)")
    add_image_placeholder(doc, "Chức năng 6 – List players sorted by club name, shirt number")
    add_image_placeholder(doc, "Chức năng 8 – Add player thành công")
    add_image_placeholder(doc, "Chức năng 8 – Add player bị chặn (shirt number trùng)")
    add_image_placeholder(doc, "Chức năng 16 – List youth players eligible for first team")
    add_image_placeholder(doc, "Chức năng 17 – Save data to files")
    add_image_placeholder(doc, "Chức năng 19 – Quit với auto-save (dirty flag)")

    # 8. Conclusion
    add_heading(doc, "8. Kết luận", 1)
    add_para(doc,
        "Dự án đã xây dựng thành công hệ thống quản lý câu lạc bộ và cầu thủ cho giải EEL "
        "theo đúng yêu cầu LAB211 J1.L.P0036, mở rộng thêm module cầu thủ trẻ và phân quyền "
        "người dùng. Kiến trúc OOP rõ ràng với tách biệt model – business – dispatcher – tools, "
        "validation tập trung qua Validatable, và cơ chế dirty flag đảm bảo không mất dữ liệu khi thoát.")
    add_para(doc, "Hạn chế:", bold=True)
    add_para(doc, "• Mật khẩu lưu plain-text trong bộ nhớ (phù hợp bài lab, chưa mã hóa).")
    add_para(doc, "• Dữ liệu lưu file text, chưa hỗ trợ database hoặc đa người dùng đồng thời.")
    add_para(doc, "• Giao diện console, chưa có GUI.")
    add_para(doc, "Hướng mở rộng:", bold=True)
    add_para(doc, "• Tích hợp MySQL/SQLite; mã hóa mật khẩu (BCrypt).")
    add_para(doc, "• Thêm chức năng promote youth player → player tự động.")
    add_para(doc, "• Xây dựng giao diện JavaFX/Spring Boot REST API.")
    add_para(doc, "Bài học rút ra: Áp dụng Computational Thinking trước khi code giúp giảm lỗi logic; "
             "tập trung validation tại một nguồn (Validatable) giúp bảo trì dễ dàng; "
             "phân quyền qua enum Permission giúp mở rộng vai trò mà không sửa nhiều code.")

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build_report()
